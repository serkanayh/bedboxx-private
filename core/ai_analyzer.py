import json
import re
import anthropic
import logging
import difflib  # Eklendi: email analizi için metin benzerliği hesaplama
from typing import Dict, Any, List, Optional, Union, Tuple
from django.conf import settings
from datetime import datetime, date, timedelta
from bs4 import BeautifulSoup  # HTML işleme için gerekli
import os # os modülünü import et

# --- Model Imports ---
# Import Market and MarketAlias here
from hotels.models import Market, MarketAlias

# --- Attachment Text Extraction Libraries --- 
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from docx import Document
except ImportError:
    Document = None
# --- End Attachment Libraries --- 

logger = logging.getLogger(__name__)

# Sabit prompt tamamen kaldırıldı. Artık veritabanından yüklenecek


class ClaudeAnalyzer:
    """
    A unified class to analyze email content (body or extracted attachment text) 
    using the Claude API.
    Also includes methods for extracting text from attachments.
    """
    
    def __init__(self, api_key=None, prompt=None):
        """
        Initialize the analyzer with API credentials and a system prompt
        """
        self.api_key = api_key or getattr(settings, 'ANTHROPIC_API_KEY', '')
        
        # Eğer prompt parametre olarak verilmezse, veritabanından yükle
        if prompt is not None:
            # Direkt olarak verilen prompt içeriğini kullan
            self.system_prompt = prompt
        else:
            # Veritabanından aktif prompt'u yüklemeye çalış
            try:
                from emails.models import Prompt
                active_prompt = Prompt.objects.filter(active=True).first()
                if active_prompt:
                    self.system_prompt = active_prompt.content
                    logger.info(f"Using active prompt from database: {active_prompt.title}")
                else:
                    logger.error("No active prompt found in database.")
                    self.system_prompt = ""
            except Exception as e:
                logger.error(f"Error loading prompt from database: {str(e)}")
                self.system_prompt = ""
                
        if not self.system_prompt:
            logger.error("No system prompt available. Analysis will likely fail.")
        
        # Configure Anthropic client
        try:
            if not self.api_key:
                raise ValueError("Anthropic API key not found in settings or provided.")
            self.claude_client = anthropic.Anthropic(api_key=self.api_key)
            logger.info("ClaudeAnalyzer initialized.")
        except ImportError:
             logger.error("Anthropic library not installed. pip install anthropic")
             self.claude_client = None
        except Exception as e:
            logger.error(f"Anthropic client failed to initialize: {e}", exc_info=True)
            self.claude_client = None
    
    # --- Text Extraction Methods (Moved from AttachmentAnalyzer) --- 
    def extract_text_from_attachment(self, file_path: str) -> Tuple[str, Optional[str]]:
        """Extracts text from a supported attachment file."""
        if not os.path.exists(file_path):
            return "", f"File not found: {os.path.basename(file_path)}"
        
        file_ext = os.path.splitext(file_path)[1].lower().strip('.')
        
        extraction_methods = {
            'pdf': self._extract_text_pdf,
            'xlsx': self._extract_text_excel,
            'xls': self._extract_text_excel,
            'docx': self._extract_text_word,
            'txt': self._extract_text_text
        }
        
        if file_ext in extraction_methods:
            return extraction_methods[file_ext](file_path)
        else:
            return "", f"Unsupported attachment format: {file_ext}"

    def _extract_text_pdf(self, file_path: str) -> Tuple[str, Optional[str]]:
        """Extracts text from PDF. Returns (text, error_message)."""
        if not PyPDF2:
            return "", 'PDF library (PyPDF2) not installed.'
        text_content = ""
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if reader.is_encrypted:
                     try:
                          reader.decrypt('') # Try decrypting with empty password
                     except Exception as decrypt_err:
                          logger.warning(f"Could not decrypt PDF {os.path.basename(file_path)}: {decrypt_err}")
                          return "", f"Could not decrypt PDF: {os.path.basename(file_path)}"
                          
                for page in reader.pages:
                    try:
                        text_content += page.extract_text() or ""
                    except Exception as page_err: # Handle potential errors during page extraction
                         logger.warning(f"Error extracting text from a page in PDF {os.path.basename(file_path)}: {page_err}")
                         continue # Skip problematic pages
                         
            logger.info(f"Extracted {len(text_content)} chars from PDF: {os.path.basename(file_path)}")
            return text_content, None
        except Exception as e:
            logger.error(f"Error extracting text from PDF {os.path.basename(file_path)}: {e}", exc_info=True)
            return "", f'Error processing PDF: {str(e)}'

    def _extract_text_excel(self, file_path: str) -> Tuple[str, Optional[str]]:
        """Extracts text content from all sheets in an Excel file."""
        if not pd:
            return "", 'Excel library (pandas) not installed.'
        all_text = ""
        try:
            excel_data = pd.read_excel(file_path, sheet_name=None)
            for sheet_name, df in excel_data.items():
                # Convert entire sheet to string, preserving some structure
                sheet_text = df.to_string()
                all_text += f"\n\n--- Sheet: {sheet_name} ---\n{sheet_text}"
            logger.info(f"Extracted {len(all_text)} chars from Excel: {os.path.basename(file_path)}")
            return all_text, None
        except Exception as e:
            logger.error(f"Error extracting text from Excel {os.path.basename(file_path)}: {e}", exc_info=True)
            return "", f'Error processing Excel: {str(e)}'

    def _extract_text_word(self, file_path: str) -> Tuple[str, Optional[str]]:
         """Extracts text from DOCX files."""
         if not file_path.lower().endswith('.docx'):
             return "", '.doc file analysis requires manual conversion to .docx.'
         if not Document:
             return "", 'Word library (python-docx) not installed.'
         text_content = "" # For paragraphs
         all_text = "" # Initialize for table text
         try:
             doc = Document(file_path)
             for para in doc.paragraphs:
                 text_content += para.text + '\n'
                 
             # Optionally extract text from tables too
             if doc.tables:
                 for table in doc.tables:
                      # Ensure all_text exists before appending
                      all_text += "\n--- TABLE START ---\n"
                      try:
                           for row in table.rows:
                                row_text = []
                                for cell in row.cells:
                                     # Extract text from paragraphs within the cell
                                     cell_text = '\n'.join([p.text for p in cell.paragraphs])
                                     row_text.append(cell_text.strip())
                                all_text += " | ".join(row_text) + "\n" # Use pipe separator
                      except Exception as table_err:
                           logger.warning(f"Error processing a table in Word doc {os.path.basename(file_path)}: {table_err}")
                      all_text += "--- TABLE END ---\n"
                         
             # Combine paragraph and table text
             final_text = text_content + all_text 
             logger.info(f"Extracted {len(final_text)} chars from Word: {os.path.basename(file_path)}")
             return final_text, None # Return combined text
             
         except Exception as e:
             logger.error(f"Error extracting text from Word {os.path.basename(file_path)}: {e}", exc_info=True)
             # Check if the error is the specific UnboundLocalError
             if isinstance(e, UnboundLocalError) and 'all_text' in str(e):
                 return "", f"Error processing Word (variable assignment issue): {str(e)}"
             else:
                 return "", f'Error processing Word: {str(e)}'

    def _extract_text_text(self, file_path: str) -> Tuple[str, Optional[str]]:
        """Reads text directly from a TXT file."""
        try:
            # Try common encodings
            encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text_content = None
            for enc in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        text_content = f.read()
                    logger.info(f"Read {len(text_content)} chars from TXT (encoding: {enc}): {os.path.basename(file_path)}")
                    break # Stop if successful
                except UnicodeDecodeError:
                    continue # Try next encoding
                except Exception as read_err: # Catch other potential errors
                     raise read_err # Re-raise other errors
                     
            if text_content is None:
                 raise ValueError("Could not decode TXT file with common encodings.")
                 
            return text_content, None
        except Exception as e:
            logger.error(f"Error reading TXT file {os.path.basename(file_path)}: {e}", exc_info=True)
            return "", f'Error reading TXT file: {str(e)}'
    # --- End Text Extraction Methods --- 
    
    # --- Email Body Cleaning --- 
    def clean_email_body(self, email_html: str, email_text: str) -> str:
        """Cleans email content, preferring HTML if available. 
        Enhanced to preserve table structures and date formats."""
        content_to_clean = email_html if email_html else email_text
        if not content_to_clean:
            return ""
            
        is_html = bool(email_html) and bool(re.search(r'<html|<body|<table|<div|<p>', email_html, re.IGNORECASE))
        
        cleaned_content = ""
        
        if is_html:
            try:
                soup = BeautifulSoup(email_html, 'html.parser')
                
                # Remove script, style, head elements
                for element in soup(["script", "style", "head", "title", "meta", "link"]):
                    element.extract()

                # Define common reply/forward markers
                REPLY_FORWARD_MARKERS_REGEX = r'(from:|sent:|to:|subject:|cc:|bcc:|forwarded message|original message|begin forwarded message|on .* wrote:)'
                
                # Find ALL markers in the email content
                marker_elements = []
                
                # Search text nodes
                for element in soup.find_all(string=re.compile(REPLY_FORWARD_MARKERS_REGEX, re.IGNORECASE | re.DOTALL)):
                    parent = element.parent
                    if parent:
                        marker_elements.append(parent)
                
                if marker_elements:
                    # Find the first marker element
                    first_marker = marker_elements[0]
                    logger.info(f"Found first marker element: {first_marker.name}")
                    
                    # --- IMPORTANT CHANGE: Don't remove tables before the last marker ---
                    # Instead of removing everything after the marker, extract only the content before the marker
                    # but preserve tables with date information
                    
                    # Check if there are any tables in the email
                    tables = soup.find_all('table')
                    tables_with_date = []
                    
                    # Look for tables containing date information (especially 'Tek Gece' or 'One Night')
                    date_patterns = ['Tek Gece', 'One Night', 'Single Day', r'\d{2}\.\d{2}\.\d{4}', r'\d{2}/\d{2}/\d{4}']
                    for table in tables:
                        table_text = table.get_text()
                        if any(re.search(pattern, table_text, re.IGNORECASE) for pattern in date_patterns):
                            tables_with_date.append(table)
                            logger.info(f"Found table with date information: {table_text[:100]}...")
                    
                    # Extract content before the marker
                    content_before_marker = []
                    for element in first_marker.previous_siblings:
                        content_before_marker.append(str(element))
                    content_before_marker.reverse()  # Reverse to get correct order
                    
                    # Add the content from important tables
                    table_content = []
                    for table in tables_with_date:
                        # Only keep tables that are not already in content_before_marker
                        if str(table) not in ''.join(content_before_marker):
                            table_content.append(str(table))
                            logger.info(f"Adding date table to content: {str(table)[:100]}...")
                    
                    # Combine content
                    combined_html = ''.join(content_before_marker) + ''.join(table_content)
                    
                    # Create a new soup from the combined content
                    new_soup = BeautifulSoup(combined_html, 'html.parser')
                    cleaned_content = new_soup.get_text(separator=' ', strip=True)
                else:
                    # No markers found, use the whole content
                    cleaned_content = soup.get_text(separator=' ', strip=True)
                
                # --- SPECIAL HANDLING FOR TABLES WITH DATES ---
                # Add special table annotation for better AI processing
                if tables_with_date:
                    # Extract actual table content in a more structured format for AI
                    structured_tables = []
                    for table in tables_with_date:
                        rows = table.find_all('tr')
                        table_data = []
                        for row in rows:
                            cols = row.find_all(['td', 'th'])
                            if cols:  # Skip empty rows
                                row_data = [col.get_text(strip=True) for col in cols]
                                # Filter out empty cells
                                row_data = [cell for cell in row_data if cell]
                                if row_data:  # Only add non-empty rows
                                    table_data.append(' | '.join(row_data))
                        
                        if table_data:
                            structured_table = "\n--------TABLE:--------\n" + "\n".join(table_data)
                            structured_tables.append(structured_table)
                    
                    # Append structured tables to cleaned content
                    if structured_tables:
                        cleaned_content += "\n\n" + "\n\n".join(structured_tables)
                        logger.info(f"Added {len(structured_tables)} structured tables to content")
                
            except Exception as e:
                logger.error(f"Error cleaning HTML email: {e}", exc_info=True)
                # Fallback to plain text
                cleaned_content = self._clean_plain_text(email_text if email_text else email_html)
        else:
            # Plain text cleaning
             cleaned_content = self._clean_plain_text(content_to_clean)
        
        # Ensure any content with "Tek Gece" or "One Night" is preserved
        special_date_patterns = [
            (r'(\d{2}\.\d{2}\.\d{4})\s*\(Tek Gece\)', r'\1 (Tek Gece)'),
            (r'(\d{2}\.\d{2}\.\d{4})\s*\(One Night\)', r'\1 (One Night)'),
            (r'(\d{2}/\d{2}/\d{4})\s*\(Tek Gece\)', r'\1 (Tek Gece)'),
            (r'(\d{2}/\d{2}/\d{4})\s*\(One Night\)', r'\1 (One Night)'),
        ]
        
        # Search for these patterns in the original content
        special_dates = []
        for pattern, _ in special_date_patterns:
            matches = re.findall(pattern, content_to_clean, re.IGNORECASE)
            special_dates.extend(matches)
        
        # If special dates found but not in cleaned content, explicitly add them
        if special_dates and not any(date in cleaned_content for date in special_dates):
            special_date_section = "\n\nSPECIAL DATES SECTION:\n"
            for date in special_dates:
                special_date_section += f"{date} (Tek Gece / One Night)\n"
            cleaned_content += special_date_section
            logger.info(f"Added special dates section with: {special_dates}")
        
        logger.debug(f"Cleaned email content: {cleaned_content[:500]}...")
        return cleaned_content
    
    def _clean_plain_text(self, text_content: str) -> str:
        """Cleans plain text email content, enhanced to preserve date formats."""
        if not text_content:
             return ""
            
        # Replace multiple newlines with a single newline
        cleaned_text = re.sub(r'\n{3,}', '\n\n', text_content)
        
        # Try to identify and remove quoted content
        quote_markers = [
            r'^>.*$',  # Basic quoting ">"
            r'^On .* wrote:.*$',  # "On DATE, NAME wrote:"
            r'^From:.*$',  # Email headers in quotes
            r'^Sent:.*$',
            r'^To:.*$',
            r'^Subject:.*$'
        ]
        
        # Only apply quote removal if we've identified proper sections
        # Don't blindly remove text matching these patterns from the body
        quote_pattern = '|'.join(f'({m})' for m in quote_markers)
        quote_lines = [line for line in cleaned_text.split('\n') if re.match(quote_pattern, line.strip(), re.IGNORECASE)]
        
        # Only apply quote removal if we have confirmed quote lines
        if quote_lines and len(quote_lines) > 3:  # Heuristic: at least 3 confirmed quote lines
            # Try to split at the first quote marker
            quote_split_pattern = re.compile(f"({quote_pattern})", re.MULTILINE | re.IGNORECASE)
            match = quote_split_pattern.search(cleaned_text)
            if match:
                # Take only the content before the quote
                cleaned_text = cleaned_text[:match.start()]
        
        # Preserve any tables, especially those with dates
        # Look for patterns that might indicate tabular data
        table_markers = [
            "OTEL/HOTEL", "TARİH/DATE", "ODA TİPİ/ROOM TYPE", "PAZARI/MARKET",  # Common headers
            r"\d{2}\.\d{2}\.\d{4}", r"\d{2}/\d{2}/\d{4}",  # Date formats
            "Tek Gece", "One Night"  # Special date indicators
        ]
        
        # Extract potential tables
        lines = cleaned_text.split('\n')
        potential_table_lines = []
        in_potential_table = False
        table_start_index = 0
        
        for i, line in enumerate(lines):
            # Check if line contains table markers
            if any(re.search(marker, line, re.IGNORECASE) for marker in table_markers):
                if not in_potential_table:
                    in_potential_table = True
                    table_start_index = i
            # End of table detection: empty line after table data
            elif in_potential_table and not line.strip():
                # Save this potential table section
                table_section = lines[table_start_index:i]
                if len(table_section) >= 2:  # At least two lines to be considered a table
                    potential_table_lines.extend(table_section)
                in_potential_table = False
        
        # If in_potential_table is still True at the end, include the last section
        if in_potential_table:
            potential_table_lines.extend(lines[table_start_index:])
        
        # Format identified table sections
        if potential_table_lines:
            formatted_tables = []
            current_table = []
            
            for line in potential_table_lines:
                current_table.append(line)
                # New table starts with an empty line
                if not line.strip() and current_table:
                    if len(current_table) > 1:  # Minimum 2 lines for a table
                        formatted_tables.append("\n".join(current_table))
                    current_table = []
            
            # Add any remaining table
            if current_table and len(current_table) > 1:
                formatted_tables.append("\n".join(current_table))
            
            # Create a special table section
            if formatted_tables:
                table_section = "\n\n--------TABLE:--------\n" + "\n".join(formatted_tables)
                
                # Check if we need to add this section (might already be in the text)
                if table_section.strip() not in cleaned_text:
                    cleaned_text += f"\n\n{table_section}"
                    logger.info(f"Added formatted table section from plain text")
        
        # Special handling for date formats with "Tek Gece" or "One Night"
        date_patterns = [
            r'(\d{2}\.\d{2}\.\d{4})\s*\(Tek Gece\s*//?\s*One Night\)',
            r'(\d{2}\.\d{2}\.\d{4})\s*\(Tek Gece\)',
            r'(\d{2}\.\d{2}\.\d{4})\s*\(One Night\)',
            r'(\d{2}/\d{2}/\d{4})\s*\(Tek Gece\)',
            r'(\d{2}/\d{2}/\d{4})\s*\(One Night\)'
        ]
        
        # Check if any special date formats exist in the text
        special_dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, cleaned_text, re.IGNORECASE)
            special_dates.extend(matches)
        
        # If special dates found, add a dedicated section to make sure AI recognizes them
        if special_dates:
            special_date_section = "\n\nSPECIAL DATES SECTION:\n"
            for date in special_dates:
                special_date_section += f"{date} (Tek Gece / One Night)\n"
            
            # Check if this section needs to be added
            if special_date_section not in cleaned_text:
                cleaned_text += special_date_section
                logger.info(f"Added special dates section in plain text: {special_dates}")
        
        return cleaned_text.strip()
        
    # --- Main AI Analysis Method --- 
    def analyze_content(self, text_content: str, context_subject: str = '') -> Dict:
        """
        Analyzes provided text content (cleaned body OR extracted attachment text) using Claude.
        Accepts optional subject for context, but doesn't include it directly in the prompt.
        """
        result_template = {
            'rows': [],
            'raw_ai_response': None,
            'error': None,
            'analysis_source': 'ai' # Mark source as AI
        }
        
        if not self.claude_client:
            logger.error("Claude client not initialized. Cannot analyze content.")
            result_template['error'] = 'AI Analyzer not configured'
            return result_template

        if not text_content:
             logger.warning("No text content provided for analysis.")
             result_template['error'] = 'No content provided'
             return result_template
        
        # Store the context subject for potential fallback use
        self.last_subject = context_subject
             
        # Limit content size before sending to AI
        max_chars = 150000 # Adjust based on typical token limits and safety margin
        if len(text_content) > max_chars:
            logger.warning(f"Content truncated for AI analysis. Original length: {len(text_content)}, truncated to {max_chars}")
            text_content = text_content[:max_chars]
            
        logger.debug(f"Sending content to Claude for analysis (first 500 chars):\n{text_content[:500]}...")
        
        try:
            message = self.claude_client.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=4096,
                temperature=0.1,
                system=self.system_prompt,
                messages=[{"role": "user", "content": text_content}]
            )
            
            raw_response = ""
            if message.content and isinstance(message.content, list):
                for block in message.content:
                    if hasattr(block, 'text'):
                        raw_response += block.text
            
            # Store the raw response globally for the instance to help fallback mechanisms
            self.last_raw_response = raw_response
            
            result_template['raw_ai_response'] = raw_response
            logger.debug(f"Raw response from Claude analysis: {raw_response}")
            
            # Add debug logging for special date formats in raw response
            date_patterns = [
                r'(\d{2}\.\d{2}\.\d{4})\s*\(Tek Gece\s*//?\s*One Night\)',
                r'(\d{2}\.\d{2}\.\d{4})\s*\(Tek Gece\)',
                r'(\d{2}\.\d{2}\.\d{4})\s*\(One Night\)'
            ]
            
            for pattern in date_patterns:
                special_dates = re.findall(pattern, text_content, re.IGNORECASE)
                if special_dates:
                    logger.info(f"Found special date formats in input: {special_dates} (pattern: {pattern})")
                
                matches = re.findall(pattern, raw_response, re.IGNORECASE)
                if matches:
                    logger.info(f"Special date format found in AI response: {matches} (pattern: {pattern})")
            
            logger.info(f"[RAW AI RESPONSE DUMP] Subject: {context_subject}:\n{raw_response[:1000]}")
            
            # Parse JSON response
            parsed_data = self._safe_json_parse(raw_response)
            
            if parsed_data is None:
                logger.error("Failed to parse JSON response from Claude analysis.")
                result_template['error'] = 'Failed to parse AI response'
                return result_template
            
            # Add raw response to each rule for potential fallback extraction
            for rule in parsed_data:
                if isinstance(rule, dict):
                    rule['raw_response'] = raw_response
            
            # Post-process data (normalize dates, validate, etc.)
            processed_rows = self.post_process_ai_rules(parsed_data)
            result_template['rows'] = processed_rows
            
            if not processed_rows:
                 logger.warning("AI analysis returned a valid response, but post-processing yielded no valid rows.")
                 # Keep success=True but rows=[] ? Or mark as error?
                 # Let's consider it a partial success for now, signal might check attachments.
                 pass 
                 
            logger.info(f"AI analysis successful. Extracted {len(processed_rows)} processed rows.")
            return result_template

        except anthropic.APIConnectionError as e:
             logger.error(f"Claude API connection error: {e}", exc_info=True)
             result_template['error'] = f'AI API connection error: {e}'
        except anthropic.RateLimitError as e:
             logger.error(f"Claude API rate limit exceeded: {e}", exc_info=True)
             result_template['error'] = f'AI API rate limit exceeded: {e}'
        except anthropic.APIStatusError as e:
             logger.error(f"Claude API status error: {e.status_code} - {e.response}", exc_info=True)
             result_template['error'] = f'AI API status error: {e.status_code}'
        except Exception as e:
            logger.error(f"An unexpected error occurred during AI content analysis: {e}", exc_info=True)
            result_template['error'] = f'Unexpected AI analysis error: {e}'
            
        return result_template 

    # --- Helper methods for AI response processing --- 
    def _safe_json_parse(self, content: str) -> Optional[List[Dict]]:
        """Claude yanıtından JSON listesini güvenli bir şekilde çıkarmayı dener.
        
        Geliştirilmiş versiyon: Metin içindeki birden fazla JSON objesini bulmaya çalışır.
        """
        logger.debug(f"safe_json_parse input (ilk 500 char): {content[:500]}")
        extracted_rows = []

        if not content:
            return None
        
        # Outer try block for the whole parsing process
        try:
            # Temel temizleme
            content = content.strip()
            if content.startswith("```json"):
                content = content[7:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()

            parsed_successfully = False

            # --- Optimizasyon: Önce geçerli bir liste olarak ayrıştırmayı dene ---
            if content.startswith('['):
                try:
                    parsed_list = json.loads(content)
                    if isinstance(parsed_list, list):
                        logger.debug(f"Doğrudan JSON listesi başarıyla ayrıştırıldı ({len(parsed_list)} öğe).")
                        return parsed_list # Return directly if successful list parse
                    else:
                        # Parsed as JSON, but not a list
                        logger.warning("Doğrudan ayrıştırılan JSON bir liste değil, sonraki adımlara geçilecek.")
                except json.JSONDecodeError as e:
                    # Failed to parse as list
                    logger.warning(f"Doğrudan liste JSON ayrıştırma hatası: {e}. Sonraki adımlara geçilecek.")
            
            # --- Optimizasyon: Tek bir obje olarak ayrıştırmayı dene ---
            # Only try this if parsing as list failed or wasn't attempted
            # Use elif to connect it logically to the list check, but ensure it's at the correct indentation
            elif content.startswith('{') and content.endswith('}'): # Correct indentation for elif
                try:
                    single_obj = json.loads(content)
                    if isinstance(single_obj, dict):
                        logger.info("Tek JSON objesi bulundu ve liste içine alındı.")
                        return [single_obj] # Return directly if successful single object parse
                except json.JSONDecodeError as e:
                    # Failed to parse as single object
                    logger.warning(f"Tek obje JSON ayrıştırma hatası: {e}. Sonraki adımlara geçilecek.")

            # --- Ana Mantık: Metin içindeki objeleri non-greedy regex ile bul ve ayrıştır ---
            # This runs if direct list/object parsing failed or wasn't applicable
            # No need for 'if not parsed_successfully:' check here, as the code flow handles it
            
            json_obj_pattern = r'({.*?})'
            potential_matches = re.findall(json_obj_pattern, content, re.DOTALL)
            logger.debug(f"Potansiyel {len(potential_matches)} JSON objesi non-greedy regex ile bulundu.")

            match_index = 0
            for match_str in potential_matches:
                match_index += 1
                match_str = match_str.strip()
                if len(match_str) < 10 or ':' not in match_str:
                    logger.debug(f"---> [Parse Debug] Skipping potential match #{match_index} (too short or no colon): {match_str[:100]}...")
                    continue
                
                logger.debug(f"---> [Parse Debug] Attempting to parse potential match #{match_index}: {match_str[:200]}...")
                try: # Inner try for parsing each regex match
                    parsed_obj = json.loads(match_str)
                    if isinstance(parsed_obj, dict):
                        logger.info(f"---> [Parse Debug] Successfully parsed potential match #{match_index}. Keys: {list(parsed_obj.keys())}")
                        extracted_rows.append(parsed_obj)
                    else: # Corrected indentation for this else, belongs to inner if
                        logger.warning(f"---> [Parse Debug] Parsed structure for match #{match_index} is not a dict: {type(parsed_obj)}")
                # Except blocks for the inner try
                except json.JSONDecodeError as e: 
                    logger.error(f"---> [Parse Debug] FAILED to parse potential match #{match_index}. Error: {e}. Object (start): {match_str[:200]}...")
                except Exception as e: # Catch other errors during individual parsing
                    logger.error(f"---> [Parse Debug] Unexpected error processing potential match #{match_index}: {e}. Object (start): {match_str[:200]}...", exc_info=True)

            if extracted_rows:
                logger.info(f"Metinden {len(extracted_rows)} JSON objesi başarıyla ayıklandı.")
                return extracted_rows

            # If we reach here, no valid JSON was parsed by any method
            logger.warning("Metin içinden geçerli JSON objesi ayıklanamadı (tüm yöntemler denendi).")
            return None

        # Catch errors related to the outer try block (e.g., initial stripping)
        except Exception as e:
            logger.error(f"_safe_json_parse içinde beklenmeyen genel hata: {e}", exc_info=True)
            return None

    def post_process_ai_rules(self, rules: List[Dict]) -> List[Dict]:
        """Validates, normalizes rules extracted by AI, and resolves market names."""
        processed = []
        current_year = datetime.now().year

        # Pre-fetch all Market objects for efficiency
        all_markets_by_name = {m.name.strip().upper(): m for m in Market.objects.all()}
        # We will query aliases dynamically inside the loop if needed

        for rule_index, rule in enumerate(rules):
            if not isinstance(rule, dict):
                logger.warning(f"Skipping non-dict item #{rule_index} in AI rules: {rule}")
                continue

            # --- Field Extraction and Validation ---
            hotel_name = rule.get('hotel_name')
            start_date_str = rule.get('start_date') # Expect YYYY-MM-DD
            end_date_str = rule.get('end_date')     # Expect YYYY-MM-DD
            sale_status = rule.get('sale_status') # Expect 'stop' or 'open'
            ai_markets_list = rule.get('markets', ['ALL']) # Expect list or ['ALL']
            
            # Special handling for 'one night' cases where AI returned null dates
            if ((start_date_str is None or end_date_str is None) and 
                rule.get('one_night_date') or rule.get('tek_gece_date') or rule.get('single_night_date')):
                # Try to use the special date field
                special_date = rule.get('one_night_date') or rule.get('tek_gece_date') or rule.get('single_night_date')
                if special_date:
                    logger.info(f"Found special 'one night' date field: {special_date}")
                    start_date_str = special_date
                    end_date_str = special_date

            # Fallback: If both dates are still null, check raw data for date patterns
            if not start_date_str or not end_date_str:
                # Look for "Tek Gece" patterns in raw response
                raw_response = rule.get('raw_response', '')
                if not raw_response and hasattr(self, 'last_raw_response'):
                    raw_response = self.last_raw_response
                
                if raw_response:
                    date_patterns = [
                        r'(\d{2}\.\d{2}\.\d{4})\s*\(Tek Gece\s*//?\s*One Night\)',
                        r'(\d{2}\.\d{2}\.\d{4})\s*\(Tek Gece\)',
                        r'(\d{2}\.\d{2}\.\d{4})\s*\(One Night\)',
                        r'(\d{1,2}\.\d{1,2}\.\d{4}).*?(tek gece|one night)',
                        r'(\d{1,2}/\d{1,2}/\d{4}).*?(tek gece|one night)'
                    ]
                    
                    for pattern in date_patterns:
                        matches = re.findall(pattern, raw_response, re.IGNORECASE)
                        if matches:
                            date_str = matches[0][0] if isinstance(matches[0], tuple) else matches[0]
                            logger.info(f"Found date in raw response using pattern '{pattern}': {date_str}")
                            
                            # Convert date to YYYY-MM-DD format
                            try:
                                if '.' in date_str:
                                    # DD.MM.YYYY format
                                    parts = date_str.split('.')
                                    if len(parts) == 3:
                                        normalized_date = f"{parts[2]}-{parts[1]}-{parts[0]}"
                                        start_date_str = normalized_date
                                        end_date_str = normalized_date
                                        logger.info(f"Normalized date to: {normalized_date}")
                                elif '/' in date_str:
                                    # DD/MM/YYYY format
                                    parts = date_str.split('/')
                                    if len(parts) == 3:
                                        normalized_date = f"{parts[2]}-{parts[1]}-{parts[0]}"
                                        start_date_str = normalized_date
                                        end_date_str = normalized_date
                                        logger.info(f"Normalized date to: {normalized_date}")
                            except Exception as e:
                                logger.error(f"Error normalizing date '{date_str}': {e}")
                            
                            # If we found a date, break out of the loop
                            if start_date_str and end_date_str:
                                break

            # Check for date in email subject
            if (not start_date_str or not end_date_str) and hasattr(self, 'last_subject'):
                subject = self.last_subject
                date_patterns = [
                    r'(\d{2}\.\d{2}\.\d{4})',
                    r'(\d{2}/\d{2}/\d{4})'
                ]
                
                for pattern in date_patterns:
                    matches = re.findall(pattern, subject)
                    if matches:
                        date_str = matches[0]
                        logger.info(f"Found date in subject: {date_str}")
                        
                        # Convert date to YYYY-MM-DD format
                        try:
                            if '.' in date_str:
                                # DD.MM.YYYY format
                                parts = date_str.split('.')
                                if len(parts) == 3:
                                    normalized_date = f"{parts[2]}-{parts[1]}-{parts[0]}"
                                    start_date_str = normalized_date
                                    end_date_str = normalized_date
                                    logger.info(f"Normalized date from subject to: {normalized_date}")
                            elif '/' in date_str:
                                # DD/MM/YYYY format
                                parts = date_str.split('/')
                                if len(parts) == 3:
                                    normalized_date = f"{parts[2]}-{parts[1]}-{parts[0]}"
                                    start_date_str = normalized_date
                                    end_date_str = normalized_date
                                    logger.info(f"Normalized date from subject to: {normalized_date}")
                        except Exception as e:
                            logger.error(f"Error normalizing date from subject '{date_str}': {e}")
                        
                        if start_date_str and end_date_str:
                            break

            if not all([hotel_name, start_date_str, end_date_str, sale_status]):
                logger.warning(f"Skipping rule #{rule_index} due to missing required fields (hotel, dates, status): {rule}")
                continue

            # --- Date Validation ---
            start_date = self._normalize_date(start_date_str, current_year)
            end_date = self._normalize_date(end_date_str, current_year)

            if not start_date or not end_date:
                logger.warning(f"Skipping rule #{rule_index} due to invalid date format after normalization: {rule}")
                continue

            # Ensure start <= end
            if start_date > end_date:
                logger.warning(f"Swapping start/end dates for rule #{rule_index}: {start_date} > {end_date}. Rule: {rule}")
                start_date, end_date = end_date, start_date

            rule['start_date'] = start_date.strftime('%Y-%m-%d')
            rule['end_date'] = end_date.strftime('%Y-%m-%d')

            # --- Room Type Normalization ---
            room_type = rule.get('room_type', 'All Room')
            if not room_type or str(room_type).strip().upper() in ['ALL ROOM TYPES', 'ALL ROOMS', 'TÜM ODA TIPLERI', '-']:
                room_type = 'All Room'
            rule['room_type'] = str(room_type).strip() # Ensure it's a string and stripped

            # --- Market Resolution & Normalization ---
            resolved_markets = set() # Stores resolved *canonical market names* (uppercase)
            if not isinstance(ai_markets_list, list):
                ai_markets_list = ['ALL'] # Treat non-lists as default

            for market_name in ai_markets_list:
                market_name_upper = str(market_name).strip().upper()
                if not market_name_upper:
                    continue

                if market_name_upper == "ALL":
                    resolved_markets.add("ALL")
                    # If 'ALL' is present, we can potentially break early,
                    # unless we need to log specific resolved markets alongside 'ALL'.
                    # For now, let's continue processing others for logging clarity.
                    continue

                # 1. Check direct Market name match
                matched_market = all_markets_by_name.get(market_name_upper)
                if matched_market:
                    resolved_markets.add(matched_market.name.strip().upper()) # Add the canonical name
                    logger.debug(f"Rule #{rule_index}: Market '{market_name}' matched directly.")
                    continue

                # 2. Check MarketAlias match (Query dynamically)
                try:
                    # Use filter with __iexact for case-insensitive alias matching
                    alias_obj = MarketAlias.objects.prefetch_related('markets').filter(alias__iexact=market_name_upper).first()
                    if alias_obj and alias_obj.markets.exists():
                        logger.debug(f"Rule #{rule_index}: Alias '{market_name}' found. Resolving associated markets.")
                        markets_added_from_alias = set()
                        for related_market in alias_obj.markets.all():
                            canonical_name = related_market.name.strip().upper()
                            resolved_markets.add(canonical_name)
                            markets_added_from_alias.add(canonical_name)
                        logger.debug(f"Rule #{rule_index}: Alias '{market_name}' resolved to -> {markets_added_from_alias}")
                        continue # Move to next market_name in ai_markets_list
                    else:
                         # If no alias found via iexact, proceed to step 3
                         pass
                except Exception as alias_lookup_err:
                     logger.error(f"Rule #{rule_index}: Error looking up alias '{market_name}': {alias_lookup_err}", exc_info=True)
                     # Continue to step 3 if alias lookup fails

                # 3. No direct market or alias match found
                logger.warning(f"Rule #{rule_index}: Market/Alias '{market_name}' not found in DB. Adding as is (uppercase).")
                resolved_markets.add(market_name_upper) # Add the original name (uppercase) if not found

            # Final market list handling
            final_markets = list(resolved_markets)
            if not final_markets:
                final_markets = ['ALL'] # Default to ALL if nothing resolved
            elif "ALL" in final_markets and len(final_markets) > 1:
                final_markets = ['ALL'] # If ALL is present with others, simplify to just ALL

            rule['markets'] = final_markets # Keep as list of canonical names
            rule['market'] = ", ".join(final_markets) # Comma-separated string of canonical names

            # --- Sale Status Validation & Standardization ---
            action = str(sale_status).lower().strip().replace('_sale', '') # Normalize
            if action not in ['stop', 'open']:
                logger.warning(f"Invalid sale_status '{sale_status}' in rule #{rule_index}, defaulting to stop. Rule: {rule}")
                action = 'stop'
            rule['sale_status'] = action # Use normalized 'stop' or 'open'
            rule['sale_type'] = action # For compatibility

            processed.append(rule)

        logger.info(f"Post-processed {len(processed)} rules from AI.")
        return processed

    def _normalize_date(self, date_str: Optional[str], current_year: int) -> Optional[date]:
        """Tries to parse date string into a date object, assuming current year if needed."""
        if not date_str or not isinstance(date_str, str):
            return None
    
        cleaned_date_str = date_str.strip()
        if not cleaned_date_str or cleaned_date_str.upper() in ["YYYY-MM-DD", "NULL"]:
            return None
            
        formats_to_try = [
            "%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d",
            "%d-%m-%Y", "%d/%m/%Y", "%d.%m.%Y",
            "%m-%d-%Y", "%m/%d/%Y", "%m.%d.%Y",
            # Formats potentially missing year
            "%d %b", "%d %B", # 15 May, 15 Mayis
            "%b %d", "%B %d", # May 15, Mayis 15
            "%d-%b", "%d-%B",
            "%b-%d", "%B-%d",
            "%d.%m", "%d/%m", "%d-%m", # 15.05, 15/05, 15-05
            "%m.%d", "%m/%d", "%m-%d"  # 05.15, 05/15, 05-15
        ]
        
        parsed_date = None
        for fmt in formats_to_try:
            try:
                parsed_date = datetime.strptime(cleaned_date_str, fmt).date()
                # If format doesn't include year, add current year
                if '%Y' not in fmt and '%y' not in fmt:
                     parsed_date = parsed_date.replace(year=current_year)
                break # Success
            except (ValueError, TypeError):
                continue # Try next format
                
        # Handle Turkish month names if standard parsing failed
        if not parsed_date:
             month_tr_pattern = r'(\d{1,2})\s+((?:Ocak|Şubat|Mart|Nisan|Mayıs|Haziran|Temmuz|Ağustos|Eylül|Ekim|Kasım|Aralık)[a-zA-ZğüşöçıİĞÜŞÖÇ]*)' 
             match = re.search(month_tr_pattern, cleaned_date_str, re.IGNORECASE)
             if match:
                 day = int(match.group(1))
                 month_name_tr = match.group(2).lower()
                 month_map_tr = {
                      'ocak': 1, 'şubat': 2, 'mart': 3, 'nisan': 4, 'mayıs': 5, 'haziran': 6,
                      'temmuz': 7, 'ağustos': 8, 'eylül': 9, 'ekim': 10, 'kasım': 11, 'aralık': 12
                 }
                 month = month_map_tr.get(month_name_tr.split()[0]) # Get first word if needed
                 if month and 1 <= day <= 31:
                      try:
                           parsed_date = date(current_year, month, day)
                      except ValueError: # Handle invalid day for month (e.g., 31 Şubat)
                           logger.warning(f"Invalid day/month combination from Turkish name: {day} {month_name_tr}")
                           parsed_date = None # Mark as invalid

        if not parsed_date:
             logger.warning(f"Could not normalize date string: '{date_str}'")
             
        return parsed_date
        
    def _adjust_date_order(self, start_date: Optional[date], end_date: Optional[date]) -> tuple[Optional[date], Optional[date]]:
        """Tarih sırasını kontrol eder, gerekirse değiştirir."""
        if start_date and end_date and start_date > end_date:
            logger.warning(f"Tarihler değiştiriliyor: {start_date}, {end_date}")
            return end_date, start_date
        return start_date, end_date
    
    def _extract_key_value(self, text: str, regex: str, default: Optional[str] = None) -> Optional[str]:
        """Metin içinden regex pattern ile değer çıkarma."""
        if match := re.search(regex, text, re.IGNORECASE):
            return match.group(1).strip()
        return default
    
    def _is_date_valid(self, date_str: str) -> bool:
        """Tarih string'inin geçerli bir format olup olmadığını kontrol eder."""
        try:
            datetime.strptime(date_str, "%Y-%m-%d")
            return True
        except ValueError:
            return False
    
    def _parse_date_range(self, date_range_str: str) -> tuple[Optional[date], Optional[date]]:
        """
        '10-12 MARCH', '10 - 15 APRIL' gibi tarih aralıklarını ayrıştırır.
        
        Args:
            date_range_str: Tarih aralığı string'i
            
        Returns:
            (start_date, end_date) tuple olarak iki tarih. Ayrıştırılamazsa (None, None).
        """
        if not date_range_str or not isinstance(date_range_str, str):
            return None, None
            
        date_range_str = date_range_str.strip().upper()
        
        # Month adı ve iki sayı (gün aralığı) içeren patternler için
        month_pattern = r'(\d{1,2})\s*[-\/to]+\s*(\d{1,2})\s+((?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC|JANUARY|FEBRUARY|MARCH|APRIL|JUNE|JULY|AUGUST|SEPTEMBER|OCTOBER|NOVEMBER|DECEMBER|OCAK|ŞUBAT|MART|NİSAN|MAYIS|HAZİRAN|TEMMUZ|AĞUSTOS|EYLÜL|EKİM|KASIM|ARALIK)[A-Za-z]*)'
        
        month_match = re.search(month_pattern, date_range_str)
        if month_match:
            start_day = int(month_match.group(1))
            end_day = int(month_match.group(2))
            month_name = month_match.group(3)
            
            # Ay adları sözlüğü 
            month_names = {
                "JANUARY": 1, "JAN": 1, "OCAK": 1,
                "FEBRUARY": 2, "FEB": 2, "ŞUBAT": 2, "SUBAT": 2,
                "MARCH": 3, "MAR": 3, "MART": 3,
                "APRIL": 4, "APR": 4, "NİSAN": 4, "NISAN": 4,
                "MAY": 5, "MAYIS": 5,
                "JUNE": 6, "JUN": 6, "HAZİRAN": 6, "HAZIRAN": 6,
                "JULY": 7, "JUL": 7, "TEMMUZ": 7,
                "AUGUST": 8, "AUG": 8, "AĞUSTOS": 8, "AGUSTOS": 8,
                "SEPTEMBER": 9, "SEP": 9, "EYLÜL": 9, "EYLUL": 9,
                "OCTOBER": 10, "OCT": 10, "EKİM": 10, "EKIM": 10,
                "NOVEMBER": 11, "NOV": 11, "KASIM": 11,
                "DECEMBER": 12, "DEC": 12, "ARALIK": 12
            }
            
            # Ay adını standardize et
            month_num = None
            for name, num in month_names.items():
                if month_name.startswith(name):
                    month_num = num
                    break
            
            if month_num and 1 <= start_day <= 31 and 1 <= end_day <= 31:
                # Geçerli yılı kullan
                current_year = datetime.now().year
                
                try:
                    start_date = date(current_year, month_num, start_day)
                    end_date = date(current_year, month_num, end_day)
                    
                    # end_date < start_date ise muhtemelen yanlış tarih sıralaması var
                    if end_date < start_date:
                        logger.warning(f"Tarih sıralaması düzeltiliyor: {start_day}-{end_day} {month_name}")
                        start_date, end_date = end_date, start_date
                        
                    return start_date, end_date
                except ValueError as e:
                    logger.warning(f"Geçersiz tarih aralığı: {start_day}-{end_day} {month_name}, hata: {e}")
        
        # Eğer standart pattern ile eşleşme bulunmazsa, iki tarihli formatlara bak
        # Örn: "10 MARCH - 15 APRIL 2023"
        two_dates_pattern = r'(\d{1,2})\s+((?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC|JANUARY|FEBRUARY|MARCH|APRIL|JUNE|JULY|AUGUST|SEPTEMBER|OCTOBER|NOVEMBER|DECEMBER|OCAK|ŞUBAT|MART|NİSAN|MAYIS|HAZİRAN|TEMMUZ|AĞUSTOS|EYLÜL|EKİM|KASIM|ARALIK)[A-Za-z]*)\s*[-\/to]+\s*(\d{1,2})\s+((?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC|JANUARY|FEBRUARY|MARCH|APRIL|JUNE|JULY|AUGUST|SEPTEMBER|OCTOBER|NOVEMBER|DECEMBER|OCAK|ŞUBAT|MART|NİSAN|MAYIS|HAZİRAN|TEMMUZ|AĞUSTOS|EYLÜL|EKİM|KASIM|ARALIK)[A-Za-z]*)'
        
        two_dates_match = re.search(two_dates_pattern, date_range_str)
        if two_dates_match:
            start_day = int(two_dates_match.group(1))
            start_month_name = two_dates_match.group(2)
            end_day = int(two_dates_match.group(3))
            end_month_name = two_dates_match.group(4)
            
            # Ay adları sözlüğü 
            month_names = {
                "JANUARY": 1, "JAN": 1, "OCAK": 1,
                "FEBRUARY": 2, "FEB": 2, "ŞUBAT": 2, "SUBAT": 2,
                "MARCH": 3, "MAR": 3, "MART": 3,
                "APRIL": 4, "APR": 4, "NİSAN": 4, "NISAN": 4,
                "MAY": 5, "MAYIS": 5,
                "JUNE": 6, "JUN": 6, "HAZİRAN": 6, "HAZIRAN": 6,
                "JULY": 7, "JUL": 7, "TEMMUZ": 7,
                "AUGUST": 8, "AUG": 8, "AĞUSTOS": 8, "AGUSTOS": 8,
                "SEPTEMBER": 9, "SEP": 9, "EYLÜL": 9, "EYLUL": 9,
                "OCTOBER": 10, "OCT": 10, "EKİM": 10, "EKIM": 10,
                "NOVEMBER": 11, "NOV": 11, "KASIM": 11,
                "DECEMBER": 12, "DEC": 12, "ARALIK": 12
            }
            
            # Ay adlarını standardize et
            start_month_num = None
            end_month_num = None
            
            for name, num in month_names.items():
                if start_month_name.startswith(name):
                    start_month_num = num
                if end_month_name.startswith(name):
                    end_month_num = num
            
            if start_month_num and end_month_num and 1 <= start_day <= 31 and 1 <= end_day <= 31:
                # Geçerli yılı kullan
                current_year = datetime.now().year
                
                try:
                    start_date = date(current_year, start_month_num, start_day)
                    end_date = date(current_year, end_month_num, end_day)
                    
                    # Eğer bitiş ayı başlangıç ayından önce ise, muhtemelen yıllar farklı
                    # Örn: Aralık 2023 - Ocak 2024
                    if end_month_num < start_month_num:
                        end_date = date(current_year + 1, end_month_num, end_day)
                        
                    return start_date, end_date
                except ValueError as e:
                    logger.warning(f"Geçersiz tarih aralığı: {start_day} {start_month_name} - {end_day} {end_month_name}, hata: {e}")
        
        # Hiçbir durumda eşleşme bulunamazsa
        return None, None
    
    def post_process_data(self, data: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """Ayrıştırılmış veriyi işleyerek normalleştirir ve filtreler."""
        if not data or not isinstance(data, list):
            return {"rows": []}
            
        processed_data = {"rows": []}
        for rule in data:
            if not isinstance(rule, dict):
                continue
                
            # Değerleri çıkar
            hotel_name = rule.get("hotel_name")
            room_type_mail_raw = rule.get("room_type") or rule.get("room_types", [])
            
            # Liste formatındaki oda tiplerini işle
            if isinstance(room_type_mail_raw, list):
                room_type_mail = ', '.join([str(rt) for rt in room_type_mail_raw if rt]) or "All Room Types"
            else:
                room_type_mail = str(room_type_mail_raw or "All Room Types").strip()
            if not room_type_mail:
                room_type_mail = "All Room Types"
            
            # Markets değerlerini işle
            markets_list = rule.get("markets", ["ALL"])
            if isinstance(markets_list, list):
                markets = ', '.join(m for m in markets_list if m and isinstance(m, str)).strip()
            else:
                markets = str(markets_list or 'ALL').strip()
            if not markets:
                markets = "ALL"
            
            # Satış durumunu işle
            sale_status_raw = str(rule.get("sale_status") or 'stop').lower().strip()
            sale_type = sale_status_raw if sale_status_raw in ['stop', 'open'] else 'stop'
            
            # Tarihleri işle
            start_dt = None
            end_dt = None
            
            # Eğer date_range (tarih aralığı) bilgisi doğrudan verilmişse
            date_range_str = rule.get("date_range")
            if date_range_str:
                start_dt, end_dt = self._parse_date_range(date_range_str)
            
            # Tarihi date_range'den alamazsak normal start_date/end_date'den almayı dene
            if not (start_dt and end_dt):
                start_dt = self._normalize_date(rule.get("start_date"))
                end_dt = self._normalize_date(rule.get("end_date"))
                start_dt, end_dt = self._adjust_date_order(start_dt, end_dt)
            
            # Tarihler geçerliyse kayıt oluştur
            if start_dt and end_dt:
                processed_data["rows"].append({
                    'hotel_name': hotel_name,
                    'room_type': room_type_mail,
                    'market': markets,
                    'start_date': start_dt.strftime('%d.%m.%Y'),
                    'end_date': end_dt.strftime('%d.%m.%Y'),
                    'sale_type': sale_type
                })
            else:
                logger.warning(f"Kural atlandı: Geçersiz/eksik tarihler - {rule.get('start_date')}, {rule.get('end_date')}")
                
        return processed_data
    
    def clean_email_content(self, email_content: str) -> str:
        """
        Clean and preprocess email content to reduce token usage while preserving structure
        
        Args:
            email_content: Raw email content
            
        Returns:
            Cleaned email content with preserved structure
        """
        if not email_content:
            return ""
        
        cleaned_content = "" # Initialize cleaned_content
        
        # Check if content is HTML
        is_html = bool(re.search(r'<html|<body|<table|<div|<p>', email_content, re.IGNORECASE))
        
        html_processed_successfully = False
        if is_html:
            try:
                # Try importing BeautifulSoup here, closer to where it's used
                from bs4 import BeautifulSoup 
                
                # Parse HTML content
                soup = BeautifulSoup(email_content, 'html.parser')
                
                # ... (rest of the HTML cleaning logic using soup) ...
                # (Assuming the existing logic inside this try block is correct)

                # Get text with some basic structure preservation
                lines = soup.get_text().split('\\n')
                cleaned_lines = []
                
                for line in lines:
                    line = line.strip()
                    if line:
                        cleaned_lines.append(line)
                
                # Join with newlines to preserve paragraph structure
                cleaned_content = '\\n\\n'.join(cleaned_lines)
                
                html_processed_successfully = True # Mark success
                
            except ImportError: # This except belongs to the HTML try block
                logger.warning("BeautifulSoup not installed. Basic HTML cleaning will be used.")
                # html_processed_successfully remains False
            except Exception as e: # This except also belongs to the HTML try block
                logger.error(f"Error processing HTML content: {e}", exc_info=True)
                # html_processed_successfully remains False
        
        # If HTML processing failed or it wasn't HTML, use basic cleaning
        if not html_processed_successfully:
            logger.debug("Performing basic cleaning for plain text or as fallback.")
            # Basic cleaning for plain text 
            patterns_to_remove = [
                r'From:.*?(?=\\n\\n|\\Z)', 
                r'Sent:.*?(?=\\n\\n|\\Z)',
                r'To:.*?(?=\\n\\n|\\Z)',
                r'Subject:.*?(?=\\n\\n|\\Z)',
                r'Cc:.*?(?=\\n\\n|\\Z)',
                r'This email and any files.*?$',
                r'DISCLAIMER.*?$',
                r'Confidentiality Notice.*?$',
                r'[^s]+@[^s]+\\.[^s]+', # Remove most email addresses - corrected regex escaping
                r'(?:\\+\\d{1,3}|\\b)\\d{3}[-.\\s]?\\d{3}[-.\\s]?\\d{4}\\b' # Remove phone numbers - corrected regex escaping
            ]
            
            temp_cleaned_content = email_content # Start with original if HTML failed
            for pattern in patterns_to_remove:
                temp_cleaned_content = re.sub(pattern, '', temp_cleaned_content, flags=re.IGNORECASE | re.DOTALL)
            
            cleaned_content = temp_cleaned_content # Assign result to cleaned_content
        
        # Final whitespace cleanup (applied to both HTML extracted text and plain text)
        cleaned_content = re.sub(r'\\n{3,}', '\\n\\n', cleaned_content) # Limit consecutive newlines
        cleaned_content = re.sub(r'\\s{2,}', ' ', cleaned_content) # Replace multiple spaces/tabs with one
        
        return cleaned_content.strip()
        
    def analyze_email_content(self, email_content, email_subject=None, email_html=None):
        """
        Analyzes the provided email content with the AI model.
        This function handles all AI model specific logic.
        
        Args:
            email_content (str): The raw text content of the email to analyze  
            email_subject (str, optional): The email subject for additional context
            email_html (str, optional): The HTML version of the email if available
        
        Returns:
            dict: Analysis result containing extracted data or error
                {
                    'rows': [{'hotel_name': '...', 'room_type': '...', 'start_date': '...', 'end_date': '...', 'markets': ['...', '...']}, {...}],
                    'error': str or None,
                    'should_analyze_attachments': bool  # Indicates if attachments should be analyzed
                }
        """
        if not email_content:
            return {'error': 'No email content provided for analysis', 'rows': [], 'raw_ai_response': None, 'has_sale_info': False, 'should_analyze_attachments': True}
        
        # Check if email references attachments
        has_attachment_references = self.check_for_attachment_references(email_content)
        
        # Determine if we should extract stop sale info from content or suggest checking attachments
        # 1. If content has stop sale info directly, we'll try to extract it
        # 2. If content refers to attachments AND doesn't have clear stop sale info in the body,
        #    we'll suggest checking the attachments instead
        
        prompt_text = self.system_prompt
        
        try:
            # Use Claude API to analyze the content
            logger.debug(f"Calling Claude API with text length: {len(email_content)}")
            
            # Add the subject if available for better context
            content_to_analyze = email_content
            if email_subject:
                content_to_analyze = f"Subject: {email_subject}\n\n{email_content}"
            
            # Call the AI API for analysis
            response = self.claude_client.messages.create(
                model="claude-3-haiku-20240307",
                max_tokens=4096,
                temperature=0.1,
                system=prompt_text,
                messages=[
                    {"role": "user", "content": content_to_analyze}
                ]
            )
            
            # Get the raw response text
            raw_response = response.content[0].text
            
            # Try to parse the JSON from the response
            try:
                # Extract JSON part from the response
                json_match = re.search(r'```json\s*(.*?)\s*```', raw_response, re.DOTALL)
                
                if json_match:
                    json_str = json_match.group(1)
                    ai_result = json.loads(json_str)
                else:
                    # No JSON block found, try to extract any JSON
                    json_match = re.search(r'\{\s*"rows"\s*:', raw_response)
                    if json_match:
                        # Extract from the beginning of the match to the end
                        json_str = raw_response[json_match.start():]
                        # Attempt to parse, may fail if there's text after the JSON
                        ai_result = json.loads(json_str)
                    else:
                        # Fallback: Try to parse the entire response as JSON
                        ai_result = json.loads(raw_response)
                
                # Check if we need to also analyze attachments
                should_analyze_attachments = has_attachment_references
                has_sale_info = False
                
                # Check if any rows were extracted
                if 'rows' in ai_result and len(ai_result['rows']) > 0:
                    # Found data in the body, mark as having sale info
                    has_sale_info = True
                    
                    # Extract stop sale dates from body
                    sale_dates = []
                    for row in ai_result['rows']:
                        if 'start_date' in row and 'end_date' in row:
                            try:
                                # Check if dates are in future - generally indicates stop sale info
                                start_date = datetime.strptime(row['start_date'], '%Y-%m-%d').date()
                                today = datetime.now().date()
                                if start_date >= today:
                                    sale_dates.append((row['start_date'], row['end_date']))
                            except (ValueError, TypeError):
                                # Date parsing failed, just count it anyway
                                sale_dates.append((row.get('start_date'), row.get('end_date')))
                    
                    # If found concrete stop sale dates, we may not need to check attachments,
                    # even if there are attachment references
                    if len(sale_dates) > 0:
                        # If clear data in body, only check attachments if specifically suggested
                        # by certain phrases that indicate important info is ONLY in the attachment
                        should_analyze_attachments = (
                            has_attachment_references and 
                            any(phrase in email_content.lower() for phrase in [
                                'ayrıntılar ekte', 'detaylar ekte', 'details in attachment',
                                'ekteki tarihlerin', 'tarihleri ekte', 'dates in attachment'
                            ])
                        )
                else:
                    # No rows found in body, check attachments if references exist
                    should_analyze_attachments = has_attachment_references
                
                return {
                    'rows': ai_result.get('rows', []),
                    'error': None,
                    'raw_ai_response': raw_response,
                    'has_sale_info': has_sale_info,
                    'should_analyze_attachments': should_analyze_attachments
                }
                
            except json.JSONDecodeError as json_err:
                logger.error(f"JSON parsing error: {json_err}. Raw response: {raw_response[:200]}...")
                return {'error': f'Failed to parse AI response: {json_err}', 'rows': [], 'raw_ai_response': raw_response, 'has_sale_info': False, 'should_analyze_attachments': has_attachment_references}
            
        except Exception as e:
            logger.error(f"An unexpected error occurred during AI analysis: {e}", exc_info=True)
            return {'error': f'Unexpected AI analysis error: {e}', 'rows': [], 'raw_ai_response': None, 'has_sale_info': False, 'should_analyze_attachments': has_attachment_references}

    def check_for_attachment_references(self, text):
        """
        E-posta içeriğinden eklere referans olup olmadığını kontrol eder.
        Örneğin: "ekte belirtilen", "ekte gönderilen", "ekli dosyada" gibi ifadeleri arar.
        
        Args:
            text (str): E-posta içerik metni
            
        Returns:
            bool: Eğer ekli dosyalara referans varsa True, yoksa False
        """
        if not text:
            return False
            
        text = text.lower()
        attachment_keywords = [
            'ekte', 'ekli', 'ekteki', 'ek olarak', 'attachment', 
            'attached', 'enclosed', 'ek dosya', 'ekli dosya',
            'ekte belirtilen', 'ekte gönderilen', 'ekte yer alan',
            'ektedir', 'eklerde', 'eklerde belirtilen'
        ]
        
        for keyword in attachment_keywords:
            if keyword in text:
                return True
                
        return False
        
    def is_stop_sale_chart_file(self, filename):
        """
        Dosya adının stop sale özeti veya grafiği olup olmadığını kontrol eder.
        Bu dosyalar genellikle işleme alınmamalıdır çünkü önceki stop sale'lerin özeti olabilir.
        
        Args:
            filename (str): Kontrol edilecek dosya adı
            
        Returns:
            bool: Eğer dosya adı stop sale chart/özet dosyasına benziyorsa True, yoksa False
        """
        if not filename:
            return False
            
        filename_lower = filename.lower()
        patterns = [
            'stop sale chart', 
            'stopsale chart', 
            'stop-sale-chart',
            'chart of stop', 
            'report of stop',
            'özet', 'summary',
            'tüm stop', 'all stop',
            'genel stop', 'general stop'
        ]
        
        for pattern in patterns:
            if pattern in filename_lower:
                return True
                
        return False

    def smart_clean_email_body(self, email_html: str, email_text: str, sender: str = None) -> str:
        """Cleans email content with smarter handling of forwarded messages and reply chains.
        Preserves important date information and tables while removing clutter."""
        # Check if email is from ecctur.com domain - if so, skip cleaning to preserve forwarded content
        if sender and '@ecctur.com' in sender.lower():
            logger.info(f"Email from ecctur.com domain detected, skipping cleaning: {sender}")
            return email_html if email_html else email_text
            
        content_to_clean = email_html if email_html else email_text
        if not content_to_clean:
            logger.warning("No email content to clean")
            return "EMPTY EMAIL CONTENT"  # Return a non-empty string to avoid ValueErrors
            
        is_html = bool(email_html) and bool(re.search(r'<html|<body|<table|<div|<p>', email_html, re.IGNORECASE))
        
        cleaned_content = ""
        
        if is_html:
            try:
                soup = BeautifulSoup(email_html, 'html.parser')
                
                # Remove script, style, head elements
                for element in soup(["script", "style", "head", "title", "meta", "link"]):
                    element.extract()

                # First try the existing method
                try:
                    # Try standard cleaning approach
                    # Define common reply/forward markers
                    REPLY_FORWARD_MARKERS_REGEX = r'(from:|sent:|to:|subject:|cc:|bcc:|forwarded message|original message|begin forwarded message|on .* wrote:)'
                    
                    # Find ALL markers in the email content
                    marker_elements = []
                    
                    # Search text nodes
                    for element in soup.find_all(string=re.compile(REPLY_FORWARD_MARKERS_REGEX, re.IGNORECASE | re.DOTALL)):
                        parent = element.parent
                        if parent:
                            marker_elements.append(parent)
                    
                    if marker_elements:
                        logger.info(f"Original method found {len(marker_elements)} markers")
                        # Keep the original cleaning logic but ensure content is never empty
                        first_marker = marker_elements[0]
                        logger.info(f"Found first marker element: {first_marker.name}")
                        
                        # Extract content before the marker
                        content_before_marker = []
                        for element in first_marker.previous_siblings:
                            content_before_marker.append(str(element))
                        content_before_marker.reverse()  # Reverse to get correct order
                        
                        # ÖNEMLİ KONTROL: İşaretçi öncesi içerik boş mu?
                        if not content_before_marker:
                            logger.warning("No content before marker, using full email")
                            cleaned_content = soup.get_text(separator=' ', strip=True)
                            if not cleaned_content.strip():
                                cleaned_content = "EMPTY CONTENT AFTER CLEANING - USING RAW HTML"
                                logger.warning("Empty cleaned content, using raw HTML")
                        else:
                            combined_html = ''.join(content_before_marker)
                            new_soup = BeautifulSoup(combined_html, 'html.parser')
                            cleaned_content = new_soup.get_text(separator=' ', strip=True)
                            
                            # Check if we've lost content
                            if not cleaned_content.strip():
                                logger.warning("Empty content after cleaning, using full email")
                                cleaned_content = soup.get_text(separator=' ', strip=True)
                    else:
                        # No markers found, use the whole content
                        cleaned_content = soup.get_text(separator=' ', strip=True)
                except Exception as e:
                    logger.error(f"Error in original HTML cleaning: {e}", exc_info=True)
                    cleaned_content = soup.get_text(separator=' ', strip=True)
                
                # Special handling for tables with dates
                try:
                    # Check for tables with date information
                    date_patterns = ['Tek Gece', 'One Night', 'Single Day', r'\d{2}\.\d{2}\.\d{4}', r'\d{2}/\d{2}/\d{4}']
                    tables = soup.find_all('table')
                    date_tables = []
                    
                    for table in tables:
                        table_text = table.get_text()
                        if any(re.search(pattern, table_text, re.IGNORECASE) for pattern in date_patterns):
                            date_tables.append(table)
                            logger.info(f"Found table with date information")
                    
                    # Add formatted table content
                    if date_tables:
                        table_texts = []
                        for table in date_tables:
                            rows = table.find_all('tr')
                            table_text = []
                            for row in rows:
                                cells = row.find_all(['td', 'th'])
                                if cells:
                                    row_text = ' | '.join(cell.get_text(strip=True) for cell in cells if cell.get_text(strip=True))
                                    if row_text:
                                        table_text.append(row_text)
                            if table_text:
                                table_texts.append("\n-------TABLE:-------\n" + "\n".join(table_text))
                        
                        # Append table texts to cleaned content
                        if table_texts:
                            cleaned_content += "\n\n" + "\n\n".join(table_texts)
                            logger.info(f"Added {len(table_texts)} structured tables to content")
                except Exception as e:
                    logger.error(f"Error processing tables: {e}", exc_info=True)
                
                # Check for date information in the full HTML that might have been lost
                try:
                    full_text = soup.get_text()
                    date_patterns = [
                        r'(\d{2}\.\d{2}\.\d{4}).*?(Tek Gece|One Night)',
                        r'(\d{2}/\d{2}/\d{4}).*?(Tek Gece|One Night)'
                    ]
                    
                    date_info = []
                    for pattern in date_patterns:
                        matches = re.findall(pattern, full_text, re.IGNORECASE)
                        for match in matches:
                            if isinstance(match, tuple):
                                date_str = ' '.join(match)
                            else:
                                date_str = match
                            # Check if this date info is already in cleaned_content
                            if date_str not in cleaned_content:
                                date_info.append(date_str)
                    
                    if date_info:
                        logger.info(f"Adding date information that was lost: {date_info}")
                        cleaned_content += "\n\nIMPORTANT DATES:\n" + "\n".join(date_info)
                except Exception as e:
                    logger.error(f"Error extracting date information: {e}", exc_info=True)
            
            except Exception as e:
                logger.error(f"Error in smart HTML cleaning: {e}", exc_info=True)
                # Fallback to plain text
                cleaned_content = self._clean_plain_text_smart(email_text if email_text else email_html)
        else:
            # Plain text cleaning
            cleaned_content = self._clean_plain_text_smart(content_to_clean)
        
        # Ensure content is not empty
        if not cleaned_content or not cleaned_content.strip():
            logger.warning("Smart cleaning produced empty content, using original content")
            # Try to extract text with basic BeautifulSoup
            if is_html:
                try:
                    soup = BeautifulSoup(content_to_clean, 'html.parser')
                    cleaned_content = soup.get_text(separator=' ', strip=True)
                except Exception:
                    cleaned_content = content_to_clean
            else:
                cleaned_content = content_to_clean
        
        # Final safety check to ensure we always return some content
        if not cleaned_content or not cleaned_content.strip():
            logger.warning("All cleaning attempts produced empty content, returning fixed string")
            cleaned_content = "RAW EMAIL CONTENT COULD NOT BE PROCESSED - PLEASE CHECK MANUALLY"
        
        logger.debug(f"Smart cleaned email content (first 200 chars): {cleaned_content[:200]}")
        return cleaned_content

    def _clean_plain_text_smart(self, text_content: str) -> str:
        """Smarter cleaning of plain text emails with better handling of forwarded content."""
        if not text_content:
            return "EMPTY TEXT CONTENT"
        
        # Define common forwarded message markers
        FORWARD_MARKERS = [
            r'^-{3,}.*?forwarded message.*?-{3,}$',
            r'^_{3,}.*?forwarded message.*?_{3,}$',
            r'^={3,}.*?forwarded message.*?={3,}$',
            r'^>{3,}.*?forwarded message.*?<{3,}$',
            r'^begin forwarded message:$',
            r'^-+\s*original message\s*-+$',
            r'^from:.*$\s*^sent:.*$\s*^to:.*$\s*^subject:.*$',  # Header block pattern
            r'^on\s+.*?wrote:$',  # Reply introduction
        ]
        
        # Join patterns for a single regex search
        forward_pattern = '|'.join(FORWARD_MARKERS)
        
        # Split into lines for processing
        lines = text_content.splitlines()
        cleaned_lines = []
        in_forwarded_section = False
        
        for i, line in enumerate(lines):
            # Check if this line starts a forwarded section
            if re.search(forward_pattern, line, re.IGNORECASE | re.MULTILINE):
                in_forwarded_section = True
                logger.info(f"Found forwarded message marker: '{line}'")
                # If we found a forward marker, only keep content before it
                break
            
            # For lines before the forward marker
            if not in_forwarded_section:
                # Skip common quote markers at line start
                if not line.strip().startswith('>'):
                    cleaned_lines.append(line)
        
        # If we haven't found any forwarded content, use the original text
        if not in_forwarded_section or not cleaned_lines:
            cleaned_lines = lines
        
        # Combine lines back into text
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Look for date information in the original text that might have been lost
        date_patterns = [
            r'\d{2}\.\d{2}\.\d{4}.*?(Tek Gece|One Night)',
            r'\d{2}/\d{2}/\d{4}.*?(Tek Gece|One Night)'
        ]
        
        date_info = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text_content, re.IGNORECASE)
            if matches:
                for match in matches:
                    if isinstance(match, tuple):
                        match = ' '.join(match)
                    if match not in cleaned_text:
                        date_info.append(match)
        
        # Add any found date information
        if date_info:
            logger.info(f"Adding date information that was lost: {date_info}")
            cleaned_text += "\n\nIMPORTANT DATES:\n" + "\n".join(date_info)
        
        # If still empty, return original
        if not cleaned_text.strip():
            return text_content
            
        return cleaned_text