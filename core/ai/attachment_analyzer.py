"""
Attachment analyzer module for email attachments.
Extracts text and data from various file types.
"""

import os
import logging
from django.conf import settings
from typing import Dict, List, Optional, Any, Union
from datetime import datetime, timedelta

# For PDF processing
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    
# For Excel processing
try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# For Word processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class AttachmentAnalyzer:
    """
    Analyzer for email attachments.
    Supports PDF, Excel, Word, and other file types.
    """
    
    def __init__(self, api_key=None):
        """
        Initialize the attachment analyzer
        
        Args:
            api_key (str, optional): API key for Claude or other AI services
        """
        self.api_key = api_key
        
    def analyze(self, file_path: str) -> Dict[str, Any]:
        """
        Analyze an attachment file (high-level method)
        """
        try:
            # Determine file type based on extension
            ext = os.path.splitext(file_path.lower())[1]
            logger.info(f"Analyzing attachment with extension: {ext}")
            
            # Extract text based on file type
            extracted_text = ""
            if ext == '.pdf':
                logger.info(f"Starting PDF analysis for: {file_path}")
                extracted_text = self._extract_pdf_text(file_path)
                
                # Log the first 500 characters of extracted text for debugging
                logger.info(f"Extracted {len(extracted_text)} characters from PDF")
                logger.info(f"PDF text sample: {extracted_text[:500]}...")
                
            elif ext in ['.docx', '.doc']:
                logger.info(f"Starting Word document analysis for: {file_path}")
                extracted_text = self._extract_word_text(file_path)
                logger.info(f"Extracted {len(extracted_text)} characters from Word document")
                logger.info(f"Word text sample: {extracted_text[:500]}...")
                
            elif ext in ['.xlsx', '.xls']:
                logger.info(f"Starting Excel spreadsheet analysis for: {file_path}")
                extracted_text = self._extract_excel_text(file_path)
                logger.info(f"Extracted {len(extracted_text)} characters from Excel")
                logger.info(f"Excel text sample: {extracted_text[:500]}...")
                
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return {'error': f'Unsupported file type: {ext}', 'hotels': []}
            
            # Process the extracted text
            if not extracted_text:
                logger.warning("No text extracted from file")
                return {'error': 'No text extracted from file', 'hotels': []}
            
            # Parse the extracted text using AI
            ai_analysis_result = self._analyze_content_with_ai(extracted_text)
            logger.info(f"AI analysis result: {ai_analysis_result}")
            
            # Process AI analysis result
            if ai_analysis_result:
                return ai_analysis_result
            
            # Fallback to regex parsing if AI failed
            regex_result = self._parse_text_with_regex(extracted_text)
            logger.info(f"Regex analysis result: {regex_result}")
            
            if regex_result and 'hotels' in regex_result and regex_result['hotels']:
                return regex_result
            else:
                return {'error': 'Failed to extract structured data from content', 'hotels': []}
            
        except Exception as e:
            logger.error(f"Error analyzing attachment: {str(e)}", exc_info=True)
            return {'error': f'Error analyzing attachment: {str(e)}', 'hotels': []}
    
    def analyze_attachment(self, file_path: str) -> Dict[str, Any]:
        """
        Analyze an attachment file (DEPRECATED - use analyze method)
        """
        logger.warning("analyze_attachment method is deprecated, use analyze method instead")
        return self.analyze(file_path)
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from an attachment based on file type
        
        Args:
            file_path (str): Path to the attachment file
            
        Returns:
            str: Extracted text content
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_excel_text(file_path)
            elif file_ext == '.docx':
                return self._extract_word_text(file_path)
            elif file_ext == '.txt':
                return self._extract_text_file(file_path)
            else:
                logger.warning(f"Unsupported file extension for text extraction: {file_ext}")
                return f"[Unsupported file type: {file_ext}]"
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return f"[Error extracting text: {str(e)}]"
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """
        Extract text from a PDF file using PyPDF2 and if that fails, fallback to pdfplumber
        """
        text = ""
        
        # First try PyPDF2
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                logger.info(f"PDF has {num_pages} pages")
                
                # Extract text from each page
                for page_num in range(num_pages):
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            
            if text.strip():
                logger.info(f"Successfully extracted {len(text)} characters with PyPDF2")
                return text
            else:
                logger.warning("PyPDF2 extracted empty text, will try pdfplumber")
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}, will try pdfplumber")
        
        # If PyPDF2 failed or returned empty text, try pdfplumber
        try:
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                num_pages = len(pdf.pages)
                logger.info(f"PDF has {num_pages} pages (pdfplumber)")
                
                # Extract text from each page
                for page_num in range(num_pages):
                    page = pdf.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            
            if text.strip():
                logger.info(f"Successfully extracted {len(text)} characters with pdfplumber")
                return text
            else:
                logger.warning("pdfplumber also extracted empty text")
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
        
        # If both methods failed, try pytesseract (OCR) as last resort
        try:
            import pytesseract
            from PIL import Image
            from pdf2image import convert_from_path
            
            logger.info("Both PDF text extraction methods failed, trying OCR with pytesseract")
            
            # Convert PDF to images
            images = convert_from_path(file_path)
            logger.info(f"Converted PDF to {len(images)} images for OCR")
            
            # Extract text from each image using OCR
            for i, image in enumerate(images):
                logger.info(f"Processing page {i+1} with OCR")
                page_text = pytesseract.image_to_string(image)
                if page_text:
                    text += page_text + "\n\n"
            
            if text.strip():
                logger.info(f"Successfully extracted {len(text)} characters with OCR")
                return text
            else:
                logger.error("All extraction methods failed, returning empty string")
                return ""
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            logger.error("All extraction methods failed, returning empty string")
            return ""
    
    def _extract_excel_text(self, file_path: str) -> str:
        """Extract text from Excel files"""
        if not EXCEL_AVAILABLE:
            return "[Excel extraction not available. Install openpyxl package.]"
        
        text = ""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n--- Sheet: {sheet_name} ---\n"
                
                for row in sheet.rows:
                    row_text = []
                    for cell in row:
                        if cell.value is not None:
                            row_text.append(str(cell.value))
                    text += " | ".join(row_text) + "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting Excel text: {str(e)}")
            return f"[Excel extraction error: {str(e)}]"
    
    def _extract_word_text(self, file_path: str) -> str:
        """Extract text from Word documents"""
        if not DOCX_AVAILABLE:
            return "[Word extraction not available. Install python-docx package.]"
        
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text += " | ".join(row_text) + "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting Word text: {str(e)}")
            return f"[Word extraction error: {str(e)}]"
    
    def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting text from file: {str(e)}")
            return f"[Text file extraction error: {str(e)}]"
    
    def analyze_text(self, file_path: str) -> str:
        """
        Analyze a text file
        """
        return self.extract_text(file_path)
        
    def extract_rules_from_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract rules from text content
        
        Args:
            text (str): Text content to extract rules from
            
        Returns:
            List[Dict]: List of rule dictionaries
        """
        rules = []
        
        # Extract hotel names
        hotels = self._extract_hotel_names(text)
        
        # Extract room types
        rooms = self._extract_room_types(text)
        
        # Extract dates
        dates = self._extract_dates(text)
        
        # Extract markets
        markets = self._extract_markets(text)
        
        # If we found hotels, create rules
        if hotels:
            # Default to the first hotel if multiple found
            hotel = hotels[0]
            
            # Default room type
            room_type = "All Room" if not rooms else rooms[0]
            
            # Default market
            market = "ALL" if not markets else markets[0]
            
            # Default dates (today and tomorrow if not found)
            start_date = datetime.now().date()
            end_date = start_date + timedelta(days=1)
            if dates and len(dates) >= 2:
                start_date = dates[0]
                end_date = dates[1]
            elif dates and len(dates) == 1:
                start_date = end_date = dates[0]
            
            # Create rule
            rule = {
                "hotel_name": hotel,
                "room_type": room_type,
                "market": market,
                "start_date": start_date.strftime("%Y-%m-%d"),
                "end_date": end_date.strftime("%Y-%m-%d"),
                "sale_type": "stop"  # Default to stop sale
            }
            
            rules.append(rule)
        
        return rules
    
    def _extract_hotel_names(self, text: str) -> List[str]:
        """
        Extract hotel names from text content
        
        Args:
            text (str): Text content to extract hotel names from
            
        Returns:
            List[str]: List of hotel names
        """
        # Simple example implementation - in a real system, this would be more sophisticated
        # using regex, ML or other techniques
        hotel_names = []
        
        # Look for common hotel name patterns
        hotel_keywords = ["hotel", "resort", "palace", "suites", "inn"]
        
        # Split text into lines
        lines = text.split("\n")
        
        for line in lines:
            for keyword in hotel_keywords:
                if keyword.lower() in line.lower():
                    # Found potential hotel name
                    words = line.strip().split()
                    if len(words) >= 2:  # At least two words for hotel name
                        hotel_name = " ".join(words[:3])  # Take up to first 3 words 
                        hotel_names.append(hotel_name)
                        break
        
        return hotel_names
    
    def _extract_room_types(self, text: str) -> List[str]:
        """
        Extract room types from text content
        
        Args:
            text (str): Text content to extract room types from
            
        Returns:
            List[str]: List of room types
        """
        # Simple example implementation
        room_types = []
        
        # Look for common room type patterns
        room_keywords = ["room", "suite", "deluxe", "standard", "junior", "family"]
        
        # Split text into lines
        lines = text.split("\n")
        
        for line in lines:
            for keyword in room_keywords:
                if keyword.lower() in line.lower():
                    # Found potential room type
                    line_parts = line.lower().split(keyword.lower())
                    if len(line_parts) > 1:
                        before = line_parts[0].strip()
                        after = line_parts[1].strip()
                        
                        # Construct room type
                        if before and len(before.split()) <= 2:
                            room_type = before + " " + keyword
                        else:
                            room_type = keyword + " " + after.split()[0] if after else keyword
                            
                        room_types.append(room_type.strip().title())
                        break
        
        return room_types
    
    def _extract_dates(self, text: str) -> List[datetime.date]:
        """
        Extract dates from text content
        
        Args:
            text (str): Text content to extract dates from
            
        Returns:
            List[datetime.date]: List of dates
        """
        import re
        from datetime import datetime
        
        dates = []
        
        # Define date patterns to search for
        date_patterns = [
            # DD.MM.YYYY or DD/MM/YYYY
            r'\b(\d{1,2})[./-](\d{1,2})[./-](20\d{2})\b',
            # YYYY-MM-DD
            r'\b(20\d{2})-(\d{1,2})-(\d{1,2})\b',
            # Month name DD, YYYY
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{1,2}),?\s+(20\d{2})\b'
        ]
        
        for pattern in date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            
            for match in matches:
                try:
                    if "January" in pattern:  # Month name pattern
                        month_name = match.group(1)
                        day = int(match.group(2))
                        year = int(match.group(3))
                        
                        month_map = {
                            "january": 1, "february": 2, "march": 3, "april": 4,
                            "may": 5, "june": 6, "july": 7, "august": 8,
                            "september": 9, "october": 10, "november": 11, "december": 12
                        }
                        
                        month = month_map.get(month_name.lower())
                        if month:
                            date_obj = datetime(year, month, day).date()
                            dates.append(date_obj)
                    elif "-" in pattern:  # YYYY-MM-DD pattern
                        year = int(match.group(1))
                        month = int(match.group(2))
                        day = int(match.group(3))
                        
                        date_obj = datetime(year, month, day).date()
                        dates.append(date_obj)
                    else:  # DD/MM/YYYY pattern
                        day = int(match.group(1))
                        month = int(match.group(2))
                        year = int(match.group(3))
                        
                        date_obj = datetime(year, month, day).date()
                        dates.append(date_obj)
                except (ValueError, IndexError):
                    # Invalid date, skip
                    pass
        
        return dates
    
    def _extract_markets(self, text: str) -> List[str]:
        """
        Extract markets from text content
        
        Args:
            text (str): Text content to extract markets from
            
        Returns:
            List[str]: List of markets
        """
        # Simple example implementation
        markets = []
        
        # Look for common market keywords
        market_keywords = [
            "market", "all market", "all markets", "pazar", "pazarlar",
            "european market", "uk market", "german market"
        ]
        
        # Common market names that might appear
        common_markets = [
            "UK", "German", "Russia", "Turkish", "Dutch", "Poland", "Belgium",
            "Scandinavia", "Balkan", "European", "Middle East", "CIS"
        ]
        
        # Split text into lines
        lines = text.split("\n")
        
        for line in lines:
            line_lower = line.lower()
            
            # Check for market keywords
            for keyword in market_keywords:
                if keyword.lower() in line_lower:
                    # Extract text around the keyword
                    parts = line_lower.split(keyword.lower())
                    if len(parts) > 1:
                        if parts[1].strip():
                            # Take the words after the keyword
                            market = parts[1].strip().split()[0].upper()
                            markets.append(market)
                        elif parts[0].strip():
                            # Take the words before the keyword
                            market = parts[0].strip().split()[-1].upper()
                            markets.append(market)
                    break
            
            # Check for common market names
            for market in common_markets:
                if market.lower() in line_lower:
                    markets.append(market.upper())
        
        # If no markets found, default to "ALL"
        if not markets:
            markets.append("ALL")
        
        return markets 

    def _analyze_content_with_ai(self, text):
        """
        Analyze the extracted text using AI
        """
        try:
            # Import only when needed to avoid circular imports
            from core.ai.analyzer import ClaudeAnalyzer
            
            logger.info("Using ClaudeAnalyzer to process attachment text")
            ai_analyzer = ClaudeAnalyzer()
            
            # Call the AI analyzer with the text
            prompt = f"""
            Please analyze the following document text and extract any stop sale or open sale information:
            
            {text[:3000]}  # Limit to 3000 chars to stay within token limits
            
            Please return a structured JSON list of hotels and their data in this exact format:
            {{
                "hotels": [
                    {{
                        "name": "HOTEL NAME",
                        "room_type": "ROOM TYPE (or 'All Room' if applicable)",
                        "market": "MARKET NAME (or 'ALL' if all markets)",
                        "date_range": "YYYY-MM-DD - YYYY-MM-DD",
                        "action": "stop_sale" or "open_sale"
                    }},
                    ...
                ]
            }}
            
            If no stop sale or open sale information is found, return an empty hotels list.
            """
            
            logger.info("Sending prompt to ClaudeAnalyzer")
            result = ai_analyzer.custom_prompt(prompt)
            logger.info(f"Raw AI analysis result: {result}")
            
            # Extract the JSON part from the result
            if result and result.get('success', False):
                content = result.get('content', '')
                logger.info(f"AI response content: {content[:500]}...")
                
                # Try to extract JSON from the response
                import re
                import json
                
                # Look for JSON pattern
                json_pattern = r'```json\s*(.*?)\s*```'
                json_matches = re.findall(json_pattern, content, re.DOTALL)
                
                if json_matches:
                    json_str = json_matches[0]
                    logger.info(f"Found JSON in AI response: {json_str[:500]}...")
                    try:
                        parsed_json = json.loads(json_str)
                        logger.info(f"Successfully parsed JSON from AI response: {parsed_json}")
                        # Ensure the result has a 'hotels' key
                        if 'hotels' not in parsed_json:
                            parsed_json = {'hotels': []}
                        return parsed_json
                    except json.JSONDecodeError as json_err:
                        logger.error(f"Error parsing JSON from AI response: {json_err}")
                
                # Try another approach - look for { at beginning of line
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('{') and line.endswith('}'):
                        logger.info(f"Found JSON-like line: {line}")
                        try:
                            parsed_json = json.loads(line)
                            logger.info(f"Successfully parsed JSON from line: {parsed_json}")
                            # Ensure the result has a 'hotels' key
                            if 'hotels' not in parsed_json:
                                parsed_json = {'hotels': []}
                            return parsed_json
                        except json.JSONDecodeError as json_err:
                            logger.error(f"Error parsing JSON from line: {json_err}")
                
                # One more attempt - try to find any JSON object in the text
                try:
                    # Try to find the first { and last } and extract everything between
                    start_idx = content.find('{')
                    end_idx = content.rfind('}')
                    
                    if start_idx != -1 and end_idx != -1 and start_idx < end_idx:
                        json_str = content[start_idx:end_idx+1]
                        logger.info(f"Extracted potential JSON: {json_str[:500]}...")
                        parsed_json = json.loads(json_str)
                        logger.info(f"Successfully parsed JSON from content: {parsed_json}")
                        # Ensure the result has a 'hotels' key
                        if 'hotels' not in parsed_json:
                            parsed_json = {'hotels': []}
                        return parsed_json
                except (json.JSONDecodeError, ValueError) as e:
                    logger.error(f"Error extracting JSON from content: {e}")
                
            logger.warning("No valid JSON found in AI response")
            return {'hotels': []}
            
        except Exception as e:
            logger.error(f"Error in AI analysis: {str(e)}", exc_info=True)
            return {'hotels': []}
    
    def _parse_text_with_regex(self, text):
        """
        Parse the extracted text using regex patterns as a fallback method
        """
        try:
            logger.info("Starting regex-based text analysis")
            
            # Initialize the result structure
            result = {
                "hotels": []
            }
            
            # Look for common patterns in the text
            
            # 1. Look for hotel names (usually followed by specific keywords)
            hotel_patterns = [
                r'([A-Za-z0-9\s\'\-\.]+)\s+Hotel\s+([A-Za-z]+\s+Sale)',
                r'Hotel\s+Name:\s*([A-Za-z0-9\s\'\-\.]+)',
                r'Otel\s+Adı:\s*([A-Za-z0-9\s\'\-\.]+)',
                r'Hotel:\s*([A-Za-z0-9\s\'\-\.]+)'
            ]
            
            # 2. Look for room types
            room_patterns = [
                r'Room\s+Type:\s*([A-Za-z0-9\s\'\-\.]+)',
                r'Oda\s+Tipi:\s*([A-Za-z0-9\s\'\-\.]+)',
                r'Room:\s*([A-Za-z0-9\s\'\-\.]+)'
            ]
            
            # 3. Look for date ranges
            date_patterns = [
                r'(\d{2}\.\d{2}\.\d{4})\s*-\s*(\d{2}\.\d{2}\.\d{4})',  # DD.MM.YYYY - DD.MM.YYYY
                r'(\d{4}-\d{2}-\d{2})\s*-\s*(\d{4}-\d{2}-\d{2})',      # YYYY-MM-DD - YYYY-MM-DD
                r'from\s+(\d{2}\.\d{2}\.\d{4})\s+to\s+(\d{2}\.\d{2}\.\d{4})'  # from DD.MM.YYYY to DD.MM.YYYY
            ]
            
            # 4. Look for sale type
            sale_type_patterns = [
                r'(Stop\s+Sale)',
                r'(Open\s+Sale)',
                r'(Stop\s+Sales)',
                r'(Open\s+Sales)',
                r'(STOP\s+SALE)',
                r'(OPEN\s+SALE)'
            ]
            
            # Extract hotel names
            hotel_names = []
            for pattern in hotel_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    logger.info(f"Found hotel matches with pattern {pattern}: {matches}")
                    for match in matches:
                        if isinstance(match, tuple):
                            hotel_names.append(match[0].strip())
                        else:
                            hotel_names.append(match.strip())
            
            # If no hotel names found, try to extract from the first few lines
            if not hotel_names:
                first_lines = text.split('\n')[:5]  # First 5 lines
                for line in first_lines:
                    # Look for lines that might contain hotel names
                    if any(keyword in line.lower() for keyword in ['hotel', 'otel', 'resort']):
                        logger.info(f"Potential hotel name line: {line}")
                        hotel_names.append(line.strip())
            
            # Extract room types
            room_types = []
            for pattern in room_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    logger.info(f"Found room type matches: {matches}")
                    room_types.extend([match.strip() for match in matches])
            
            # If no specific room types found, use "All Room" as default
            if not room_types:
                room_types = ["All Room"]
                logger.info("No specific room types found, using 'All Room' as default")
            
            # Extract date ranges
            date_ranges = []
            for pattern in date_patterns:
                matches = re.findall(pattern, text)
                if matches:
                    logger.info(f"Found date range matches: {matches}")
                    for match in matches:
                        start_date = match[0]
                        end_date = match[1]
                        
                        # Convert to YYYY-MM-DD format if needed
                        if '.' in start_date:
                            parts = start_date.split('.')
                            if len(parts) == 3:
                                day, month, year = parts
                                start_date = f"{year}-{month}-{day}"
                                
                        if '.' in end_date:
                            parts = end_date.split('.')
                            if len(parts) == 3:
                                day, month, year = parts
                                end_date = f"{year}-{month}-{day}"
                        
                        date_ranges.append(f"{start_date} - {end_date}")
            
            # If no date ranges found, try to find individual dates
            if not date_ranges:
                date_pattern = r'(\d{2}\.\d{2}\.\d{4}|\d{4}-\d{2}-\d{2})'
                dates = re.findall(date_pattern, text)
                if len(dates) >= 2:
                    # Use the first two dates as start and end
                    start_date = dates[0]
                    end_date = dates[1]
                    
                    # Convert to YYYY-MM-DD format if needed
                    if '.' in start_date:
                        parts = start_date.split('.')
                        if len(parts) == 3:
                            day, month, year = parts
                            start_date = f"{year}-{month}-{day}"
                            
                    if '.' in end_date:
                        parts = end_date.split('.')
                        if len(parts) == 3:
                            day, month, year = parts
                            end_date = f"{year}-{month}-{day}"
                    
                    date_ranges.append(f"{start_date} - {end_date}")
                    logger.info(f"Created date range from individual dates: {start_date} - {end_date}")
                elif len(dates) == 1:
                    # Use the same date for start and end
                    date = dates[0]
                    
                    # Convert to YYYY-MM-DD format if needed
                    if '.' in date:
                        parts = date.split('.')
                        if len(parts) == 3:
                            day, month, year = parts
                            date = f"{year}-{month}-{day}"
                    
                    date_ranges.append(f"{date} - {date}")
                    logger.info(f"Created single-day date range: {date}")
            
            # Extract sale type
            sale_type = None
            for pattern in sale_type_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    logger.info(f"Found sale type matches: {matches}")
                    match = matches[0]
                    if isinstance(match, tuple):
                        match = match[0]
                    
                    if 'stop' in match.lower():
                        sale_type = "stop_sale"
                    else:
                        sale_type = "open_sale"
                    break
            
            # If no sale type explicitly found, look for keywords
            if not sale_type:
                if 'stop' in text.lower() or 'durdur' in text.lower() or 'kapalı' in text.lower() or 'kapali' in text.lower():
                    sale_type = "stop_sale"
                elif 'open' in text.lower() or 'açık' in text.lower() or 'acik' in text.lower():
                    sale_type = "open_sale"
                else:
                    # Default to stop_sale if unclear
                    sale_type = "stop_sale"
                    logger.info("No clear sale type found, defaulting to stop_sale")
            
            # Create hotel entries
            if hotel_names:
                for hotel_name in hotel_names:
                    for room_type in room_types:
                        # Use default date range if none found
                        if not date_ranges:
                            # Use current date and next day as default
                            today = datetime.now().date()
                            tomorrow = today + timedelta(days=1)
                            date_range = f"{today.strftime('%Y-%m-%d')} - {tomorrow.strftime('%Y-%m-%d')}"
                            logger.info(f"No date ranges found, using default: {date_range}")
                            date_ranges = [date_range]
                            
                        for date_range in date_ranges:
                            hotel_entry = {
                                "name": hotel_name,
                                "room_type": room_type,
                                "market": "ALL",  # Default to ALL market
                                "date_range": date_range,
                                "action": sale_type
                            }
                            result["hotels"].append(hotel_entry)
                            logger.info(f"Added hotel entry: {hotel_entry}")
            
            # If we found data, return the result
            if result["hotels"]:
                logger.info(f"Regex analysis found {len(result['hotels'])} hotel entries")
                return result
            else:
                logger.warning("Regex analysis couldn't extract structured data")
                return {'hotels': []}
            
        except Exception as e:
            logger.error(f"Error in regex text analysis: {str(e)}", exc_info=True)
            return {'hotels': []} 