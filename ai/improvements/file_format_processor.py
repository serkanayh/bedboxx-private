"""
File Format Processor Module for StopSale Automation System

This module provides support for extracting information from various file formats
attached to emails, including PDF, Excel, and Word documents.
"""

import os
import re
import logging
import tempfile
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime

# Set up logging
logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Word document processing will be limited.")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logger.warning("pandas not installed. Excel processing will be limited.")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning("PyPDF2 not installed. PDF processing will be limited.")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not installed. Advanced PDF table extraction will be limited.")


class FileFormatProcessor:
    """Base class for processing different file formats"""
    
    def __init__(self):
        """Initialize the processor"""
        self.supported_formats = self._get_supported_formats()
        
    def _get_supported_formats(self) -> Dict[str, bool]:
        """Get dictionary of supported formats and their availability"""
        return {
            'pdf': PYPDF2_AVAILABLE or PDFPLUMBER_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'xlsx': PANDAS_AVAILABLE,
            'xls': PANDAS_AVAILABLE,
            'csv': PANDAS_AVAILABLE,
        }
    
    def is_supported(self, filename: str) -> bool:
        """
        Check if a file format is supported
        
        Args:
            filename: Name of the file to check
            
        Returns:
            bool: True if the format is supported, False otherwise
        """
        ext = os.path.splitext(filename)[1].lower().lstrip('.')
        return ext in self.supported_formats and self.supported_formats[ext]
    
    def process_file(self, file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Process a file and extract its content and structured data
        
        Args:
            file_path: Path to the file
            
        Returns:
            tuple: (extracted_text, structured_data)
                extracted_text: Plain text content of the file
                structured_data: List of dictionaries containing structured data
        """
        ext = os.path.splitext(file_path)[1].lower().lstrip('.')
        
        if not self.is_supported(file_path):
            logger.warning(f"Unsupported file format: {ext}")
            return "", []
        
        if ext == 'pdf':
            return self._process_pdf(file_path)
        elif ext == 'docx':
            return self._process_docx(file_path)
        elif ext in ['xlsx', 'xls']:
            return self._process_excel(file_path)
        elif ext == 'csv':
            return self._process_csv(file_path)
        else:
            logger.warning(f"Unhandled file format: {ext}")
            return "", []
    
    def _process_pdf(self, file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Process a PDF file
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            tuple: (extracted_text, structured_data)
        """
        extracted_text = ""
        structured_data = []
        
        # Try using pdfplumber first for better table extraction
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    # Extract text from each page
                    for page in pdf.pages:
                        extracted_text += page.extract_text() + "\n\n"
                        
                        # Extract tables
                        tables = page.extract_tables()
                        for table in tables:
                            # Convert table to structured data
                            if table and len(table) > 1:  # Has header and data rows
                                headers = [str(h).strip() for h in table[0]]
                                for row in table[1:]:
                                    row_data = {headers[i]: str(cell).strip() 
                                              for i, cell in enumerate(row) 
                                              if i < len(headers) and cell}
                                    if row_data:
                                        structured_data.append(row_data)
                
                logger.info(f"Extracted {len(structured_data)} table rows from PDF using pdfplumber")
                return extracted_text, structured_data
            
            except Exception as e:
                logger.error(f"Error processing PDF with pdfplumber: {str(e)}")
                # Fall back to PyPDF2
        
        # Fall back to PyPDF2 if pdfplumber failed or is not available
        if PYPDF2_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        extracted_text += page.extract_text() + "\n\n"
                
                logger.info(f"Extracted text from PDF using PyPDF2")
                # PyPDF2 doesn't extract tables well, so we'll try to parse tables from text
                structured_data = self._extract_tables_from_text(extracted_text)
                return extracted_text, structured_data
            
            except Exception as e:
                logger.error(f"Error processing PDF with PyPDF2: {str(e)}")
        
        return extracted_text, structured_data
    
    def _process_docx(self, file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Process a Word document
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            tuple: (extracted_text, structured_data)
        """
        extracted_text = ""
        structured_data = []
        
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not installed. Cannot process Word document.")
            return extracted_text, structured_data
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                # Get headers from first row
                if len(table.rows) > 1:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    
                    # Process data rows
                    for row in table.rows[1:]:
                        row_data = {headers[i]: cell.text.strip() 
                                  for i, cell in enumerate(row.cells) 
                                  if i < len(headers) and cell.text.strip()}
                        if row_data:
                            structured_data.append(row_data)
            
            logger.info(f"Extracted {len(structured_data)} table rows from Word document")
            return extracted_text, structured_data
        
        except Exception as e:
            logger.error(f"Error processing Word document: {str(e)}")
            return extracted_text, structured_data
    
    def _process_excel(self, file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Process an Excel file
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            tuple: (extracted_text, structured_data)
        """
        extracted_text = ""
        structured_data = []
        
        if not PANDAS_AVAILABLE:
            logger.warning("pandas not installed. Cannot process Excel file.")
            return extracted_text, structured_data
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert sheet to text
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False) + "\n\n"
                extracted_text += sheet_text
                
                # Convert to structured data
                for _, row in df.iterrows():
                    row_dict = row.to_dict()
                    # Clean up NaN values
                    row_dict = {k: str(v) for k, v in row_dict.items() if pd.notna(v)}
                    if row_dict:
                        structured_data.append(row_dict)
            
            logger.info(f"Extracted {len(structured_data)} rows from Excel file")
            return extracted_text, structured_data
        
        except Exception as e:
            logger.error(f"Error processing Excel file: {str(e)}")
            return extracted_text, structured_data
    
    def _process_csv(self, file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Process a CSV file
        
        Args:
            file_path: Path to the CSV file
            
        Returns:
            tuple: (extracted_text, structured_data)
        """
        extracted_text = ""
        structured_data = []
        
        if not PANDAS_AVAILABLE:
            logger.warning("pandas not installed. Cannot process CSV file.")
            return extracted_text, structured_data
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                logger.error("Failed to read CSV file with any encoding")
                return extracted_text, structured_data
            
            # Convert to text
            extracted_text = df.to_string(index=False)
            
            # Convert to structured data
            for _, row in df.iterrows():
                row_dict = row.to_dict()
                # Clean up NaN values
                row_dict = {k: str(v) for k, v in row_dict.items() if pd.notna(v)}
                if row_dict:
                    structured_data.append(row_dict)
            
            logger.info(f"Extracted {len(structured_data)} rows from CSV file")
            return extracted_text, structured_data
        
        except Exception as e:
            logger.error(f"Error processing CSV file: {str(e)}")
            return extracted_text, structured_data
    
    def _extract_tables_from_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Attempt to extract table-like structures from text
        
        Args:
            text: The text to parse
            
        Returns:
            list: List of dictionaries containing structured data
        """
        structured_data = []
        
        # Look for common table patterns
        # This is a simplified approach and may not work for all tables
        
        # Try to find lines that look like headers (containing multiple words separated by whitespace)
        lines = text.split('\n')
        for i, line in enumerate(lines):
            if i + 1 >= len(lines):
                continue
                
            # Check if this line could be a header
            words = [w for w in re.split(r'\s{2,}', line.strip()) if w]
            if len(words) >= 3:  # At least 3 columns
                # Check the next few lines to see if they follow the same pattern
                potential_headers = words
                data_rows = []
                
                for j in range(i + 1, min(i + 10, len(lines))):
                    row_data = [w for w in re.split(r'\s{2,}', lines[j].strip()) if w]
                    if len(row_data) == len(potential_headers):
                        data_rows.append(row_data)
                
                # If we found at least one data row, consider this a table
                if data_rows:
                    for row in data_rows:
                        row_dict = {potential_headers[k]: row[k] 
                                   for k in range(len(potential_headers))
                                   if k < len(row)}
                        if row_dict:
                            structured_data.append(row_dict)
        
        return structured_data


class EmailAttachmentProcessor:
    """Class for processing email attachments"""
    
    def __init__(self):
        """Initialize the processor"""
        self.file_processor = FileFormatProcessor()
    
    def process_attachment(self, attachment_path: str) -> Dict[str, Any]:
        """
        Process an email attachment
        
        Args:
            attachment_path: Path to the attachment file
            
        Returns:
            dict: Processed attachment data
        """
        if not os.path.exists(attachment_path):
            logger.error(f"Attachment file not found: {attachment_path}")
            return {
                "success": False,
                "error": "File not found",
                "file_path": attachment_path,
                "content": "",
                "structured_data": []
            }
        
        if not self.file_processor.is_supported(attachment_path):
            logger.warning(f"Unsupported attachment format: {attachment_path}")
            return {
                "success": False,
                "error": "Unsupported file format",
                "file_path": attachment_path,
                "content": "",
                "structured_data": []
            }
        
        try:
            content, structured_data = self.file_processor.process_file(attachment_path)
            
            return {
                "success": True,
                "file_path": attachment_path,
                "content": content,
                "structured_data": structured_data
            }
        
        except Exception as e:
            logger.error(f"Error processing attachment: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "file_path": attachment_path,
                "content": "",
                "structured_data": []
            }
    
    def process_attachments(self, attachment_paths: List[str]) -> List[Dict[str, Any]]:
        """
        Process multiple email attachments
        
        Args:
            attachment_paths: List of paths to attachment files
            
        Returns:
            list: List of processed attachment data
        """
        results = []
        
        for path in attachment_paths:
            result = self.process_attachment(path)
            results.append(result)
        
        return results


def install_dependencies():
    """Install required dependencies if not already installed"""
    try:
        import pip
        
        # Check and install dependencies
        if not DOCX_AVAILABLE:
            print("Installing python-docx...")
            pip.main(['install', 'python-docx'])
        
        if not PANDAS_AVAILABLE:
            print("Installing pandas...")
            pip.main(['install', 'pandas'])
        
        if not PYPDF2_AVAILABLE:
            print("Installing PyPDF2...")
            pip.main(['install', 'PyPDF2'])
        
        if not PDFPLUMBER_AVAILABLE:
            print("Installing pdfplumber...")
            pip.main(['install', 'pdfplumber'])
        
        print("Dependencies installed successfully.")
        
    except Exception as e:
        print(f"Error installing dependencies: {str(e)}")


if __name__ == "__main__":
    # Example usage
    processor = EmailAttachmentProcessor()
    
    # Check if dependencies are installed
    missing_deps = []
    if not DOCX_AVAILABLE:
        missing_deps.append("python-docx")
    if not PANDAS_AVAILABLE:
        missing_deps.append("pandas")
    if not PYPDF2_AVAILABLE:
        missing_deps.append("PyPDF2")
    if not PDFPLUMBER_AVAILABLE:
        missing_deps.append("pdfplumber")
    
    if missing_deps:
        print(f"Missing dependencies: {', '.join(missing_deps)}")
        install = input("Do you want to install missing dependencies? (y/n): ")
        if install.lower() == 'y':
            install_dependencies()
    
    # Example processing
    test_file = input("Enter path to a test file (PDF, DOCX, XLSX, CSV): ")
    if os.path.exists(test_file):
        result = processor.process_attachment(test_file)
        print(f"Processing result: {'Success' if result['success'] else 'Failed'}")
        print(f"Extracted {len(result['structured_data'])} structured data items")
        print(f"Content preview: {result['content'][:200]}...")
    else:
        print("File not found.")
