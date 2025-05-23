from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, JsonResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.clickjacking import xframe_options_sameorigin, xframe_options_exempt
from django.views.decorators.http import require_POST
from django.utils import timezone
from django.conf import settings
from django.core.paginator import Paginator
from django.db.models import Q
from django.contrib import messages
from django.urls import reverse
from django.db import transaction
from urllib.parse import urlencode

from .models import Email, EmailRow, UserLog, EmailAttachment, EmailMarketMatch, EmailContractMatch
from hotels.models import Hotel, Room, Market, JuniperContractMarket, HotelLearning
from core.models import AIPerformanceMetric

import json
import logging
import os
import tempfile
from datetime import datetime, timedelta
import traceback
import re
import io
import pandas as pd
from docx import Document

# Set up logger
logger = logging.getLogger(__name__)

# Try to import the AI analyzer
try:
    from ai.analyzer import ClaudeAnalyzer
    AI_ANALYZER_AVAILABLE = True
    logger.info("ClaudeAnalyzer successfully imported")
except ImportError as e:
    logger.error(f"Failed to import ClaudeAnalyzer: {e}")
    AI_ANALYZER_AVAILABLE = False


@login_required
def email_list(request):
    """
    View for listing emails
    """
    # Get query parameters
    status = request.GET.get('status', '')
    search = request.GET.get('search', '')
    sort_order = request.GET.get('sort', 'asc')  # Default is ascending (oldest first)
    date_pattern = request.GET.get('date_pattern', '')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')
    sender = request.GET.get('sender', '')
    
    # If date_pattern is provided, it overrides start_date and end_date
    if date_pattern:
        try:
            from .filters import get_date_filter_params
            pattern_start_date, pattern_end_date = get_date_filter_params(date_pattern)
            if pattern_start_date and pattern_end_date:
                start_date = pattern_start_date
                end_date = pattern_end_date
        except ImportError:
            # If filters module is not available, just use provided dates
            pass
    
    # Base queryset
    queryset = Email.objects.all()
    
    # Apply filters
    if status:
        queryset = queryset.filter(status=status)
    
    if search:
        queryset = queryset.filter(
            Q(subject__icontains=search) | 
            Q(sender__icontains=search) |
            Q(body_text__icontains=search)
        )
    
    # Apply date filters
    if start_date:
        try:
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
            queryset = queryset.filter(received_date__date__gte=start_date_obj)
        except ValueError:
            # If date format is invalid, ignore this filter
            pass
    
    if end_date:
        try:
            end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
            queryset = queryset.filter(received_date__date__lte=end_date_obj)
        except ValueError:
            # If date format is invalid, ignore this filter
            pass
    
    # Apply sender filter
    if sender:
        queryset = queryset.filter(sender__icontains=sender)
    
    # Order by received date (oldest first by default, user can change)
    if sort_order == 'desc':
        queryset = queryset.order_by('-received_date')  # Newest first
    else:
        queryset = queryset.order_by('received_date')   # Oldest first
    
    # Hesaplanan özellikler: Önceden hesaplayıp template'e gönder
    for email in queryset:
        email.total_count = email.total_rules_count
        
        # Güncellenmiş matched_rules_count mantığını kullan
        # Juniper otel eşleşmesi olan ve juniper_rooms ilişkisinde kayıt olan satırlar
        rows_with_rooms = email.rows.filter(
            juniper_hotel__isnull=False
        ).filter(
            juniper_rooms__isnull=False
        ).distinct()
        
        # Juniper otel eşleşmesi olan ve ALL ROOM tipi olan satırlar
        all_room_rows = email.rows.filter(
            juniper_hotel__isnull=False, 
            room_type__iregex=r'all\s*room'  # "ALL ROOM", "All Room", "ALLROOM" gibi varyasyonlar
        ).distinct()
        
        # İki sorgu sonucunu birleştir ve tekrar sayısını say (distinct)
        email.matched_count = (rows_with_rooms | all_room_rows).distinct().count()
        
        # Eşleşme oranı rengini belirle
        if email.total_count == 0:
            email.match_status_color = "secondary"  # Gri renk (eşleşme yok)
        elif email.matched_count == email.total_count:
            email.match_status_color = "success"    # Yeşil renk (tam eşleşme)
        elif email.matched_count >= email.total_count * 0.75:
            email.match_status_color = "warning"    # Turuncu renk (yüksek eşleşme)
        else:
            email.match_status_color = "danger"     # Kırmızı renk (düşük eşleşme)
    
    # Paginate
    paginator = Paginator(queryset, 20)  # Show 20 emails per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'status_filter': status,
        'search_query': search,
        'sort_order': sort_order,  # Pass current sort order to template
    }
    
    return render(request, 'emails/email_list.html', context)


@login_required
def email_detail(request, email_id):
    """
    View for displaying email details
    """
    email = get_object_or_404(Email, id=email_id)
    rows = email.rows.all().order_by('id')
    attachments = email.attachments.all().order_by('id')

    # Filtreleme
    match = request.GET.get('match')
    if match == 'danger':
        rows = rows.filter(juniper_rooms__isnull=True)
    elif match == 'success':
        rows = [r for r in rows if hasattr(r, 'get_match_status') and r.get_match_status() == 'success']
    elif match == 'warning':
        rows = [r for r in rows if hasattr(r, 'get_match_status') and r.get_match_status() == 'warning']

    # --- ADD analysis result to each attachment --- 
    attachment_analysis_results = email.attachment_analysis_results or {}
    for att in attachments:
        try:
            att_id_str = str(att.id)
            att.analysis_result = attachment_analysis_results.get(att_id_str, None)
        except Exception as e:
            logger.error(f"Error assigning analysis result to attachment {att.id} for email {email.id}: {e}")
            att.analysis_result = None
    # --- END ADD --- 

    # Process HTML content for display with embedded images
    processed_html = process_html_for_display(email)

    # AI önerilerini oluştur
    ai_suggestions = []
    
    # Eşleşmesi olmayan satırlar için öneriler üret
    unmatched_rows = [r for r in rows if (r.status == 'pending' or r.status == 'matching' or r.status == 'hotel_not_found' or r.status == 'room_not_found') and (not r.juniper_hotel or (not r.juniper_rooms.exists() and not r.room_type.lower().strip() in ['all room', 'all rooms', 'all room types']))]
    
    if unmatched_rows:
        # EmailHotelMatch ve RoomTypeMatch modellerini kullanarak öneriler oluştur
        from .models import EmailHotelMatch, RoomTypeMatch
        from hotels.models import Hotel, Room # Import Hotel and Room models
        from difflib import SequenceMatcher # Import SequenceMatcher
        from .views import get_hotel_suggestions, get_room_suggestions # Import suggestion functions

        # Yeni otel önerileri listesi
        hotel_suggestions_list = []

        for row in unmatched_rows: # Loop through all unmatched rows
            if row.hotel_name and not row.juniper_hotel:
                # Get hotel suggestions using the helper function
                best_hotel_match, current_hotel_suggestions = get_hotel_suggestions(row.hotel_name)

                # Eger en iyi otel eşleşmesi bulunduysa ve güven skoru yüksekse (örneğin %80 üzeri)
                if best_hotel_match and getattr(best_hotel_match, 'match_score', 0) >= 65:
                    row.juniper_hotel = best_hotel_match
                    # Durumu da eşleşme bulundu olarak güncelle (opsiyonel, isteğe bağlı olarak bırakılabilir)
                    # row.status = 'matching' # or another appropriate status
                    row.save()
                    logger.info(f"Row {row.id} için otomatik otel eşleştirme yapıldı: {best_hotel_match.juniper_hotel_name}")

                if current_hotel_suggestions:
                    # Add row_id to each suggestion and append to the main list
                    for suggestion_hotel in current_hotel_suggestions:
                        suggestion = {
                            'row_id': row.id,
                            'hotel_id': suggestion_hotel.id,
                            'hotel_name': suggestion_hotel.juniper_hotel_name,
                            'room_type': "Sistem tarafından belirlenecek", # This will be updated by room suggestions later
                            'room_ids': [],
                            'confidence': getattr(suggestion_hotel, 'match_score', 0), # Use match_score from get_hotel_suggestions
                            'reason': f"Otel adı benzerliği: {getattr(suggestion_hotel, 'match_score', 0)}%"
                        }
                        hotel_suggestions_list.append(suggestion)

        # Remove duplicates based on row_id and hotel_id, keeping the one with higher confidence if duplicated
        unique_suggestions = {}
        for suggestion in hotel_suggestions_list:
            key = (suggestion['row_id'], suggestion['hotel_id'])
            if key not in unique_suggestions or suggestion['confidence'] > unique_suggestions[key]['confidence']:
                unique_suggestions[key] = suggestion

        # Convert back to list and sort by confidence (descending)
        ai_suggestions = sorted(unique_suggestions.values(), key=lambda x: x['confidence'], reverse=True)

        # Limit to a reasonable number of suggestions overall, maybe top 10?
        ai_suggestions = ai_suggestions[:10]

        # Note: Room suggestions in email_detail view might be complex as one email has multiple rows.
        # For now, focus on hotel suggestions here.

        # Oda eşleştirmesi önerilerini de hesapla ve otomatik ata
        for row in unmatched_rows:
            # Sadece oteli eşleşmiş ancak odası eşleşmemiş satırlar için oda önerisi yap
            is_all_room_type = row.room_type.lower().strip() in ['all room', 'all rooms', 'all room types']
            if row.juniper_hotel and not row.juniper_rooms.exists() and not is_all_room_type:
                best_room_match, room_suggestions, search_pattern = get_room_suggestions(row.room_type, row.juniper_hotel)
                
                # Eger en iyi oda eşleşmesi bulunduysa
                if best_room_match:
                    row.juniper_rooms.set([best_room_match]) # Set expects an iterable
                    # Durumu pending olarak güncelle (veya başka bir uygun durum)
                    if row.status != 'approved': # Eğer zaten onaylanmamışsa durumu güncelle
                        row.status = 'pending'
                    row.save()
                    logger.info(f"Row {row.id} için otomatik oda eşleştirmesi yapıldı: {best_room_match.juniper_room_type}")
                    # Propagate this room match to other rows with the same room type (optional)
                    # propagate_room_matching(row)


    has_analyzed_attachments = rows.filter(extracted_from_attachment=True).exists() if hasattr(rows, 'filter') else any(getattr(r, 'extracted_from_attachment', False) for r in rows)
    first_matched_row = rows.filter(juniper_hotel__isnull=False).first() if hasattr(rows, 'filter') else None
    unique_hotel_names = list(rows.exclude(hotel_name__isnull=True).exclude(hotel_name__exact='').values_list('hotel_name', flat=True).distinct().order_by('hotel_name')) if hasattr(rows, 'exclude') else []
    
    # Check if any row has 'approved' status
    has_approved_rows = email.rows.filter(status='approved').exists() if hasattr(email.rows, 'filter') else any(r.status == 'approved' for r in email.rows.all())

    # Tüm otellerin ve pazarların listesini ekle (Manuel Kural Ekleme modalı için)
    from hotels.models import Hotel, Room, Market  # Make sure models are imported here
    hotels = Hotel.objects.all().order_by('juniper_hotel_name')
    markets = Market.objects.all().order_by('name')

    context = {
        'email': email,
        'rows': rows,
        'attachments': attachments,
        'has_analyzed_attachments': has_analyzed_attachments,
        'first_matched_row': first_matched_row,
        'unique_hotel_names': unique_hotel_names,
        'ai_suggestions': ai_suggestions,  # AI önerilerini context'e ekle
        'processed_html': processed_html,  # Add processed HTML to context
        'has_approved_rows': has_approved_rows,  # Onaylanmış satırlar var mı?
        'hotels': hotels,  # Manuel kural ekleme için oteller listesi
        'markets': markets,  # Manuel kural ekleme için pazarlar listesi
    }
    return render(request, 'emails/email_detail.html', context)


@login_required
@xframe_options_exempt
def email_attachment_view(request, attachment_id):
    """
    View for displaying email attachment (allows embedding anywhere)
    """
    attachment = get_object_or_404(EmailAttachment, id=attachment_id)
    
    # Check if the file exists
    if not os.path.exists(attachment.file.path):
        return HttpResponse("Attachment file not found", status=404)
    
    # Determine content type - use the model's properties to get accurate content type
    content_type = attachment.content_type
    if not content_type:
        content_type = 'application/octet-stream'  # Default content type
    
    # Get decoded filename for proper display
    decoded_filename = attachment.decoded_filename
    file_basename = os.path.basename(decoded_filename)
    
    # Determine file type based on properties
    is_pdf = attachment.is_pdf
    is_image = attachment.is_image
    is_text = attachment.is_text
    
    # Define safely previewable content types
    safe_preview_types = [
        'application/pdf',
        'text/plain',
        'text/html',
        'image/jpeg',
        'image/png',
        'image/gif',
        'image/svg+xml'
    ]
    
    # Handle problematic file formats
    force_download = False
    
    # If content type isn't safely previewable, force download by default
    # PDFs are always attempted to be displayed inline regardless
    if not is_pdf and (content_type not in safe_preview_types and not is_image and not is_text):
        force_download = True
    
    # Get force_preview parameter from query string
    force_preview = request.GET.get('force_preview', '').lower() == 'true'
    if force_preview:
        force_download = False
    
    try:
        with open(attachment.file.path, 'rb') as f:
            file_content = f.read()
            
            response = HttpResponse(file_content, content_type=content_type)
            
            # Use 'attachment' disposition for download (to prevent 415 errors)
            # Use 'inline' for preview attempt
            if force_download:
                response['Content-Disposition'] = f'attachment; filename="{file_basename}"'
            else:
                response['Content-Disposition'] = f'inline; filename="{file_basename}"'
            
            # Override Content-Type for PDFs to ensure correct handling
            if is_pdf:
                response['Content-Type'] = 'application/pdf'
            elif is_image and not content_type.startswith('image/'):
                # Try to determine image type if content type isn't already set
                if attachment.file_extension == '.jpg' or attachment.file_extension == '.jpeg':
                    response['Content-Type'] = 'image/jpeg'
                elif attachment.file_extension == '.png':
                    response['Content-Type'] = 'image/png'
                elif attachment.file_extension == '.gif':
                    response['Content-Type'] = 'image/gif'
                elif attachment.file_extension == '.svg':
                    response['Content-Type'] = 'image/svg+xml'
            
            return response
    except FileNotFoundError:
        logger.error(f"File not found during serving: {attachment.file.path}")
        return HttpResponse("Attachment file physically not found on server", status=404)
    except Exception as e:
        logger.error(f"Error serving attachment {attachment_id}: {e}", exc_info=True)
        return HttpResponse("Error serving attachment", status=500)


@login_required
def process_email_with_ai(request, email_id):
    """
    View for processing an email with AI
    """
    email = get_object_or_404(Email, id=email_id)
    
    if not AI_ANALYZER_AVAILABLE:
        messages.error(request, "AI analyzer is not available")
        return redirect('emails:email_detail', email_id=email_id)
    
    # Check if the email already has rows
    if email.rows.exists():
        messages.warning(request, "This email has already been processed")
        return redirect('emails:email_detail', email_id=email_id)
    
    try:
        # Create an instance of the AI analyzer
        analyzer = ClaudeAnalyzer()
        
        # Process the email content
        start_time = timezone.now()
        result = analyzer.analyze_email(email.body_text, email.subject)
        end_time = timezone.now()
        
        # Log AI performance
        AIPerformanceMetric.objects.create(
            email=email,
            processing_time=(end_time - start_time).total_seconds(),
            model_name=analyzer.get_model_name(),
            prompt_tokens=analyzer.get_prompt_tokens(),
            completion_tokens=analyzer.get_completion_tokens(),
            success=result.get('success', False)
        )
        
        if result.get('success', False):
            rows_data = result.get('data', [])
            created_rows = [] # Keep track of created rows for batch matching
            for row_data in rows_data:
                
                # --- Parse dates using the helper function --- 
                start_date_str = row_data.get('start_date')
                end_date_str = row_data.get('end_date')
                
                parsed_start_date = parse_ai_date(start_date_str)
                parsed_end_date = parse_ai_date(end_date_str)
                # --- End date parsing --- 
                
                try:
                    email_row = EmailRow.objects.create(
                        email=email,
                        hotel_name=row_data.get('hotel_name', ''),
                        room_type=row_data.get('room_type', ''),
                        market=row_data.get('market', ''), # Assuming market is a string from AI
                        start_date=parsed_start_date, # Use parsed date
                        end_date=parsed_end_date,   # Use parsed date
                        sale_type=row_data.get('sale_type', 'stop'),
                        status='matching', # Start with matching status for the task
                        ai_extracted=True
                    )
                    created_rows.append(email_row.id)

                except Exception as create_error:
                     logger.error(f"Error creating EmailRow for data: {row_data}. Error: {create_error}", exc_info=True)
                     continue # Skip to the next row_data

            # --- Trigger batch matching task --- 
            if created_rows:
                logger.info(f"Scheduling BATCH matching task for email {email.id} with {len(created_rows)} rows.")
                from .tasks import match_email_rows_batch_task # Import task here
                match_email_rows_batch_task.delay(email.id, created_rows)
                email.status = 'processing' # Or a more specific status
                email.save()
            else:
                 logger.warning(f"AI analysis successful for email {email.id} but no rows were created.")
                 email.status = 'processed_nodata' 
                 email.save()

            messages.success(request, f"AI processing initiated for email {email.id}. Rows are being matched.") # Updated message
        else:
            # If AI processing failed but email has attachments, try to analyze attachments
            if email.has_attachments and email.attachments.exists():
                success = process_email_attachments(email, request.user)
                if success:
                    messages.success(request, "Email processed using attachment analysis")
                else:
                    messages.error(request, f"Failed to process email with AI: {result.get('error', 'Unknown error')}")
            else:
                messages.error(request, f"Failed to process email with AI: {result.get('error', 'Unknown error')}")
    
    except Exception as e:
        logger.error(f"Error processing email with AI: {e}", exc_info=True)
        messages.error(request, f"Error processing email: {str(e)}")
    
    return redirect('emails:email_detail', email_id=email_id)


def process_email_attachments(email, user):
    """
    Process email attachments to extract stop sale/open sale information
    """
    try:
        # Import the AttachmentAnalyzer class
        from core.ai.attachment_analyzer import AttachmentAnalyzer
        
        # Initialize the attachment analyzer
        attachment_analyzer = AttachmentAnalyzer()
        
        # Check if email has attachments
        attachments = email.attachments.all()
        if not attachments.exists():
            logger.warning(f"No attachments found for email {email.id}")
            return False
        
        # Track if any attachment was successfully analyzed
        any_success = False
        # --- Define allowed file extensions --- 
        allowed_extensions = ('.pdf', '.docx', '.xlsx', '.xls', '.doc')
        
        created_row_ids = []
        for attachment in attachments:
            # --- Check file extension using the attachment's file_extension property --- 
            # This properly handles MIME-encoded filenames by decoding them first
            file_ext = attachment.file_extension
            if file_ext.endswith('?='):
                file_ext = file_ext[:-2]
                logger.warning(f"Fixed problematic file extension from {attachment.file_extension} to {file_ext} for {attachment.filename}")
                
            if not file_ext in allowed_extensions:
                logger.info(f"Skipping attachment analysis for {attachment.filename} (Email ID: {email.id}) - Unsupported file type: {file_ext}")
                continue # Skip to the next attachment
            
            # Skip if file doesn't exist
            if not os.path.exists(attachment.file.path):
                logger.warning(f"Attachment file not found: {attachment.file.path}")
                continue
            
            logger.info(f"Analyzing attachment: {attachment.filename} (Email ID: {email.id}) - Type OK.") # Log analysis start
            # Analyze the attachment
            result = attachment_analyzer.analyze(attachment.file.path)
            
            # Skip if analysis failed
            if 'error' in result:
                logger.warning(f"Failed to analyze attachment {attachment.id}: {result['error']}")
                continue
            
            # Process extracted rules - YENİ DÖNÜŞ YAPISI (hotels dizisi)
            hotels_data = result.get('hotels', [])
            
            if not hotels_data:
                logger.warning(f"No hotel data found in attachment {attachment.id}")
                continue
                
            logger.info(f"Found {len(hotels_data)} hotel entries in attachment {attachment.id}")
            
            # Process each hotel entry
            for hotel_data in hotels_data:
                try:
                    # Extract hotel data
                    hotel_name = hotel_data.get('name', '')
                    room_type = hotel_data.get('room_type', 'All Room')
                    market = hotel_data.get('market', 'ALL')
                    date_range = hotel_data.get('date_range', '')
                    action = hotel_data.get('action', 'stop_sale')
                    
                    # Parse date range
                    start_date = None
                    end_date = None
                    
                    if date_range:
                        date_parts = date_range.split(' - ')
                        if len(date_parts) >= 2:
                            start_date = date_parts[0].strip()
                            end_date = date_parts[1].strip()
                        elif len(date_parts) == 1:
                            # If only one date, use it for both start and end
                            start_date = end_date = date_parts[0].strip()
                    
                    # If dates are still not set, use default
                    if not start_date or not end_date:
                        today = datetime.now().date()
                        tomorrow = today + timedelta(days=1)
                        start_date = today.strftime('%Y-%m-%d')
                        end_date = tomorrow.strftime('%Y-%m-%d')
                    
                    # Determine sale status
                    sale_status = 'stop' if action.lower() == 'stop_sale' else 'open'
                    
                    # Create EmailRow
                    email_row = EmailRow.objects.create(
                        email=email,
                        hotel_name=hotel_name,
                        room_type=room_type,
                        market=market,
                        start_date=start_date,
                        end_date=end_date,
                        sale_status=sale_status,
                        source='attachment',
                        status='pending',
                        ai_extracted=True,
                        extracted_from_attachment=True,
                        created_by=user
                    )
                    
                    created_row_ids.append(email_row.id)
                    any_success = True
                    logger.info(f"Created EmailRow {email_row.id} from attachment {attachment.id}")
                    
                except Exception as row_err:
                    logger.error(f"Error creating EmailRow from attachment {attachment.id}: {str(row_err)}", exc_info=True)
                    continue

        # Update email status based on whether any attachment yielded rows
        if any_success:
            # Update email status to indicate successful attachment processing
            email.status = 'processed_attachments'
            
            # Schedule matching task if rows were created
            if created_row_ids:
                logger.info(f"Created {len(created_row_ids)} rows from attachments, scheduling matching task...")
                try:
                    from emails.tasks import match_email_rows_batch_task
                    match_email_rows_batch_task.delay(email.id, created_row_ids)
                except Exception as task_err:
                    logger.error(f"Error scheduling matching task: {str(task_err)}", exc_info=True)
        else:
            logger.warning(f"Attachment analysis completed for email {email.id}, but no valid data found in any attachment.")
            email.status = 'processed_nodata'
            
        email.save() # Save email status change
        
        # Log the action
        UserLog.objects.create(
            user=user,
            action_type='process_email_attachments',
            email=email,
            ip_address='127.0.0.1',  # Default for non-request context
            details=f'Processed via attachment analysis. Created {len(created_row_ids)} rows.'
        )
        
        return any_success
    
    except Exception as e:
        logger.error(f"Error processing email attachments: {e}", exc_info=True)
        return False


def parse_date_range(date_range):
    """
    Parse a date range string in format 'YYYY-MM-DD - YYYY-MM-DD'
    Returns tuple of (start_date, end_date) as datetime.date objects
    """
    try:
        if ' - ' in date_range:
            parts = date_range.split(' - ')
            start_date = datetime.strptime(parts[0], '%Y-%m-%d').date()
            end_date = datetime.strptime(parts[1], '%Y-%m-%d').date()
            return start_date, end_date
    except:
        pass
    
    today = timezone.now().date()
    return today, today + timedelta(days=7)


@login_required
def approve_row(request, row_id):
    """
    View for approving a row
    """
    from .models import EmailHotelMatch
    row = get_object_or_404(EmailRow, id=row_id)
    
    if row.status != 'pending':
        messages.error(request, "Only pending rows can be approved")
        return redirect('emails:email_detail', email_id=row.email.id)
    
    is_all_room_type = row.room_type.strip().upper() in ['ALL ROOM', 'ALL ROOMS', 'ALL ROOM TYPES']
    if not row.juniper_hotel or (not row.juniper_rooms.exists() and not is_all_room_type):
        messages.error(request, "Cannot approve row without hotel match and either a room match or 'All Room' type.")
        return redirect('emails:email_detail', email_id=row.email.id)
    
    row.status = 'approved'
    row.processed_by = request.user
    row.processed_at = timezone.now()
    row.save()
    
    UserLog.objects.create(
        user=request.user,
        action_type='approve_row',
        email=row.email,
        email_row=row,
        ip_address=request.META.get('REMOTE_ADDR')
    )
    
    # --- E-posta ve otel eşleşme veritabanına ekle/güncelle ---
    try:
        # E-posta göndericisini ayıkla
        sender_email = row.email.sender
        if '<' in sender_email and '>' in sender_email:
            # Format: "İsim Soyisim <email@domain.com>"
            sender_email = sender_email.split('<')[1].split('>')[0].strip()
        elif '@' in sender_email:
            # Format: email@domain.com
            sender_email = sender_email.strip()
            
        # Öğrenilen eşleştirmeyi kaydet/güncelle
        hotel_match, created = EmailHotelMatch.objects.get_or_create(
            sender_email=sender_email,
            hotel=row.juniper_hotel
        )
        
        if not created:
            # Kaydı güncelleyerek güven puanını artır
            hotel_match.increase_confidence()
            logger.info(f"Öğrenilen eşleştirme güncellendi: {sender_email} -> {row.juniper_hotel.juniper_hotel_name} (Güven: {hotel_match.confidence_score})")
        else:
            logger.info(f"Yeni öğrenilen eşleştirme kaydedildi: {sender_email} -> {row.juniper_hotel.juniper_hotel_name}")
            
        # Also learn hotel name to juniper hotel mapping
        learn_hotel_matching(row, request.user)
            
    except Exception as e:
        logger.error(f"Öğrenilen eşleştirme kaydı sırasında hata: {str(e)}")
    # --- E-posta ve otel eşleşme güncellemesi sonu ---
    
    # --- Market eşleşmelerini öğren ---
    if row.original_market_name and row.markets.exists():
        learn_market_matching(row, request.user)
    # --- Market eşleşme güncellemesi sonu ---
    
    # --- Kontrat eşleşmelerini öğren ---
    if row.selected_contracts:
        learn_contract_matching(row, request.user)
    # --- Kontrat eşleşme güncellemesi sonu ---
    
    # --- Oda tipi grup eşleşmelerini öğren ---
    if row.juniper_rooms.exists() and row.room_type:
        success, message = learn_room_type_group_matching(row, request.user)
        if success:
            logger.info(f"Oda tipi grup öğrenme: {message}")
    # --- Oda tipi grup eşleşme güncellemesi sonu ---
    
    email = row.email
    all_rows = email.rows.all()
    if not all_rows.filter(status__in=['pending', 'matching', 'hotel_not_found', 'room_not_found']).exists():
        if all_rows.filter(status='approved').count() == all_rows.count():
            email.status = 'approved'
            email.processed_by = request.user
            email.processed_at = timezone.now()
            email.save()
        elif all_rows.filter(status='rejected').count() == all_rows.count():
             email.status = 'rejected'
             email.processed_by = request.user
             email.processed_at = timezone.now()
             email.save()
        else:
             email.status = 'processed'
             email.processed_by = request.user
             email.processed_at = timezone.now()
             email.save()

    messages.success(request, "Row approved successfully")
    # Filtre parametrelerini koru
    query_string = request.META.get('QUERY_STRING', '')
    if query_string:
        return redirect(f"{reverse('emails:email_detail', args=[row.email.id])}?{query_string}")
    return redirect('emails:email_detail', email_id=row.email.id)


@login_required
def reject_row(request, row_id):
    """Reject a single row."""
    row = get_object_or_404(EmailRow, id=row_id)
    row.status = 'rejected'
    row.processed_by = request.user
    row.processed_at = timezone.now()
    row.save()
    
    # Log the action
    UserLog.objects.create(
        user=request.user,
        action_type='reject_row',
        email=row.email,
        email_row=row,
        ip_address=request.META.get('REMOTE_ADDR'),
        details=f"Rejected row {row_id}"
    )
    
    # Update email status if all rows have been processed
    email = row.email
    pending_rows = email.rows.filter(status__in=['pending', 'matching', 'hotel_not_found', 'room_not_found']).count()
    
    if pending_rows == 0:
        # All rows have been processed, update the email status
        email.status = 'rejected'  # Keep 'rejected' for general rejection
        email.processed_by = request.user
        email.processed_at = timezone.now()
        email.save()
    
    messages.success(request, f"Row {row_id} has been rejected.")
    return redirect('emails:email_detail', email_id=row.email.id)


@login_required
def reject_row_hotel_not_found(request, row_id):
    """Reject a single row with reason: JP Hotel Not Found."""
    row = get_object_or_404(EmailRow, id=row_id)
    row.status = 'rejected_hotel_not_found'
    row.reject_reason = 'JP Hotel Not Found'
    row.processed_by = request.user
    row.processed_at = timezone.now()
    row.save()
    
    # Log the action
    UserLog.objects.create(
        user=request.user,
        action_type='reject_row_hotel_not_found',
        email=row.email,
        email_row=row,
        ip_address=request.META.get('REMOTE_ADDR'),
        details=f"Rejected row {row_id} - JP Hotel Not Found"
    )
    
    # Update email status if all rows have been processed
    email = row.email
    pending_rows = email.rows.filter(status__in=['pending', 'matching', 'hotel_not_found', 'room_not_found']).count()
    
    if pending_rows == 0:
        # All rows have been processed, update the email status
        email.status = 'rejected_hotel_not_found'  # Use specific rejection reason
        email.processed_by = request.user
        email.processed_at = timezone.now()
        email.save()
    
    messages.success(request, f"Row {row_id} has been rejected: JP Hotel Not Found.")
    return redirect('emails:email_detail', email_id=row.email.id)


@login_required
def reject_row_room_not_found(request, row_id):
    """Reject a single row with reason: JP Room Not Found."""
    row = get_object_or_404(EmailRow, id=row_id)
    row.status = 'rejected_room_not_found'
    row.reject_reason = 'JP Room Not Found'
    row.processed_by = request.user
    row.processed_at = timezone.now()
    row.save()
    
    # Log the action
    UserLog.objects.create(
        user=request.user,
        action_type='reject_row_room_not_found',
        email=row.email,
        email_row=row,
        ip_address=request.META.get('REMOTE_ADDR'),
        details=f"Rejected row {row_id} - JP Room Not Found"
    )
    
    # Update email status if all rows have been processed
    email = row.email
    pending_rows = email.rows.filter(status__in=['pending', 'matching', 'hotel_not_found', 'room_not_found']).count()
    
    if pending_rows == 0:
        # All rows have been processed, update the email status
        email.status = 'rejected_room_not_found'  # Use specific rejection reason
        email.processed_by = request.user
        email.processed_at = timezone.now()
        email.save()
    
    messages.success(request, f"Row {row_id} has been rejected: JP Room Not Found.")
    return redirect('emails:email_detail', email_id=row.email.id)


@login_required
def send_to_robot(request, row_id):
    """
    View for sending a row to the RPA robot
    """
    row = get_object_or_404(EmailRow, id=row_id)
    
    if row.status != 'approved':
        messages.error(request, "Only approved rows can be sent to the robot")
        return redirect('emails:email_detail', email_id=row.email.id)
    
    row.status = 'sent_to_robot'
    row.save()
    
    UserLog.objects.create(
        user=request.user,
        action_type='send_to_robot',
        email=row.email,
        email_row=row,
        ip_address=request.META.get('REMOTE_ADDR')
    )
    
    messages.success(request, "Row sent to robot successfully")
    return redirect('emails:email_detail', email_id=row.email.id)


@csrf_exempt
def webhook_robot_callback(request):
    """
    Webhook endpoint for RPA robot callback
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST method is allowed'}, status=405)
    
    try:
        data = json.loads(request.body)
        row_id = data.get('row_id')
        success = data.get('success', False)
        message = data.get('message', '')
        
        if not row_id:
            return JsonResponse({'error': 'Row ID is required'}, status=400)
        
        row = get_object_or_404(EmailRow, id=row_id)
        
        if success:
            row.status = 'robot_processed'
        else:
            row.status = 'error'
        
        row.save()
        
        UserLog.objects.create(
            user=None,  # No user for webhook
            action_type='robot_callback',
            email=row.email,
            email_row=row,
            ip_address=request.META.get('REMOTE_ADDR'),
            details=message
        )
        
        return JsonResponse({'status': 'success'})
    
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    
    except Exception as e:
        logger.error(f"Error in robot callback: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def match_hotel(request, row_id):
    """
    View for matching a row to a hotel
    """
    row = get_object_or_404(EmailRow, id=row_id)
    
    hotels = Hotel.objects.all().order_by('juniper_hotel_name')
    
    # Get hotel suggestions
    best_hotel_match, hotel_suggestions = get_hotel_suggestions(row.hotel_name) # Otel önerilerini al

    if request.method == 'POST':
        hotel_id = request.POST.get('hotel_id')
        
        if hotel_id:
            try:
                hotel = Hotel.objects.get(id=hotel_id)
                row.juniper_hotel = hotel
                row.save()
                
                UserLog.objects.create(
                    user=request.user,
                    action_type='match_hotel',
                    email=row.email,
                    email_row=row,
                    ip_address=request.META.get('REMOTE_ADDR'),
                    details=f"Manual hotel match to: {hotel.juniper_hotel_name}" # Log detayını güncelle
                )
                
                messages.success(request, f"Row matched to hotel: {hotel.juniper_hotel_name}")
            except Hotel.DoesNotExist:
                messages.error(request, "Selected hotel not found")
        else:
            messages.error(request, "No hotel selected")
        
        return redirect('emails:email_detail', email_id=row.email.id)
    
    context = {
        'row': row,
        'hotels': hotels,
        'hotel_suggestions': hotel_suggestions, # Otel önerilerini context'e ekle
        'best_hotel_match': best_hotel_match, # En iyi öneriyi context'e ekle
    }
    
    return render(request, 'emails/match_hotel.html', context)


def get_room_suggestions(room_type, hotel):
    """
    E-postadaki oda tipine göre benzer odaları bulmak için algoritma.
    Otel için tanımlanmış oda tipi gruplarını kullanarak eşleşme yapar.
    
    Args:
        room_type (str): E-postadaki oda tipi
        hotel (Hotel): Seçilen otel
        
    Returns:
        tuple: (best_match, suggestions, search_pattern)
            - best_match: En yüksek puanlı oda eşleşmesi
            - suggestions: Önerilen benzer odalar listesi
            - search_pattern: Kullanılan arama deseni
    """
    from difflib import SequenceMatcher
    import re
    from hotels.models import RoomTypeGroup, RoomTypeVariant, RoomTypeGroupLearning
    import logging
    logger = logging.getLogger(__name__)
    
    # Eğer "all room/rooms" vb. bir ifade varsa, özel durum
    if room_type.lower() in ["all room", "all rooms", "all room types", "tüm odalar"]:
        return None, [], "ALL_ROOMS"
    
    # Oda tipini temizle ve arama için hazırla
    clean_room_type = room_type.strip().upper()
    
    # Öncelikle öğrenilmiş oda tipi eşleşmelerini kontrol et
    room_learnings = RoomTypeGroupLearning.objects.filter(
        hotel=hotel,
        mail_room_type__iexact=clean_room_type
    ).order_by('-confidence')
    
    if room_learnings.exists():
        # Öneri oluşturur Confidence 70% ve üzeri ise
        top_learning = room_learnings.first()
        if top_learning and top_learning.confidence >= 0.7:
            if top_learning.juniper_room:
                logger.info(f"Öğrenilen oda eşleşmesi kullanılıyor: {clean_room_type} -> {top_learning.juniper_room.juniper_room_type}")
                return top_learning.juniper_room, [top_learning.juniper_room], "LEARNING"
            elif top_learning.group:
                # Grup eşleşmesi bulunduysa, o gruptaki tüm varyantları getir
                variant_rooms = []
                for variant in top_learning.group.variants.all():
                    matching_rooms = hotel.rooms.filter(juniper_room_type__icontains=variant.variant_room_name)
                    variant_rooms.extend(matching_rooms)
                
                if variant_rooms:
                    best_match = variant_rooms[0]
                    logger.info(f"Öğrenilen oda grubu kullanılıyor: {clean_room_type} -> {top_learning.group.name}")
                    return best_match, variant_rooms, top_learning.group.name
    
    # Otele ait odaları al
    available_rooms = Room.objects.filter(hotel=hotel)
    if not available_rooms:
        return None, [], None
    
    # Otele özel oda grubu eşleşmelerini kontrol et
    room_groups = RoomTypeGroup.objects.filter(hotel=hotel)
    
    # Öncelikle, e-postadaki oda tipi ile birebir eşleşen bir RoomTypeGroup var mı diye kontrol et
    direct_group_match = room_groups.filter(name__iexact=clean_room_type).first()
    if direct_group_match:
        # Doğrudan grup eşleşmesi bulundu, bu gruptan tüm varyantları getir
        logger.info(f"Doğrudan oda grubu eşleşmesi: {clean_room_type} -> {direct_group_match.name}")
        variant_rooms = []
        
        for variant in direct_group_match.variants.all():
            # Her varyant için eşleşen odaları bul
            matching_rooms = available_rooms.filter(juniper_room_type__icontains=variant.variant_room_name)
            variant_rooms.extend(matching_rooms)
        
        if variant_rooms:
            # Varyant odaları bulunduysa, bunları öner
            return variant_rooms[0], variant_rooms, direct_group_match.name
    
    # Doğrudan eşleşme yoksa, bulanık eşleşme yapalım
    best_group_match = None
    best_group_score = 0
    
    for group in room_groups:
        score = SequenceMatcher(None, clean_room_type, group.name).ratio()
        
        # Anahtar kelimelere dayalı ek puanlar ekle
        key_words = re.findall(r'\b\w+\b', clean_room_type)
        for word in key_words:
            if len(word) > 3 and word in group.name:
                score += 0.1
        
        # En yüksek puanlı grup eşleşmesini bul
        if score > best_group_score:
            best_group_score = score
            best_group_match = group
    
    # Yeterince yüksek puanlı grup eşleşmesi varsa, bu gruptan odaları öner
    if best_group_match and best_group_score >= 0.6:
        logger.info(f"Bulanık oda grubu eşleşmesi: {clean_room_type} -> {best_group_match.name} (%{int(best_group_score*100)})")
        variant_rooms = []
        
        for variant in best_group_match.variants.all():
            matching_rooms = available_rooms.filter(juniper_room_type__icontains=variant.variant_room_name)
            variant_rooms.extend(matching_rooms)
        
        if variant_rooms:
            return variant_rooms[0], variant_rooms, best_group_match.name
    
    # Grup eşleşmesi bulunamadıysa, doğrudan odalar üzerinde eşleşme yapalım
    # En iyi eşleşen odayı bul (eski mantık)
    highest_score = 0
    best_match = None
    
    for room in available_rooms:
        # Calculate similarity score using SequenceMatcher
        juniper_room_type = room.juniper_room_type.strip().upper()
        score = SequenceMatcher(None, clean_room_type, juniper_room_type).ratio()
        
        # Update best match if score is higher
        if score > highest_score:
            highest_score = score
            best_match = room
    
    if best_match is None:
        return None, [], None
        
    # Extract pattern from best match room name
    # Look for patterns after "PAX" or "SNG" in the room name
    pattern = None
    juniper_room_name = best_match.juniper_room_type.upper()
    
    # Try to find pattern after "PAX" or "SNG"
    pax_match = re.search(r'(\d+)\s*PAX\s*(.*?)(?:\s|$)', juniper_room_name, re.IGNORECASE)
    sng_match = re.search(r'SNG\s*(.*?)(?:\s|$)', juniper_room_name, re.IGNORECASE)
    
    if pax_match:
        pattern = pax_match.group(2).strip()
    elif sng_match:
        pattern = sng_match.group(1).strip()
    
    # If no pattern found after PAX/SNG, try to use the part after room type (SUITE, ROOM, etc.)
    if not pattern or pattern == '':
        type_match = re.search(r'(SUITE|ROOM|DBL|DOUBLE|TWIN|SINGLE)\s*(.*?)(?:\s|$)', juniper_room_name, re.IGNORECASE)
        if type_match:
            pattern = type_match.group(2).strip()
    
    suggestions = []
    
    # If a pattern was found, find related rooms
    if pattern and pattern != '':
        for room in available_rooms:
            if pattern.upper() in room.juniper_room_type.upper() and room != best_match:
                suggestions.append(room)
    
    # Add the best match to the beginning of the suggestions if not already there
    if best_match and best_match not in suggestions:
        suggestions.insert(0, best_match)
    
    logger.info(f"Doğrudan oda eşleşmesi: {clean_room_type} -> {best_match.juniper_room_type} (%{int(highest_score*100)})")
    return best_match, suggestions, pattern


def propagate_room_matching(row, user=None):
    """
    After a successful room match, propagate the match to other rows in the same email with the same room type.
    
    Args:
        row: The EmailRow with a successful room match
        user: The user who initiated the match (for logging)
    
    Returns:
        int: Number of rows updated
    """
    # Only proceed if we have a successful room match
    if not row.juniper_hotel or not row.juniper_rooms.exists() or not row.room_type:
        return 0

    # Look for other rows in the same email with the same room type but no room matches
    same_type_rows = EmailRow.objects.filter(
        email=row.email,
        room_type=row.room_type,
        juniper_hotel__isnull=False  # Only match rows that already have a hotel match
    ).exclude(
        id=row.id  # Exclude the current row
    ).exclude(
        juniper_rooms__isnull=False  # Exclude rows that already have room matches
    )
    
    if not same_type_rows.exists():
        return 0
    
    # Get the matched rooms for the original row
    matched_rooms = list(row.juniper_rooms.all())
    count = 0
    
    # Apply the room matches to the other rows
    for target_row in same_type_rows:
        # Only propagate if the hotel matches or is related (same Juniper code)
        if target_row.juniper_hotel.juniper_code == row.juniper_hotel.juniper_code:
            # Set the same room matches
            target_row.juniper_rooms.set(matched_rooms)
            
            # Set status to pending if it's not already approved
            if target_row.status != 'approved':
                target_row.status = 'pending'
                
            target_row.save()
            count += 1
            
            # Log the action if a user is provided
            if user:
                UserLog.objects.create(
                    user=user,
                    action_type='propagate_room_match',
                    email=row.email,
                    email_row=target_row,
                    ip_address='0.0.0.0',  # Default when not in request context
                    details=f"Room match propagated from row {row.id} to row {target_row.id}"
                )
    
    return count


@login_required
def match_room(request, row_id):
    """
    View for matching a row to one or more rooms (multi-select)
    """
    from .models import RoomTypeMatch
    from hotels.models import RoomTypeGroup
    row = get_object_or_404(EmailRow, id=row_id)
    if not row.juniper_hotel:
        messages.error(request, "Must match hotel before matching room")
        return redirect('emails:email_detail', email_id=row.email.id)
    
    # Get all rooms for this hotel
    rooms = Room.objects.filter(hotel=row.juniper_hotel).order_by('juniper_room_type')
    
    # Check if we have a room type group match
    room_type_group = None
    if row.room_type:
        # 1. Check for direct room type group match
        clean_room_type = row.room_type.strip().upper()
        room_type_group = RoomTypeGroup.objects.filter(
            hotel=row.juniper_hotel,
            name__iexact=clean_room_type
        ).first()
        
        # 2. If not found, try case-insensitive contains match
        if not room_type_group:
            room_type_group = RoomTypeGroup.objects.filter(
                hotel=row.juniper_hotel,
                name__icontains=clean_room_type
            ).first()
    
    # If we found a room type group, get all rooms matching its variants
    group_room_suggestions = []
    if room_type_group:
        logger.info(f"Room type group match found: {room_type_group.name} for '{row.room_type}'")
        variants = room_type_group.variants.all()
        
        # Find all rooms matching any variant in the group
        for variant in variants:
            matching_rooms = Room.objects.filter(
                hotel=row.juniper_hotel,
                juniper_room_type__icontains=variant.variant_room_name
            )
            group_room_suggestions.extend(matching_rooms)
    
    # Get suggestions using the enhanced algorithm (fallback)
    best_match, suggestions, search_pattern = get_room_suggestions(row.room_type, row.juniper_hotel)
    
    # If we have group suggestions, prioritize them
    if group_room_suggestions:
        # Remove duplicates while preserving order
        seen = set()
        unique_group_suggestions = []
        for room in group_room_suggestions:
            if room.id not in seen:
                seen.add(room.id)
                unique_group_suggestions.append(room)
        
        # Use room type group suggestions, fall back to regular suggestions if empty
        suggestions = unique_group_suggestions or suggestions
        if suggestions:
            best_match = suggestions[0]
    
    if request.method == 'POST':
        room_ids = request.POST.getlist('room_ids')
        all_rooms_selected = 'all_rooms' in request.POST
        
        if all_rooms_selected and room_type_group:
            # If "all rooms in group" is selected and we have a room type group,
            # automatically select all rooms from the group's variants
            selected_rooms = []
            for variant in room_type_group.variants.all():
                variant_rooms = Room.objects.filter(
                    hotel=row.juniper_hotel,
                    juniper_room_type__icontains=variant.variant_room_name
                )
                selected_rooms.extend(variant_rooms)
            
            # Remove duplicates
            selected_rooms = list(set(selected_rooms))
            
            if selected_rooms:
                row.juniper_rooms.set(selected_rooms)
                row.status = 'pending' if row.status != 'approved' else 'approved'
                row.save()
                
                UserLog.objects.create(
                    user=request.user,
                    action_type='match_room_group',
                    email=row.email,
                    email_row=row,
                    ip_address=request.META.get('REMOTE_ADDR'),
                    details=f"Matched all variants from group: {room_type_group.name}"
                )
                
                messages.success(request, f"All rooms from group '{room_type_group.name}' matched successfully.")
                return redirect('emails:email_detail', email_id=row.email.id)
            else:
                messages.error(request, "No rooms found for the selected group variants.")
        
        elif room_ids:
            selected_rooms = Room.objects.filter(id__in=room_ids, hotel=row.juniper_hotel)
            row.juniper_rooms.set(selected_rooms)
            # Oda eşleştirmesi sonrası status 'pending' yapılır (manuel mapping ile aynı mantık)
            if row.status != 'approved':
                row.status = 'pending'
            row.save()
            # Öğrenen sistem için RoomTypeMatch kaydı
            for room in selected_rooms:
                RoomTypeMatch.objects.get_or_create(email_room_type=row.room_type, juniper_room=room)
            UserLog.objects.create(
                user=request.user,
                action_type='match_room',
                email=row.email,
                email_row=row,
                ip_address=request.META.get('REMOTE_ADDR')
            )
            
            # Propagate this room match to other rows with the same room type
            updated_count = propagate_room_matching(row, request.user)
            
            # Learn from market and contract matches if data is available
            if row.original_market_name and row.markets.exists():
                learn_market_matching(row, request.user)
            
            if row.selected_contracts:
                learn_contract_matching(row, request.user)
            
            # Oda tipi grup öğrenimi
            if row.juniper_rooms.exists() and row.room_type:
                success, message = learn_room_type_group_matching(row, request.user)
                if success:
                    logger.info(f"Oda tipi grup öğrenme: {message}")
            
            if updated_count > 0:
                messages.success(request, f"Room(s) matched successfully and applied to {updated_count} other rules with the same room type.")
            else:
                messages.success(request, "Room(s) matched successfully.")
        else:
            row.juniper_rooms.clear()
            messages.error(request, "No room selected.")
        return redirect('emails:email_detail', email_id=row.email.id)
    
    context = {
        'row': row,
        'rooms': rooms,
        'best_match': best_match,
        'suggestions': suggestions,
        'search_pattern': search_pattern,
        'room_type_group': room_type_group,
        'group_variants': room_type_group.variants.all() if room_type_group else None,
    }
    return render(request, 'emails/match_room.html', context)


@login_required
def confirm_attachment_analysis(request, email_id):
    """
    View for confirming attachment analysis results
    """
    email = get_object_or_404(Email, id=email_id)
    
    attachment_rows = email.rows.filter(from_attachment=True)
    
    if not attachment_rows.exists():
        messages.error(request, "No attachment analysis results found for this email")
        return redirect('emails:email_detail', email_id=email.id)
    
    if request.method == 'POST':
        row_ids = request.POST.getlist('row_ids')
        
        if not row_ids:
            messages.error(request, "No rows selected for confirmation")
            return redirect('emails:email_detail', email_id=email.id)
        
        selected_rows = EmailRow.objects.filter(id__in=row_ids)
        
        email.rows.filter(from_attachment=True).exclude(id__in=row_ids).delete()
        
        for row in selected_rows:
            row.status = 'pending'
            row.save()
        
        email.status = 'processed'
        email.save()
        
        UserLog.objects.create(
            user=request.user,
            action_type='confirm_attachment_analysis',
            email=email,
            ip_address=request.META.get('REMOTE_ADDR'),
            details=f"Confirmed {len(selected_rows)} rows from attachment analysis"
        )
        
        messages.success(request, f"Successfully confirmed {len(selected_rows)} rows from attachment analysis")
        return redirect('emails:email_detail', email_id=email.id)
    
    context = {
        'email': email,
        'attachment_rows': attachment_rows,
    }
    
    return render(request, 'emails/confirm_attachment_analysis.html', context)


@login_required
def manual_mapping(request, row_id):
    """
    View for manually mapping email row data
    """
    import json
    row = get_object_or_404(EmailRow, id=row_id)
    
    hotels = Hotel.objects.all().order_by('juniper_hotel_name')
    markets = Market.objects.all().order_by('name')
    
    if request.method == 'POST':
        hotel_id = request.POST.get('hotel_id')
        room_types = request.POST.getlist('room_types')
        market_ids = request.POST.getlist('market_ids')  # Çoklu pazar seçimi
        contracts = request.POST.getlist('contracts')    # Çoklu kontrat seçimi
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        sale_type = request.POST.get('sale_type')
        original_market_name = request.POST.get('original_market_name')  # Add this to capture original market name

        if not hotel_id or not room_types or not market_ids or not start_date or not end_date or not sale_type:
            messages.error(request, "All fields are required")
            return redirect('emails:manual_mapping', row_id=row.id)
            
        try:
            hotel = Hotel.objects.get(id=hotel_id)
            
            # Çoklu pazarları al
            selected_markets = Market.objects.filter(id__in=market_ids)
            if not selected_markets.exists():
                messages.error(request, "At least one market must be selected")
                return redirect('emails:manual_mapping', row_id=row.id)
            
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
            end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
            
            row.juniper_hotel = hotel
            row.hotel_name = hotel.juniper_hotel_name
            row.start_date = start_date_obj
            row.end_date = end_date_obj
            row.sale_type = sale_type
            
            # Store original market name for learning
            if original_market_name:
                row.original_market_name = original_market_name
            
            # Odaları temizle ve yeniden ayarla
            if hasattr(row, 'juniper_rooms'):
                row.juniper_rooms.clear()
                row.room_type = None 

            if 'all' in room_types:
                row.room_type = "All Room" 
            else:
                selected_rooms = Room.objects.filter(hotel=hotel, id__in=room_types)
                if selected_rooms.exists():
                    row.juniper_rooms.set(selected_rooms)
                    row.room_type = ", ".join(selected_rooms.values_list('juniper_room_type', flat=True))
                else:
                    logger.warning(f"Invalid room IDs {room_types} selected for hotel {hotel_id} in manual mapping for row {row.id}")
                    row.room_type = "Invalid Selection" 
            
            # Pazarları ayarla
            row.markets.set(selected_markets)
            
            # Kontratları kaydet (özel bir alan veya meta veri olarak)
            if contracts:
                row.selected_contracts = ", ".join(contracts)
            else:
                row.selected_contracts = None

            if row.status != 'approved': 
                row.status = 'pending'

            row.save()
            
            # Learn from market matching
            if original_market_name and selected_markets.exists():
                learn_market_matching(row, request.user)
            
            # Learn from contract matching
            if row.selected_contracts:
                learn_contract_matching(row, request.user)
            
            # Oda tipi grup öğrenimi
            if not is_all_rooms and selected_rooms.exists() and row.room_type:
                success, message = learn_room_type_group_matching(row, request.user)
                if success:
                    logger.info(f"Oda tipi grup öğrenme: {message}")
            
            UserLog.objects.create(
                user=request.user,
                action_type='manual_mapping',
                email=row.email,
                email_row=row,
                ip_address=request.META.get('REMOTE_ADDR'),
                details=f'Manual mapping completed with {len(selected_markets)} markets and {len(contracts)} contracts'
            )
            
            messages.success(request, "Row data mapped successfully")
            return redirect('emails:email_detail', email_id=row.email.id) 
    
        except (Hotel.DoesNotExist, Market.DoesNotExist, Room.DoesNotExist) as e:
            logger.error(f"Entity not found during manual mapping for row {row.id}: {e}", exc_info=True)
            messages.error(request, f"Selected entity not found: {e}")
        except ValueError as e:
            logger.error(f"Invalid date format during manual mapping for row {row.id}: {e}", exc_info=True)
            messages.error(request, f"Invalid date format: {e}")
        except Exception as e:
            logger.error(f"Error in manual mapping for row {row.id}: {e}", exc_info=True)
            messages.error(request, f"An unexpected error occurred: {str(e)}")
        
        return redirect('emails:manual_mapping', row_id=row.id)
    
    # Mevcut odaları yükle
    rooms = []
    if row.juniper_hotel:
        rooms = Room.objects.filter(hotel=row.juniper_hotel).order_by('juniper_room_type')
    
    # Seçili oda ID'lerini al
    selected_room_ids = []
    if hasattr(row, 'juniper_rooms'):
        selected_room_ids = list(row.juniper_rooms.values_list('id', flat=True))
    elif hasattr(row, 'juniper_room') and row.juniper_room:
        selected_room_ids = [row.juniper_room.id]
    
    # Seçili pazar ID'lerini al
    selected_market_ids = []
    if hasattr(row, 'markets'):
        selected_market_ids = list(row.markets.values_list('id', flat=True))
    
    # Mevcut kontratları yükle
    available_contracts = []
    selected_contracts = []
    contracts = []  # Boş liste olarak başlatıyoruz
    
    if row.juniper_hotel:
        # Otele ait kontratları yükle - SQLite için uyumlu sorgu
        contract_markets = JuniperContractMarket.objects.filter(hotel=row.juniper_hotel).values_list('contract_name', flat=True).distinct()
        available_contracts = list(contract_markets)
        
        # Seçili kontratları al (eğer varsa)
        if hasattr(row, 'selected_contracts') and row.selected_contracts:
            selected_contracts = [c.strip() for c in row.selected_contracts.split(',')]
            contracts = selected_contracts
            contracts_info_value = selected_contracts
            logger.info(f"Row {row.id}: Using manually selected contracts: {contracts}")
        else:
            # Otomatik kontrat eşleştirmesini kullan
            try:
                contracts_info = row.get_matching_contracts_info
                if isinstance(contracts_info, tuple) and len(contracts_info) > 0:
                    contract_names = contracts_info[0]
                    if contract_names and contract_names != "-":
                        contracts.append(contract_names)
                    else:
                        contracts.append("Summer 2025 EURO")  # Default değer
                        
                    # Contract info değerini belirle - eğer varsa
                    if len(contracts_info) > 1:
                        contract_count_info = contracts_info[1]
                        # Eğer contrat sayısı "X/X" formatında ise (yani tüm contratlar kullanılıyorsa)
                        is_all_contracts = False
                        
                        if contract_count_info:
                            # Parantezleri kaldır
                            cleaned_count = contract_count_info.strip('()')
                            if '/' in cleaned_count:
                                parts = cleaned_count.split('/')
                                # Eğer X/X formatında ise (ör: 1/1, 2/2, 3/3)
                                if len(parts) == 2 and parts[0] == parts[1] and parts[0].isdigit() and int(parts[0]) > 0:
                                    is_all_contracts = True
                            
                            # All contracts veya contract listesi
                            contracts_info_value = "All Contracts" if is_all_contracts else contracts
                    else:
                        contracts.append("Summer 2025 EURO")  # Default değer
                else:
                    contracts.append("Summer 2025 EURO")  # Default değer
            except Exception as e:
                logger.error(f"Error processing contracts_info for row {row.id}: {str(e)}")
                contracts.append("Summer 2025 EURO")  # Hata durumunda default değer
                contracts_info_value = "All Contracts"  # Hata durumunda default değer

    # JSON'a çevir
    selected_contracts_json = json.dumps(selected_contracts)

    context = {
        'row': row,
        'hotels': hotels,
        'rooms': rooms, 
        'markets': markets,
        'selected_room_ids': selected_room_ids,
        'selected_market_ids': selected_market_ids,
        'available_contracts': available_contracts,
        'selected_contracts': selected_contracts,  # Liste olarak gönderdik
        'selected_contracts_json': selected_contracts_json,  # JSON olarak da gönderdik
        'is_all_room': row.room_type and row.room_type.strip().upper() in ['ALL ROOM', 'ALL ROOMS', 'ALL ROOM TYPES']
    }

    return render(request, 'emails/manual_mapping.html', context)


@login_required
def get_rooms_by_hotel(request, hotel_id):
    """
    AJAX view to get rooms for a hotel
    """
    try:
        hotel = Hotel.objects.get(id=hotel_id)
        rooms = Room.objects.filter(hotel=hotel).order_by('juniper_room_type')
        
        rooms_data = [
            {'id': room.id, 'juniper_room_type': room.juniper_room_type, 'name': room.juniper_room_type, 'code': room.room_code}
            for room in rooms
        ]
        
        return JsonResponse({'rooms': rooms_data})
    
    except Hotel.DoesNotExist:
        return JsonResponse({'error': 'Hotel not found'}, status=404)
    
    except Exception as e:
        logger.error(f"Error getting rooms by hotel: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def get_contracts_by_hotel(request, hotel_id):
    """
    AJAX view to get contracts for a hotel
    """
    try:
        hotel = Hotel.objects.get(id=hotel_id)
        
        # Otele ait benzersiz kontrat isimlerini al - SQLite için uyumlu sorgu
        contracts = JuniperContractMarket.objects.filter(hotel=hotel).values_list('contract_name', flat=True).distinct()
        
        return JsonResponse({'contracts': list(contracts)})
    
    except Hotel.DoesNotExist:
        return JsonResponse({'error': 'Hotel not found'}, status=404)
    
    except Exception as e:
        logger.error(f"Error getting contracts by hotel: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


def parse_ai_date(date_str, default_date=None):
    """Parse date string from AI, trying multiple formats."""
    if not date_str:
        return default_date or timezone.now().date()
    
    formats_to_try = [
        '%d.%m.%Y',  
        '%Y-%m-%d',  
        '%d/%m/%Y',  
        '%m/%d/%Y',  
    ]
    
    for fmt in formats_to_try:
        try:
            return datetime.strptime(str(date_str), fmt).date()
        except (ValueError, TypeError):
            continue
            
    logger.warning(f"Could not parse date string '{date_str}' using formats {formats_to_try}. Using default.")
    return default_date or timezone.now().date()


@login_required
def reanalyze_email(request, email_id):
    """
    Deletes existing analysis results (EmailRows) and resets the email status
    to 'pending' to trigger automatic re-analysis via the post_save signal.
    """
    email = get_object_or_404(Email, id=email_id)
    
    try:
        with transaction.atomic():
            logger.info(f"User {request.user.username} initiated re-analysis for email {email.id} - {email.subject}")
            UserLog.objects.create(
                user=request.user,
                action_type='reanalyze_email',
                email=email,
                ip_address=request.META.get('REMOTE_ADDR'),
                details=f'Deleting {email.rows.count()} existing rows and resetting status to pending.'
            )
            
            email.rows.all().delete()
            
            email.status = 'pending'
            email.processed_by = None 
            email.processed_at = None 
            email.robot_status = 'pending' 
            email.save()
        
        messages.success(request, f"Email '{email.subject}' has been queued for re-analysis.")
    except Exception as e:
        logger.error(f"Error during re-analysis trigger for email {email.id}: {e}", exc_info=True)
        messages.error(request, f"An error occurred while trying to re-analyze the email: {str(e)}")

    return redirect('emails:email_detail', email_id=email.id)


@login_required
@require_POST
def confirm_match_ajax(request, row_id):
    row = get_object_or_404(EmailRow, id=row_id)
    row.status = 'approved'
    row.save()
    logger.info(f"Match confirmed for row {row_id} by user {request.user.username}")
    return JsonResponse({'status': 'success', 'message': 'Match confirmed successfully.'})


@login_required
def select_alternative_ajax(request, row_id):
    row = get_object_or_404(EmailRow, id=row_id)
    logger.info(f"Alternative selection requested for row {row_id} by user {request.user.username}")
    alternatives = [
        {'hotel_name': 'Alternative Hotel 1', 'hotel_code': 'ALT001', 'score': 85},
        {'hotel_name': 'Alternative Hotel 2', 'hotel_code': 'ALT002', 'score': 80},
    ]
    return JsonResponse({'status': 'info', 'message': 'Alternative selection UI/logic not fully implemented.'})


@login_required
@require_POST
def mark_not_found_ajax(request, row_id):
    row = get_object_or_404(EmailRow, id=row_id)
    row.status = 'hotel_not_found'
    row.juniper_hotel = None
    row.juniper_rooms.clear()
    row.save()
    logger.info(f"Row {row_id} marked as not found by user {request.user.username}")
    return JsonResponse({'status': 'success', 'message': 'Row marked as not found.'})


@login_required
def create_alias_ajax(request, row_id):
    row = get_object_or_404(EmailRow, id=row_id)
    logger.info(f"Alias creation requested for row {row_id} by user {request.user.username}")
    return JsonResponse({'status': 'info', 'message': 'Alias creation UI/logic not fully implemented.'})


def get_hotel_suggestions(hotel_name):
    """
    E-postadaki otel adına göre benzer otelleri bulmak için algoritma
    
    Args:
        hotel_name (str): E-postadaki otel adı
        
    Returns:
        tuple: (best_match, suggestions)
            - best_match: En yüksek puanlı otel eşleşmesi
            - suggestions: Önerilen benzer oteller listesi
    """
    from difflib import SequenceMatcher
    from hotels.models import HotelLearning
    import logging
    import re
    logger = logging.getLogger(__name__)
    
    if not hotel_name:
        return None, []
    
    # Tüm otelleri al
    available_hotels = Hotel.objects.all()
    if not available_hotels:
        return None, []
    
    # Otel adını temizle ve normalize et
    clean_hotel_name = hotel_name.strip().upper()
    
    # Otel adını normalize etmek için yardımcı fonksiyon
    def normalize_hotel_name(name):
        # Yaygın kelimeleri kaldır: HOTEL, RESORT, SPA, vb.
        common_words = ["HOTEL", "RESORT", "SPA", "PALACE", "RESIDENCE", "SUITES"]
        normalized = name
        
        # "TEST" ve "DO NOT USE" gibi etiketleri kaldır
        test_patterns = ["TEST", "DO NOT USE", "-"]
        
        # Önce yaygın otel kelimelerini kaldır
        for word in common_words:
            normalized = re.sub(r'\b' + word + r'\b', '', normalized)
        
        # Sonra test ifadelerini kaldır
        for pattern in test_patterns:
            normalized = normalized.replace(pattern, '')
            
        # Çoklu boşlukları tek boşluğa dönüştür ve kenar boşluklarını kırp
        normalized = re.sub(r'\s+', ' ', normalized).strip()
        
        return normalized
    
    # Normalize edilmiş otel adları
    norm_email_hotel = normalize_hotel_name(clean_hotel_name)
    
    # Önce öğrenme kayıtlarında önceden eşleştirilmiş bir otel adı var mı kontrol et
    hotel_learnings = HotelLearning.objects.filter(mail_hotel_name__iexact=clean_hotel_name)
    if hotel_learnings.exists():
        # Yüksek güvenilirliğe sahip bir kayıt varsa, doğrudan kullan
        top_learned_match = hotel_learnings.order_by('-confidence', '-frequency').first()
        if top_learned_match and top_learned_match.confidence >= 0.7:
            learned_hotel = top_learned_match.hotel
            learned_hotel.match_score = int(top_learned_match.confidence * 100)
            logger.info(f"HotelLearning'den bulundu: '{clean_hotel_name}' -> '{learned_hotel.juniper_hotel_name}' (%{learned_hotel.match_score})")
            return learned_hotel, [learned_hotel]
    
    # En iyi eşleşen oteli bul
    highest_score = 0
    best_match = None
    suggestions = []
    
    # Her otel için normalize edilmiş puanlama yap
    for hotel in available_hotels:
        juniper_hotel_name = hotel.juniper_hotel_name.strip().upper()
        
        # Normalize edilmiş Juniper otel adı
        norm_juniper_hotel = normalize_hotel_name(juniper_hotel_name)
        
        # Tam eşleşme durumu (orijinal isimlerle)
        if clean_hotel_name == juniper_hotel_name:
            hotel.match_score = 100
            return hotel, [hotel]
        
        # Tam eşleşme durumu (normalize edilmiş isimlerle)
        if norm_email_hotel == norm_juniper_hotel and len(norm_email_hotel) >= 3:
            hotel.match_score = 95
            logger.info(f"Normalize sonrası tam eşleşme: '{clean_hotel_name}' -> '{juniper_hotel_name}' (%95)")
            return hotel, [hotel]
        
        # Bulanık eşleşmeler için farklı yaklaşımlar
        
        # 1. Orijinal stringler üzerinde SequenceMatcher
        full_name_score = SequenceMatcher(None, clean_hotel_name, juniper_hotel_name).ratio() * 0.5  # %50 ağırlık
        
        # 2. Normalize edilmiş stringler üzerinde SequenceMatcher
        norm_name_score = SequenceMatcher(None, norm_email_hotel, norm_juniper_hotel).ratio() * 0.5  # %50 ağırlık
        
        # 3. Temel otel adı (ilk kelime) karşılaştırması
        base_score = 0
        try:
            email_base = clean_hotel_name.split(' ')[0]
            juniper_base = juniper_hotel_name.split(' ')[0]
            
            # Temel isimler tam eşleşiyorsa ve belirli bir uzunluğun üzerindeyse yüksek puan ver
            if email_base == juniper_base and len(email_base) >= 4:
                base_score = 0.4  # Temel isim eşleşmesi için yüksek bonus
        except:
            pass
        
        # 4. "TEST" ve "DO NOT USE" içeren otel isimleri için ceza puanı
        test_penalty = 0
        if "TEST" in juniper_hotel_name:
            # Eğer otelin base name'i eşleşiyorsa cezayı azalt
            if base_score > 0:
                test_penalty = 0.05  # Hafif ceza
            else:
                test_penalty = 0.2   # Daha ağır ceza
        
        # Toplam skoru hesapla
        score = full_name_score + norm_name_score + base_score - test_penalty
        
        # Skoru 0-1 aralığında sınırla
        score = max(0, min(score, 1.0))
        
        # Her Hotel nesnesi için match_score özelliğini ayarla
        hotel.match_score = int(score * 100)
        
        # Puanı yüksekse en iyi eşleşmeyi güncelle
        if score > highest_score:
            highest_score = score
            best_match = hotel
        
        # Benzerlik puanı 0.55'ten yüksekse önerilere ekle (threshold'u biraz düşür)
        if score >= 0.55:
            suggestions.append(hotel)
    
    # En iyi eşleşmeyi önerilerin başına ekle (eğer zaten eklenmemişse)
    if best_match and best_match not in suggestions and highest_score >= 0.55:
        suggestions.insert(0, best_match)
    
    # Puanlara göre sırala
    suggestions.sort(key=lambda x: x.match_score, reverse=True)
    
    # En iyi eşleşme 75% ve üzeri ise otomatik eşleştir
    if best_match and best_match.match_score >= 75:
        logger.info(f"Yüksek skorlu otomatik eşleşme: '{clean_hotel_name}' -> '{best_match.juniper_hotel_name}' (%{best_match.match_score})")
        return best_match, suggestions
    
    return best_match, suggestions


@login_required
def smart_match(request, row_id):
    """
    Tek sayfada hem otel hem de oda eşleştirmesini gerçekleştiren akıllı eşleştirme görünümü
    """
    from .models import RoomTypeMatch, EmailHotelMatch
    row = get_object_or_404(EmailRow, id=row_id)
    
    # Tüm otelleri al
    hotels = Hotel.objects.all().order_by('juniper_hotel_name')
    
    # Otel önerilerini al
    best_hotel_match, hotel_suggestions = get_hotel_suggestions(row.hotel_name)
    
    # Şu anki otele ait odaları ve seçilen odaları hazırla
    rooms = []
    selected_room_ids = []
    best_room_match = None
    room_suggestions = []
    search_pattern = None
    
    # Eğer bir otel seçilmişse oda eşleştirme önerilerini göster
    if row.juniper_hotel:
        rooms = Room.objects.filter(hotel=row.juniper_hotel).order_by('juniper_room_type')
        
        # Seçili odaları al
        if hasattr(row, 'juniper_rooms'):
            selected_room_ids = list(row.juniper_rooms.values_list('id', flat=True))
        
        # Oda önerilerini al
        best_room_match, room_suggestions, search_pattern = get_room_suggestions(row.room_type, row.juniper_hotel)
    
    if request.method == 'POST':
        hotel_id = request.POST.get('hotel_id')
        room_ids = request.POST.getlist('room_ids')
        
        if hotel_id:
            try:
                hotel = Hotel.objects.get(id=hotel_id)
                row.juniper_hotel = hotel
                
                # Eğer oda seçimi yapıldıysa
                if room_ids:
                    selected_rooms = Room.objects.filter(id__in=room_ids, hotel=hotel)
                    row.juniper_rooms.set(selected_rooms)
                    
                    # Öğrenen sistem için RoomTypeMatch kaydı
                    for room in selected_rooms:
                        RoomTypeMatch.objects.get_or_create(email_room_type=row.room_type, juniper_room=room)
                else:
                    # Oda seçilmediyse önceki seçimleri temizle
                    row.juniper_rooms.clear()
                
                # Eşleştirme sonrası durumu 'pending' yap (onay bekliyor)
                if row.status != 'approved':
                    row.status = 'pending'
                
                row.save()
                
                # Propagate this room match to other rows with the same room type
                updated_count = propagate_room_matching(row, request.user)
                
                # Learn from hotel matching
                from .models import EmailHotelMatch
                if row.email and row.email.sender and row.juniper_hotel:
                    try:
                        sender_email = row.email.sender
                        if '<' in sender_email and '>' in sender_email:
                            sender_email = sender_email.split('<')[1].split('>')[0].strip()
                        elif '@' in sender_email:
                            sender_email = sender_email.strip()
                        
                        hotel_match, created = EmailHotelMatch.objects.get_or_create(
                            sender_email=sender_email,
                            hotel=hotel
                        )
                        
                        if not created:
                            hotel_match.increase_confidence()
                            
                        # Also learn hotel name to juniper hotel mapping
                        learn_hotel_matching(row, request.user)
                    except Exception as e:
                        logger.error(f"Error learning hotel match during apply_suggestion: {str(e)}")
                
                # Learn from market matching
                if row.original_market_name and row.markets.exists():
                    learn_market_matching(row, request.user)
                
                if row.selected_contracts:
                    learn_contract_matching(row, request.user)
                
                # Oda tipi grup öğrenimi
                if row.juniper_rooms.exists() and row.room_type:
                    success, message = learn_room_type_group_matching(row, request.user)
                    if success:
                        logger.info(f"Oda tipi grup öğrenme: {message}")
                
                if updated_count > 0:
                    messages.success(request, f"Eşleştirme başarıyla tamamlandı ve aynı oda tipine sahip {updated_count} diğer kurala uygulandı.")
                else:
                    messages.success(request, "Eşleştirme başarıyla tamamlandı")
                return redirect('emails:email_detail', email_id=row.email.id)
                
            except Hotel.DoesNotExist:
                messages.error(request, "Seçilen otel bulunamadı")
        else:
            messages.error(request, "Otel seçimi yapılmadı")
    
    context = {
        'row': row,
        'hotels': hotels,
        'hotel_suggestions': hotel_suggestions,
        'best_hotel_match': best_hotel_match,
        'rooms': rooms,
        'best_room_match': best_room_match,
        'room_suggestions': room_suggestions,
        'search_pattern': search_pattern,
        'selected_room_ids': selected_room_ids,
    }
    
    return render(request, 'emails/smart_match.html', context)


@login_required
def get_rooms_by_hotel_ajax(request, hotel_id, row_id=None):
    """
    AJAX ile otel ID'sine göre oda listesini getir
    """
    import logging
    from difflib import SequenceMatcher
    logger = logging.getLogger(__name__)
    
    try:
        logger.debug(f"get_rooms_by_hotel_ajax çağrıldı: hotel_id={hotel_id}, row_id={row_id}")
        hotel = Hotel.objects.get(id=hotel_id)
        rooms = Room.objects.filter(hotel=hotel).order_by('juniper_room_type')
        
        # Eğer row_id verilmişse, bu satırın oda tipi için önerileri getir
        best_room_match = None
        room_suggestions = []
        search_pattern = None
        row = None
        
        if row_id:
            try:
                row = EmailRow.objects.get(id=row_id)
                logger.debug(f"EmailRow bulundu: {row.id}, oda tipi: {row.room_type}")
                best_room_match, room_suggestions, search_pattern = get_room_suggestions(row.room_type, hotel)
            except EmailRow.DoesNotExist:
                logger.error(f"EmailRow bulunamadı, ID: {row_id}")
                return JsonResponse({'success': False, 'message': f'Satır bulunamadı (ID: {row_id})'}, status=404)
            except Exception as e:
                logger.error(f"Oda önerileri hesaplanırken hata: {str(e)}")
                return JsonResponse({'success': False, 'message': f'Oda önerileri hesaplanırken hata: {str(e)}'}, status=500)
        
        # Odaları JSON formatına çevir
        rooms_data = []
        for room in rooms:
            match_score = 0
            if row:
                match_score = int(SequenceMatcher(None, row.room_type, room.juniper_room_type).ratio() * 100)
            
            rooms_data.append({
                'id': room.id,
                'name': room.juniper_room_type,
                'code': room.room_code,
                'is_best_match': best_room_match and room.id == best_room_match.id,
                'is_suggestion': room in room_suggestions if room_suggestions else False,
                'match_score': match_score
            })
        
        # Öneri bilgilerini oluştur
        suggestions_data = []
        if best_room_match:
            suggestions_data.append({
                'id': best_room_match.id,
                'name': best_room_match.juniper_room_type,
                'code': best_room_match.room_code,
                'is_best_match': True
            })
        
        if room_suggestions:
            for room in room_suggestions:
                if not best_room_match or room.id != best_room_match.id:
                    suggestions_data.append({
                        'id': room.id,
                        'name': room.juniper_room_type,
                        'code': room.room_code,
                        'is_best_match': False
                    })
        
        logger.debug(f"Başarılı: {len(rooms_data)} oda, {len(suggestions_data)} öneri bulundu")
        return JsonResponse({
            'success': True,
            'rooms': rooms_data,
            'best_match': {
                'id': best_room_match.id,
                'name': best_room_match.juniper_room_type,
                'code': best_room_match.room_code
            } if best_room_match else None,
            'suggestions': suggestions_data,
            'search_pattern': search_pattern
        })
    except Hotel.DoesNotExist:
        logger.error(f"Hotel bulunamadı, ID: {hotel_id}")
        return JsonResponse({'success': False, 'message': f'Otel bulunamadı (ID: {hotel_id}).'}, status=404)
    except Exception as e:
        import traceback
        logger.error(f"Odalar getirilirken hata: {str(e)}")
        logger.error(traceback.format_exc())
        return JsonResponse({'success': False, 'message': f'Beklenmeyen hata: {str(e)}'}, status=500)


@login_required
def email_bulk_approve(request):
    """
    View for approving multiple emails at once
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST method is allowed'}, status=405)
    
    email_ids = request.POST.getlist('email_ids')
    
    if not email_ids:
        messages.error(request, "No emails selected for approval")
        return redirect('emails:email_list')
    
    try:
        emails = Email.objects.filter(id__in=email_ids)
        approved_count = 0
        
        for email in emails:
            # Only process emails that are in a pending-type status
            if email.status in ['pending', 'processing', 'processed']:
                # Mark all rows as approved
                rows = email.rows.filter(status__in=['pending', 'matching', 'hotel_not_found', 'room_not_found'])
                for row in rows:
                    if row.juniper_hotel:  # Only approve rows that have at least a hotel match
                        row.status = 'approved'
                        row.processed_by = request.user
                        row.processed_at = timezone.now()
                        row.save()
                        
                        UserLog.objects.create(
                            user=request.user,
                            action_type='bulk_approve_row',
                            email=email,
                            email_row=row,
                            ip_address=request.META.get('REMOTE_ADDR')
                        )
                        
                        # Learn from hotel matching
                        from .models import EmailHotelMatch
                        if row.email and row.email.sender and row.juniper_hotel:
                            try:
                                sender_email = row.email.sender
                                if '<' in sender_email and '>' in sender_email:
                                    sender_email = sender_email.split('<')[1].split('>')[0].strip()
                                elif '@' in sender_email:
                                    sender_email = sender_email.strip()
                                
                                hotel_match, created = EmailHotelMatch.objects.get_or_create(
                                    sender_email=sender_email,
                                    hotel=row.juniper_hotel
                                )
                                
                                if not created:
                                    hotel_match.increase_confidence()
                                    
                                # Also learn hotel name to juniper hotel mapping
                                learn_hotel_matching(row, request.user)
                            except Exception as e:
                                logger.error(f"Error learning hotel match during bulk approve: {str(e)}")
                        
                        # Learn from market matching
                        if row.original_market_name and row.markets.exists():
                            learn_market_matching(row, request.user)
                        
                        # Learn from contract matching
                        if row.selected_contracts:
                            learn_contract_matching(row, request.user)
                            
                        # Oda tipi grup öğrenimi
                        if row.juniper_rooms.exists() and row.room_type:
                            success, message = learn_room_type_group_matching(row, request.user)
                            if success:
                                logger.info(f"Oda tipi grup öğrenme: {message}")
                
                # Update email status if all rows are now approved
                if not email.rows.filter(status__in=['pending', 'matching', 'hotel_not_found', 'room_not_found']).exists():
                    email.status = 'approved'
                    email.processed_by = request.user
                    email.processed_at = timezone.now()
                    email.save()
                    approved_count += 1
        
        if approved_count > 0:
            messages.success(request, f"Successfully approved {approved_count} emails")
        else:
            messages.warning(request, "No emails were approved. Check if the selected emails have valid hotel matches.")
    
    except Exception as e:
        logger.error(f"Error in bulk approve: {e}", exc_info=True)
        messages.error(request, f"An error occurred: {str(e)}")
    
    return redirect('emails:email_list')


@login_required
def email_bulk_reject(request):
    """Bulk reject selected emails."""
    if request.method == 'POST':
        try:
            # Try to parse JSON data first
            if request.content_type == 'application/json':
                data = json.loads(request.body)
                email_ids = data.get('ids', [])
            else:
                # Then try form data
                email_ids = request.POST.getlist('email_ids')
            
            emails = Email.objects.filter(id__in=email_ids)
            processed_count = 0
            
            for email in emails:
                rows = EmailRow.objects.filter(email=email)
                for row in rows:
                    row.status = 'rejected'
                    row.processed_by = request.user
                    row.processed_at = timezone.now()
                    row.save()
                
                # Update email status directly
                email.status = 'rejected'  # Keep 'rejected' for general rejection
                email.processed_by = request.user
                email.processed_at = timezone.now()
                email.save()
                processed_count += 1
                
                # Log the action
                UserLog.objects.create(
                    user=request.user,
                    action_type='bulk_reject',
                    email=email,
                    ip_address=request.META.get('REMOTE_ADDR'),
                    details=f"Bulk rejected email {email.id}"
                )
            
            # Return JSON response for AJAX requests, redirect for normal form submissions
            if request.content_type == 'application/json':
                return JsonResponse({
                    'success': True,
                    'message': f"{processed_count} emails have been rejected.",
                    'details': {
                        'success': [email.id for email in emails],
                        'errors': []
                    }
                })
            else:
                messages.success(request, f"{processed_count} emails have been rejected.")
                return redirect('emails:email_list')
                
        except json.JSONDecodeError:
            if request.content_type == 'application/json':
                return JsonResponse({'error': 'Invalid JSON'}, status=400)
            messages.error(request, "Invalid request data")
        except Exception as e:
            logger.error(f"Error in bulk reject: {e}", exc_info=True)
            if request.content_type == 'application/json':
                return JsonResponse({'error': str(e)}, status=500)
            messages.error(request, f"An error occurred: {str(e)}")
    
    return redirect('emails:email_list')


@login_required
def email_bulk_reject_hotel_not_found(request):
    """Bulk reject selected emails with reason: JP Hotel Not Found."""
    if request.method == 'POST':
        try:
            # Try to parse JSON data first
            if request.content_type == 'application/json':
                data = json.loads(request.body)
                email_ids = data.get('ids', [])
            else:
                # Then try form data
                email_ids = request.POST.getlist('email_ids')
            
            emails = Email.objects.filter(id__in=email_ids)
            processed_count = 0
            
            for email in emails:
                rows = EmailRow.objects.filter(email=email)
                for row in rows:
                    row.status = 'rejected_hotel_not_found'
                    row.reject_reason = 'JP Hotel Not Found'
                    row.processed_by = request.user
                    row.processed_at = timezone.now()
                    row.save()
                
                # Update email status directly
                email.status = 'rejected_hotel_not_found'  # Use specific status for hotel not found
                email.processed_by = request.user
                email.processed_at = timezone.now()
                email.save()
                processed_count += 1
                
                # Log the action
                UserLog.objects.create(
                    user=request.user,
                    action_type='bulk_reject_hotel_not_found',
                    email=email,
                    ip_address=request.META.get('REMOTE_ADDR'),
                    details=f"Bulk rejected email {email.id} - JP Hotel Not Found"
                )
            
            # Return JSON response for AJAX requests, redirect for normal form submissions
            if request.content_type == 'application/json':
                return JsonResponse({
                    'success': True,
                    'message': f"{processed_count} emails have been rejected: JP Hotel Not Found",
                    'details': {
                        'success': [email.id for email in emails],
                        'errors': []
                    }
                })
            else:
                messages.success(request, f"{processed_count} emails have been rejected: JP Hotel Not Found.")
                return redirect('emails:email_list')
                
        except json.JSONDecodeError:
            if request.content_type == 'application/json':
                return JsonResponse({'error': 'Invalid JSON'}, status=400)
            messages.error(request, "Invalid request data")
        except Exception as e:
            logger.error(f"Error in bulk reject hotel not found: {e}", exc_info=True)
            if request.content_type == 'application/json':
                return JsonResponse({'error': str(e)}, status=500)
            messages.error(request, f"An error occurred: {str(e)}")
    
    return redirect('emails:email_list')


@login_required
def email_bulk_reject_room_not_found(request):
    """Bulk reject selected emails with reason: JP Room Not Found."""
    if request.method == 'POST':
        try:
            # Try to parse JSON data first
            if request.content_type == 'application/json':
                data = json.loads(request.body)
                email_ids = data.get('ids', [])
            else:
                # Then try form data
                email_ids = request.POST.getlist('email_ids')
            
            emails = Email.objects.filter(id__in=email_ids)
            processed_count = 0
            
            for email in emails:
                rows = EmailRow.objects.filter(email=email)
                for row in rows:
                    row.status = 'rejected_room_not_found'
                    row.reject_reason = 'JP Room Not Found'
                    row.processed_by = request.user
                    row.processed_at = timezone.now()
                    row.save()
                
                # Update email status directly
                email.status = 'rejected_room_not_found'  # Use specific status for room not found
                email.processed_by = request.user
                email.processed_at = timezone.now()
                email.save()
                processed_count += 1
                
                # Log the action
                UserLog.objects.create(
                    user=request.user,
                    action_type='bulk_reject_room_not_found',
                    email=email,
                    ip_address=request.META.get('REMOTE_ADDR'),
                    details=f"Bulk rejected email {email.id} - JP Room Not Found"
                )
            
            # Return JSON response for AJAX requests, redirect for normal form submissions
            if request.content_type == 'application/json':
                return JsonResponse({
                    'success': True,
                    'message': f"{processed_count} emails have been rejected: JP Room Not Found",
                    'details': {
                        'success': [email.id for email in emails],
                        'errors': []
                    }
                })
            else:
                messages.success(request, f"{processed_count} emails have been rejected: JP Room Not Found.")
                return redirect('emails:email_list')
                
        except json.JSONDecodeError:
            if request.content_type == 'application/json':
                return JsonResponse({'error': 'Invalid JSON'}, status=400)
            messages.error(request, "Invalid request data")
        except Exception as e:
            logger.error(f"Error in bulk reject room not found: {e}", exc_info=True)
            if request.content_type == 'application/json':
                return JsonResponse({'error': str(e)}, status=500)
            messages.error(request, f"An error occurred: {str(e)}")
    
    return redirect('emails:email_list')


@login_required
def email_bulk_delete(request):
    """View for deleting multiple emails at once"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST method is allowed'}, status=405)
    
    email_ids = request.POST.getlist('email_ids')
    
    if not email_ids:
        messages.error(request, "No emails selected for deletion")
        return redirect('emails:email_list')
    
    try:
        emails = Email.objects.filter(id__in=email_ids)
        deleted_count = 0
        
        for email in emails:
            # Log the deletion before deleting
            UserLog.objects.create(
                user=request.user,
                action_type='delete_email',
                email=email,
                ip_address=request.META.get('REMOTE_ADDR'),
                details=f"Deleted email ID: {email.id}, Subject: {email.subject}"
            )
            
            # Delete associated rows first to avoid potential orphaned data
            email.rows.all().delete()
            
            # Delete the email
            email.delete()
            deleted_count += 1
        
        messages.success(request, f"Successfully deleted {deleted_count} emails")
    
    except Exception as e:
        logger.error(f"Error in bulk delete: {e}", exc_info=True)
        messages.error(request, f"An error occurred: {str(e)}")
    
    return redirect('emails:email_list')


@login_required
def approve_email(request, email_id):
    """
    View for approving all rows in an email
    """
    email = get_object_or_404(Email, id=email_id)
    
    if email.status != 'pending':
        messages.error(request, "Only pending emails can be approved")
        return redirect('emails:email_detail', email_id=email.id)
    
    # Mark all rows with a hotel match as approved
    rows = email.rows.filter(status__in=['pending', 'matching', 'hotel_not_found', 'room_not_found'])
    for row in rows:
        if row.juniper_hotel:  # Only approve rows that have at least a hotel match
            row.status = 'approved'
            row.processed_by = request.user
            row.processed_at = timezone.now()
            row.save()
            
            UserLog.objects.create(
                user=request.user,
                action_type='approve_row',
                email=email,
                email_row=row,
                ip_address=request.META.get('REMOTE_ADDR')
            )
            
            # Learn from hotel matching
            from .models import EmailHotelMatch
            if row.email and row.email.sender and row.juniper_hotel:
                try:
                    sender_email = row.email.sender
                    if '<' in sender_email and '>' in sender_email:
                        sender_email = sender_email.split('<')[1].split('>')[0].strip()
                    elif '@' in sender_email:
                        sender_email = sender_email.strip()
                    
                    hotel_match, created = EmailHotelMatch.objects.get_or_create(
                        sender_email=sender_email,
                        hotel=row.juniper_hotel
                    )
                    
                    if not created:
                        hotel_match.increase_confidence()
                        
                    # Also learn hotel name to juniper hotel mapping
                    learn_hotel_matching(row, request.user)
                except Exception as e:
                    logger.error(f"Error learning hotel match during email approve: {str(e)}")
            
            # Learn from market matching
            if row.original_market_name and row.markets.exists():
                learn_market_matching(row, request.user)
            
            # Learn from contract matching
            if row.selected_contracts:
                learn_contract_matching(row, request.user)
    
    # Update email status if all rows are now approved
    if not email.rows.filter(status__in=['pending', 'matching', 'hotel_not_found', 'room_not_found']).exists():
        email.status = 'approved'
        email.processed_by = request.user
        email.processed_at = timezone.now()
        email.save()
    
    messages.success(request, "All matching rows in email approved successfully")
    return redirect('emails:email_detail', email_id=email.id)


@login_required
def reject_email(request, email_id):
    """Reject all rows in an email."""
    email = get_object_or_404(Email, id=email_id)
    rows = EmailRow.objects.filter(email=email)
    
    for row in rows:
        row.status = 'rejected'
        row.processed_by = request.user
        row.processed_at = timezone.now()
        row.save()
    
    # Log the action
    UserLog.objects.create(
        user=request.user,
        action_type='reject_email',
        email=email,
        ip_address=request.META.get('REMOTE_ADDR'),
        details=f"Rejected all rows in email {email_id}"
    )
    
    # Update email status directly
    email.status = 'rejected'  # Keep 'rejected' for general rejection
    email.processed_by = request.user
    email.processed_at = timezone.now()
    email.save()
    
    messages.success(request, f"All rows in email {email_id} have been rejected.")
    
    # AJAX request için JSON yanıtı döndür, normal request için redirect
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({
            'success': True, 
            'message': f"All rows in email {email_id} have been rejected.",
            'status': email.status
        })
    else:
        return redirect('emails:email_detail', email_id=email_id)


@login_required
@require_POST
def reject_email_hotel_not_found(request, email_id):
    """Reject all rows in an email with reason: JP Hotel Not Found."""
    email = get_object_or_404(Email, id=email_id)
    
    # Block_email parametresini kontrol et - yalnızca POST'tan gelecek
    block_email = request.POST.get('block_email') == 'true'
    
    rows = EmailRow.objects.filter(email=email)
    
    for row in rows:
        row.status = 'rejected_hotel_not_found'
        row.reject_reason = 'JP Hotel Not Found'
        row.processed_by = request.user
        row.processed_at = timezone.now()
        row.save()
    
    # Log the action
    UserLog.objects.create(
        user=request.user,
        action_type='reject_email_hotel_not_found',
        email=email,
        ip_address=request.META.get('REMOTE_ADDR'),
        details=f"Rejected all rows in email {email_id} - JP Hotel Not Found" + (" (Blocked)" if block_email else "")
    )
    
    # Update email status based on block parameter
    if block_email:
        email.status = 'blocked_hotel_not_found'  # Yeni durum: bloklanmış ve otel bulunamadı
        
        # E-posta adresini blok listesine ekle
        sender_email = email.sender
        if '<' in sender_email and '>' in sender_email:
            # Format: "Name Surname <email@domain.com>" - Extract just the email part
            sender_email = sender_email.split('<')[1].split('>')[0].strip()
        elif '@' in sender_email:
            # Format: email@domain.com - Clean any extra whitespace
            sender_email = sender_email.strip()
            
        # Add to EmailBlockList if not already blocked
        from .models import EmailBlockList
        EmailBlockList.objects.get_or_create(
            sender_email=sender_email,
            defaults={
                'reason': 'JP Hotel Not Found',
                'blocked_by': request.user,
                'original_email': email,
                'is_active': True
            }
        )
        logger.info(f"Added {sender_email} to email block list due to hotel not found.")
    else:
        email.status = 'rejected_hotel_not_found'  # Normal red durumu
        
    email.processed_by = request.user
    email.processed_at = timezone.now()
    email.save()
    
    message_text = f"All rows in email {email_id} have been rejected: JP Hotel Not Found" + (" and blocked from future analysis" if block_email else "")
    messages.success(request, message_text)
    
    # Normal request için redirect
    return redirect('emails:email_detail', email_id=email_id)


@login_required
def reject_email_room_not_found(request, email_id):
    """Reject all rows in an email with reason: JP Room Not Found."""
    email = get_object_or_404(Email, id=email_id)
    rows = EmailRow.objects.filter(email=email)
    
    for row in rows:
        row.status = 'rejected_room_not_found'
        row.reject_reason = 'JP Room Not Found'
        row.processed_by = request.user
        row.processed_at = timezone.now()
        row.save()
    
    # Log the action
    UserLog.objects.create(
        user=request.user,
        action_type='reject_email_room_not_found',
        email=email,
        ip_address=request.META.get('REMOTE_ADDR'),
        details=f"Rejected all rows in email {email_id} - JP Room Not Found"
    )
    
    # Update email status directly
    email.status = 'rejected_room_not_found'
    email.processed_by = request.user
    email.processed_at = timezone.now()
    email.save()
    
    messages.success(request, f"All rows in email {email_id} have been rejected: JP Room Not Found.")
    
    # AJAX request için JSON yanıtı döndür, normal request için redirect
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({
            'success': True, 
            'message': f"All rows in email {email_id} have been rejected: JP Room Not Found",
            'status': email.status
        })
    else:
        return redirect('emails:email_detail', email_id=email_id)


@login_required
@require_POST
def apply_suggestion(request, row_id):
    """
    AI önerisini satıra uygulayan API endpoint'i
    """
    try:
        row = get_object_or_404(EmailRow, id=row_id)
        
        # Gelen veriyi parse et
        data = json.loads(request.body)
        hotel_id = data.get('hotel_id')
        room_ids = data.get('room_ids', [])
        
        # Hotel ve Room modellerini import et
        from hotels.models import Hotel, Room
        
        # Oteli bul ve satıra ekle
        hotel = get_object_or_404(Hotel, id=hotel_id)
        row.juniper_hotel = hotel
        
        # Odaları ekle (eğer mevcutsa)
        if room_ids:
            rooms = Room.objects.filter(id__in=room_ids, hotel=hotel)
            row.juniper_rooms.set(rooms)
            
            # Oda tipini güncelle
            room_names = [room.juniper_room_type for room in rooms]
            row.room_type = ", ".join(room_names)
        
        # Durumu pending olarak güncelle
        row.status = 'pending'
        row.save()
        
        # İşlemi kaydet
        UserLog.objects.create(
            user=request.user,
            action_type='apply_suggestion',
            email=row.email,
            email_row=row,
            ip_address=request.META.get('REMOTE_ADDR'),
            details=f"AI önerisi uygulandı - Otel: {hotel.juniper_hotel_name}"
        )
        
        # Propagate this room match to other rows with the same room type
        updated_count = propagate_room_matching(row, request.user)
        
        # Learn from hotel matching
        from .models import EmailHotelMatch
        if row.email and row.email.sender and row.juniper_hotel:
            try:
                sender_email = row.email.sender
                if '<' in sender_email and '>' in sender_email:
                    sender_email = sender_email.split('<')[1].split('>')[0].strip()
                elif '@' in sender_email:
                    sender_email = sender_email.strip()
                
                hotel_match, created = EmailHotelMatch.objects.get_or_create(
                    sender_email=sender_email,
                    hotel=row.juniper_hotel
                )
                
                if not created:
                    hotel_match.increase_confidence()
                    
                # Also learn hotel name to juniper hotel mapping
                learn_hotel_matching(row, request.user)
            except Exception as e:
                logger.error(f"Error learning hotel match during apply_suggestion: {str(e)}")
        
        # Learn from market matching
        if row.original_market_name and row.markets.exists():
            learn_market_matching(row, request.user)
        
        # Learn from contract matching
        if row.selected_contracts:
            learn_contract_matching(row, request.user)
        
        # Success message with propagation info
        message = 'Öneri başarıyla uygulandı'
        if updated_count > 0:
            message += f' ve aynı oda tipine sahip {updated_count} diğer kurala yayıldı'
        
        # Başarılı yanıt
        return JsonResponse({
            'success': True,
            'message': message,
            'hotel': hotel.juniper_hotel_name,
            'rooms': [room.juniper_room_type for room in rooms] if room_ids else [],
            'propagated_count': updated_count
        })
    
    except json.JSONDecodeError:
        logger.error("JSON parse error in apply_suggestion")
        return JsonResponse({'success': False, 'error': 'Geçersiz JSON formatı'}, status=400)
    
    except Exception as e:
        logger.error(f"Error in apply_suggestion: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
def check_new_emails_api(request):
    """
    API endpoint to check for new emails since a given timestamp
    """
    try:
        last_check_time = request.GET.get('last_check')
        if not last_check_time:
            # If no timestamp is provided, default to checking the last 5 minutes
            last_check_time = (timezone.now() - timedelta(minutes=5)).isoformat()
        
        # Parse the timestamp
        try:
            last_check_datetime = datetime.fromisoformat(last_check_time)
            # Ensure it's timezone aware
            if last_check_datetime.tzinfo is None:
                last_check_datetime = timezone.make_aware(last_check_datetime)
        except (ValueError, TypeError):
            # Invalid timestamp format, default to 5 minutes ago
            last_check_datetime = timezone.now() - timedelta(minutes=5)
        
        # Query for new emails received after the timestamp
        new_emails = Email.objects.filter(received_date__gt=last_check_datetime).order_by('-received_date')
        
        # Prepare the response data
        emails_data = []
        for email in new_emails:
            emails_data.append({
                'id': email.id,
                'subject': email.subject,
                'sender': email.sender,
                'received_date': email.received_date.isoformat(),
                'status': email.status,
                'has_attachments': email.has_attachments,
                'url': reverse('emails:email_detail', args=[email.id])
            })
        
        # Return the response with new emails and current server time
        return JsonResponse({
            'success': True,
            'new_emails_count': len(emails_data),
            'new_emails': emails_data,
            'server_time': timezone.now().isoformat()
        })
    
    except Exception as e:
        logger.error(f"Error in check_new_emails_api: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e),
            'server_time': timezone.now().isoformat()
        }, status=500)


@login_required
def bulk_action_rows(request, action):
    """
    Handle bulk actions on multiple rows (rules) at once.
    Supports:
    - approve: Approve selected rows
    - reject: Reject selected rows with general rejection
    - reject-hotel-not-found: Reject selected rows with JP Hotel Not Found reason
    - reject-room-not-found: Reject selected rows with JP Room Not Found reason
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST method is allowed'}, status=405)
    
    try:
        # Parse request data based on content type
        if request.content_type == 'application/json':
            data = json.loads(request.body)
            row_ids = data.get('ids', [])
            block_email = data.get('block_email', False)  # Get block_email flag
        else:
            row_ids = request.POST.getlist('row_ids')
            block_email = request.POST.get('block_email') == 'true'  # Get block_email flag
        
        if not row_ids:
            return JsonResponse({'error': 'No rows selected'}, status=400)
        
        # Get selected rows
        rows = EmailRow.objects.filter(id__in=row_ids)
        
        # Apply action based on parameter
        processed_count = 0
        emails_updated = set()  # Keep track of emails to update
        
        for row in rows:
            try:
                # Apply specific action to each row
                if action == 'approve':
                    if row.juniper_hotel:  # Only approve rows with a hotel match
                        row.status = 'approved'
                        row.processed_by = request.user
                        row.processed_at = timezone.now()
                        processed_count += 1
                elif action == 'reject':
                    row.status = 'rejected'
                    row.processed_by = request.user
                    row.processed_at = timezone.now()
                    processed_count += 1
                elif action == 'reject-hotel-not-found':
                    row.status = 'rejected_hotel_not_found'
                    row.reject_reason = 'JP Hotel Not Found'
                    row.processed_by = request.user
                    row.processed_at = timezone.now()
                    processed_count += 1
                elif action == 'reject-room-not-found':
                    row.status = 'rejected_room_not_found'
                    row.reject_reason = 'JP Room Not Found'
                    row.processed_by = request.user
                    row.processed_at = timezone.now()
                    processed_count += 1
                else:
                    continue  # Skip unsupported actions
                
                row.save()
                
                # Log the action
                UserLog.objects.create(
                    user=request.user,
                    action_type=f'bulk_{action}',
                    email=row.email,
                    email_row=row,
                    ip_address=request.META.get('REMOTE_ADDR'),
                    details=f"Bulk action '{action}' on row {row.id}" + (" (Blocked)" if action == 'reject-hotel-not-found' and block_email else "")
                )
                
                # Add this row's email to the set of emails to update
                emails_updated.add(row.email)
                
            except Exception as e:
                logger.error(f"Error processing row {row.id} with action {action}: {e}", exc_info=True)
        
        # Now update all affected emails' statuses
        for email in emails_updated:
            # Check remaining rows status to decide the email status
            pending_rows = email.rows.filter(status__in=['pending', 'matching', 'hotel_not_found', 'room_not_found']).count()
            
            if pending_rows == 0:
                # All rows have been processed
                if action == 'approve':
                    if email.rows.filter(status='approved').count() == email.rows.count():
                        email.status = 'approved'
                elif action == 'reject':
                    email.status = 'rejected'
                elif action == 'reject-hotel-not-found':
                    # If block_email is true, use the blocked status
                    if block_email:
                        email.status = 'blocked_hotel_not_found'
                        
                        # E-posta adresini blok listesine ekle
                        sender_email = email.sender
                        if '<' in sender_email and '>' in sender_email:
                            # Format: "Name Surname <email@domain.com>" - Extract just the email part
                            sender_email = sender_email.split('<')[1].split('>')[0].strip()
                        elif '@' in sender_email:
                            # Format: email@domain.com - Clean any extra whitespace
                            sender_email = sender_email.strip()
                            
                        # Add to EmailBlockList if not already blocked
                        from .models import EmailBlockList
                        EmailBlockList.objects.get_or_create(
                            sender_email=sender_email,
                            defaults={
                                'reason': 'JP Hotel Not Found',
                                'blocked_by': request.user,
                                'original_email': email,
                                'is_active': True
                            }
                        )
                        logger.info(f"Added {sender_email} to email block list during bulk action.")
                    else:
                        email.status = 'rejected_hotel_not_found'
                elif action == 'reject-room-not-found':
                    email.status = 'rejected_room_not_found'
                
                email.processed_by = request.user
                email.processed_at = timezone.now()
                email.save()
        
        # Prepare success response
        action_display = {
            'approve': 'approved',
            'reject': 'rejected',
            'reject-hotel-not-found': 'rejected: JP Hotel Not Found' + (" and blocked from future analysis" if block_email else ""),
            'reject-room-not-found': 'rejected: JP Room Not Found'
        }.get(action, action)
        
        return JsonResponse({
            'success': True,
            'message': f"{processed_count} rows have been {action_display}.",
            'details': {
                'success': [r.id for r in rows if r.status == action.replace('-', '_') or 
                            (action == 'approve' and r.status == 'approved')],
                'errors': []
            }
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        logger.error(f"Error in bulk action {action}: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


# Belirli bir ekin görüntüleme/indirme fonksiyonunu bulmak için
@login_required
def download_attachment(request, attachment_id):
    """
    View for downloading an attachment
    """
    attachment = get_object_or_404(EmailAttachment, id=attachment_id)
    
    if not attachment.file:
        messages.error(request, "Attachment file not found")
        return redirect('emails:email_detail', email_id=attachment.email.id)
    
    # Check if file exists on disk
    if not os.path.exists(attachment.file.path):
        messages.error(request, "Attachment file not found on disk")
        return redirect('emails:email_detail', email_id=attachment.email.id)
    
    # Log attachment access
    UserLog.objects.create(
        user=request.user,
        action_type='download_attachment',
        email=attachment.email,
        ip_address=request.META.get('REMOTE_ADDR'),
        details=f"Downloaded attachment: {attachment.filename}"
    )
    
    # Create response with file
    file_path = attachment.file.path
    with open(file_path, 'rb') as f:
        file_data = f.read()
        
    # Determine content type
    content_type = attachment.content_type
    if not content_type:
        content_type = 'application/octet-stream'  # Default content type
    
    # Determine if this file is analyzed (only PDF and Word files are analyzed)
    file_ext = os.path.splitext(attachment.filename.lower())[1]
    is_analyzed = file_ext in ['.pdf', '.doc', '.docx']
    
    # Set filename
    response = HttpResponse(file_data, content_type=content_type)
    response['Content-Disposition'] = f'attachment; filename="{attachment.filename}"'
    
    return response


def process_html_for_display(email):
    """Process HTML content to make it displayable in a browser"""
    if not email.body_html:
        return None
        
    # Process CID references if attachments exist
    html_content = email.body_html
    attachments = email.attachments.all()
    
    if attachments:
        import re
        
        # Create a mapping of Content-IDs to attachment URLs
        for attachment in attachments:
            # Create a URL to the attachment
            attachment_id = attachment.id
            attachment_url = reverse('emails:email_attachment_view', args=[attachment_id])
            
            # Try to replace CID references with the attachment URL using different patterns
            patterns = []
            
            # Use the stored content_id if available
            if attachment.content_id:
                patterns.append(f'src=["\']cid:{re.escape(attachment.content_id)}["\']')
            
            # Filename based pattern (fallback)
            filename = attachment.filename
            patterns.append(f'src=["\']cid:{re.escape(filename)}["\']')
            
            # Generic pattern for image CIDs (last resort)
            patterns.append(f'src=["\']cid:image[^"\']*["\']')
            
            # Apply all patterns
            for pattern in patterns:
                html_content = re.sub(pattern, f'src="{attachment_url}"', html_content)
    
    # Add base target to open links in new tab
    html_content = html_content.replace('<head>', '<head><base target="_blank">')
    
    # Handle character encoding issues
    html_content = html_content.replace('\x00', '')
    
    return html_content


@login_required
def mark_email_juniper_manual(request, email_id):
    """Mark email as processed manually in Juniper"""
    email = get_object_or_404(Email, id=email_id)
    
    # Update the email status
    email.status = 'juniper_manual'
    email.processed_by = request.user
    email.processed_at = timezone.now()
    email.save()
    
    # Log the action
    UserLog.objects.create(
        user=request.user,
        action_type='mark_juniper_manual',
        email=email,
        ip_address=request.META.get('REMOTE_ADDR'),
        details=f"Marked email {email_id} as Juniper(M) - manually processed in Juniper"
    )
    
    messages.success(request, f"Email has been marked as Juniper(M) - manually processed in Juniper.")
    
    # AJAX request için JSON yanıtı döndür, normal request için redirect
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({
            'success': True, 
            'message': f"Email has been marked as Juniper(M).",
            'status': email.status
        })
    else:
        return redirect('emails:email_detail', email_id=email_id)


@login_required
def mark_email_juniper_robot(request, email_id):
    """Mark email as processed by robot in Juniper"""
    email = get_object_or_404(Email, id=email_id)
    
    # Update the email status
    email.status = 'juniper_robot'
    email.processed_by = request.user
    email.processed_at = timezone.now()
    email.save()
    
    # Log the action
    UserLog.objects.create(
        user=request.user,
        action_type='mark_juniper_robot',
        email=email,
        ip_address=request.META.get('REMOTE_ADDR'),
        details=f"Marked email {email_id} as Juniper(R) - processed by robot in Juniper"
    )
    
    messages.success(request, f"Email has been marked as Juniper(R) - processed by robot in Juniper.")
    
    # AJAX request için JSON yanıtı döndür, normal request için redirect
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({
            'success': True, 
            'message': f"Email has been marked as Juniper(R).",
            'status': email.status
        })
    else:
        return redirect('emails:email_detail', email_id=email_id)


@login_required
def export_rules_for_robot(request, email_id):
    """
    Onaylanan e-postadaki kuralları robot için JSON formatında bir dosyaya export eder.
    Dosya adı: robot_rules_EMAIL_ID_COUNT.json formatında olur.
    """
    import json
    import os
    import traceback
    from django.conf import settings
    
    try:
        email = get_object_or_404(Email, id=email_id)
        
        # E-posta onaylanmış mı kontrol et
        if email.status != 'approved':
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': False,
                    'error': 'Bu e-posta onaylanmamış. Önce e-postayı onaylamanız gerekiyor.'
                })
            else:
                messages.error(request, "Bu e-posta onaylanmamış. Önce e-postayı onaylamanız gerekiyor.")
                return redirect('emails:email_detail', email_id=email.id)
        
        # Onaylanan satırları al
        approved_rows = email.rows.filter(status='approved')
        
        if not approved_rows.exists():
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': False,
                    'error': 'Bu e-postada onaylanmış kural bulunamadı.'
                })
            else:
                messages.error(request, "Bu e-postada onaylanmış kural bulunamadı.")
                return redirect('emails:email_detail', email_id=email.id)
        
        # Robot için JSON kuralları oluştur
        rules = []
        
        for row in approved_rows:
            try:
                # Juniper otel ve oda bilgilerini kontrol et
                if not row.juniper_hotel:
                    continue
                
                # Tüm odalar kontrolü
                is_all_room_type = row.room_type.strip().upper() in ['ALL ROOM', 'ALL ROOMS', 'ALL ROOM TYPES']
                
                # Eğer tüm odalar değilse ve oda eşleşmesi yoksa, atla
                if not is_all_room_type and not row.juniper_rooms.exists():
                    continue
                
                # Oda tipleri listesi
                room_types = []
                if is_all_room_type:
                    room_types.append("ALL")
                else:
                    for room in row.juniper_rooms.all():
                        room_types.append(room.juniper_room_type)
                
                # Pazarları al
                markets = ["ALL"]
                if row.markets.exists():
                    markets = []
                    for market in row.markets.all():
                        if market.juniper_code:
                            markets.append(market.juniper_code)
                        else:
                            markets.append(market.name)
                
                # Kontrat bilgilerini al
                contracts = []
                contracts_info_value = "All Contracts"  # Default değer
                
                # Önce selected_contracts alanını kontrol et
                if hasattr(row, 'selected_contracts') and row.selected_contracts:
                    # Manuel olarak seçilmiş kontratları kullan
                    selected_contracts = [c.strip() for c in row.selected_contracts.split(',')]
                    contracts = selected_contracts
                    contracts_info_value = selected_contracts
                    logger.info(f"Row {row.id}: Using manually selected contracts: {contracts}")
                else:
                    # Otomatik kontrat eşleştirmesini kullan
                    try:
                        contracts_info = row.get_matching_contracts_info
                        if isinstance(contracts_info, tuple) and len(contracts_info) > 0:
                            contract_names = contracts_info[0]
                            if contract_names and contract_names != "-":
                                contracts.append(contract_names)
                            else:
                                contracts.append("Summer 2025 EURO")  # Default değer
                                
                            # Contract info değerini belirle - eğer varsa
                            if len(contracts_info) > 1:
                                contract_count_info = contracts_info[1]
                                # Eğer contrat sayısı "X/X" formatında ise (yani tüm contratlar kullanılıyorsa)
                                is_all_contracts = False
                                
                                if contract_count_info:
                                    # Parantezleri kaldır
                                    cleaned_count = contract_count_info.strip('()')
                                    if '/' in cleaned_count:
                                        parts = cleaned_count.split('/')
                                        # Eğer X/X formatında ise (ör: 1/1, 2/2, 3/3)
                                        if len(parts) == 2 and parts[0] == parts[1] and parts[0].isdigit() and int(parts[0]) > 0:
                                            is_all_contracts = True
                                
                                # All contracts veya contract listesi
                                contracts_info_value = "All Contracts" if is_all_contracts else contracts
                        else:
                            contracts.append("Summer 2025 EURO")  # Default değer
                    except Exception as e:
                        logger.error(f"Error processing contracts_info for row {row.id}: {str(e)}")
                        contracts.append("Summer 2025 EURO")  # Hata durumunda default değer
                        contracts_info_value = "All Contracts"  # Hata durumunda default değer
                    
                # Satış tipini belirle (stop/open)
                sale_type = row.sale_type.lower() if row.sale_type else 'stop'
                
                # Juniper linkini oluştur
                juniper_link = ''
                if sale_type == 'stop':
                    juniper_link = f"https://juniper.bedboxx.com/intranet/cupo/StopSalesMesv2.aspx?alojamiento={row.juniper_hotel.juniper_code}&VoC=C&idCon=0&idIframePadre=framev1747331566914&idVentana=v1747331569096&idIframe=framev1747331569096#"
                else:  # open sale
                    juniper_link = f"https://juniper.bedboxx.com/intranet/cupo/StopSalesMesv2.aspx?alojamiento={row.juniper_hotel.juniper_code}&VoC=C&idIframePadre=framev1747314523196&abrir=si&idVentana=v1747314527049&idIframe=framev1747314527049#"
                
                # Kural bilgilerini oluştur
                rule = {
                    "email_id": email.id,
                    "juniper_hotel_name": row.juniper_hotel.juniper_hotel_name,
                    "juniper_hotel_code": row.juniper_hotel.juniper_code,
                    "start_date": row.start_date.strftime('%d/%m/%Y') if row.start_date else None,
                    "end_date": row.end_date.strftime('%d/%m/%Y') if row.end_date else None,
                    "sale_type": sale_type,
                    "juniper_room_types": room_types,
                    "markets": markets,
                    "contracts": contracts,
                    "contracts_info": contracts_info_value,
                    "juniper_link": juniper_link
                }
                
                rules.append(rule)
            except Exception as row_error:
                logger.error(f"Error processing row {row.id} for export: {str(row_error)}")
                logger.error(traceback.format_exc())
                # Continue processing other rows
                continue
        
        if not rules:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': False,
                    'error': 'Geçerli kural bulunamadı. Lütfen tüm kuralların doğru şekilde eşleştirildiğinden emin olun.'
                })
            else:
                messages.error(request, "Geçerli kural bulunamadı. Lütfen tüm kuralların doğru şekilde eşleştirildiğinden emin olun.")
                return redirect('emails:email_detail', email_id=email.id)
        
        # Son JSON formatını oluştur
        final_json = {
            "rules": rules
        }
        
        # JSON dosyasını kaydet
        json_dir = getattr(settings, 'ROBOT_RULES_PATH', os.path.join(settings.BASE_DIR.parent, 'json_files'))
        os.makedirs(json_dir, exist_ok=True)
        
        # Dosya adını oluştur: robot_rules_EMAIL_ID_COUNT.json
        file_name = f"robot_rules_{email_id}_{len(rules)}.json"
        file_path = os.path.join(json_dir, file_name)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(final_json, f, indent=2, ensure_ascii=False)
        
        # E-postanın durumunu güncelle
        email.robot_status = 'processing'
        email.status = 'sent_to_robot'  # Status'u 'sent_to_robot' olarak güncelle
        email.save()
        
        # İşlem kaydını oluştur
        UserLog.objects.create(
            user=request.user,
            action_type='export_rules_for_robot',
            email=email,
            ip_address=request.META.get('REMOTE_ADDR'),
            details=f"Exported {len(rules)} rules to {file_name}"
        )
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({
                'success': True,
                'message': f"{len(rules)} kural başarıyla {file_name} dosyasına kaydedildi.",
                'file_name': file_name,
                'rules_count': len(rules)
            })
        else:
            messages.success(request, f"{len(rules)} kural başarıyla {file_name} dosyasına kaydedildi.")
            return redirect('emails:email_detail', email_id=email.id)
            
    except Exception as e:
        logger.error(f"Error in export_rules_for_robot for email {email_id}: {str(e)}")
        logger.error(traceback.format_exc())
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({
                'success': False,
                'error': f"Kurallar export edilirken bir hata oluştu: {str(e)}"
            }, status=500)
        else:
            messages.error(request, f"Kurallar export edilirken bir hata oluştu: {str(e)}")
            return redirect('emails:email_detail', email_id=email.id)


def learn_market_matching(email_row, user=None):
    """
    Create or update EmailMarketMatch entries based on manual matching done by users.
    This function should be called whenever a row with market data is approved or updated.
    
    Args:
        email_row (EmailRow): The email row containing market data
        user (User, optional): The user who initiated the action (for logging)
    
    Returns:
        int: Number of market matches created or updated
    """
    from .models import EmailMarketMatch
    import logging
    logger = logging.getLogger(__name__)
    
    # Skip if row doesn't have an original market name or no markets are associated
    if not email_row.original_market_name or not email_row.markets.exists():
        return 0
    
    count = 0
    original_market_name = email_row.original_market_name.strip().upper()
    
    # Process each matched market
    for juniper_market in email_row.markets.all():
        try:
            # Get or create the market match record
            market_match, created = EmailMarketMatch.objects.get_or_create(
                email_market_name=original_market_name,
                juniper_market=juniper_market
            )
            
            # Update the confidence score for existing matches
            if not created:
                market_match.increase_confidence()
                logger.info(f"Updated market match: {original_market_name} -> {juniper_market.name} (Score: {market_match.confidence_score})")
            else:
                logger.info(f"Created new market match: {original_market_name} -> {juniper_market.name}")
                
            count += 1
            
        except Exception as e:
            logger.error(f"Error learning market match: {str(e)}")
    
    return count


def learn_contract_matching(email_row, user=None):
    """
    Create or update EmailContractMatch entries based on manual contract selections.
    This function should be called whenever contracts are manually selected for a row.
    
    Args:
        email_row (EmailRow): The email row containing selected contracts
        user (User, optional): The user who initiated the action (for logging)
    
    Returns:
        bool: True if the learning was successful, False otherwise
    """
    from .models import EmailContractMatch
    import logging
    logger = logging.getLogger(__name__)
    
    # Skip if row doesn't have the necessary data
    if not email_row.juniper_hotel or not email_row.selected_contracts:
        return False
    
    try:
        # Format contract names as a comma-separated string
        contract_names = email_row.selected_contracts.strip()
        
        # Look for an existing match with similar parameters
        existing_match = EmailContractMatch.objects.filter(
            source_hotel_name=email_row.hotel_name,
            source_room_type=email_row.room_type,
            juniper_hotel=email_row.juniper_hotel,
            matched_contracts=contract_names
        ).first()
        
        if existing_match:
            # Update existing match
            existing_match.increase_confidence()
            logger.info(f"Updated contract match for {email_row.hotel_name}/{email_row.room_type} -> {contract_names}")
            
            # Update the ManyToMany relationships if they've changed
            if email_row.juniper_rooms.exists():
                existing_match.juniper_rooms.set(email_row.juniper_rooms.all())
            
            if email_row.markets.exists():
                existing_match.juniper_markets.set(email_row.markets.all())
                
                # If we have market data, update the source_market_names field
                market_names = ", ".join([m.name for m in email_row.markets.all()])
                if market_names:
                    existing_match.source_market_names = market_names
                    existing_match.save()
        else:
            # Create new match
            new_match = EmailContractMatch.objects.create(
                email_row_sample=email_row,
                source_hotel_name=email_row.hotel_name,
                source_room_type=email_row.room_type,
                source_market_names=", ".join([m.name for m in email_row.markets.all()]) if email_row.markets.exists() else None,
                juniper_hotel=email_row.juniper_hotel,
                matched_contracts=contract_names
            )
            
            # Add many-to-many relationships
            if email_row.juniper_rooms.exists():
                new_match.juniper_rooms.set(email_row.juniper_rooms.all())
            
            if email_row.markets.exists():
                new_match.juniper_markets.set(email_row.markets.all())
                
            logger.info(f"Created new contract match for {email_row.hotel_name}/{email_row.room_type} -> {contract_names}")
        
        return True
        
    except Exception as e:
        logger.error(f"Error learning contract match: {str(e)}")
        return False

# Add this new function after learn_contract_matching at the end of the file
def learn_hotel_matching(email_row, user=None):
    """
    Create or update HotelLearning entries based on manual hotel matching done by users.
    This function creates/updates hotel name to Juniper hotel mappings.
    
    Args:
        email_row (EmailRow): The email row containing hotel data
        user (User, optional): The user who initiated the action (for logging)
    
    Returns:
        bool: True if the learning was successful, False otherwise
    """
    from hotels.models import HotelLearning
    import logging
    logger = logging.getLogger(__name__)
    
    # Skip if row doesn't have hotel name or juniper hotel
    if not email_row.hotel_name or not email_row.juniper_hotel:
        return False
    
    try:
        # Clean the hotel name
        clean_hotel_name = email_row.hotel_name.strip().upper()
        
        # Get or create the hotel learning record
        hotel_learning, created = HotelLearning.objects.get_or_create(
            mail_hotel_name=clean_hotel_name,
            hotel=email_row.juniper_hotel
        )
        
        # Update the confidence score
        if not created:
            hotel_learning.increase_confidence()
            logger.info(f"Updated hotel learning: {clean_hotel_name} -> {email_row.juniper_hotel.juniper_hotel_name} (Score: {hotel_learning.confidence})")
        else:
            logger.info(f"Created new hotel learning: {clean_hotel_name} -> {email_row.juniper_hotel.juniper_hotel_name}")
        
        return True
    
    except Exception as e:
        logger.error(f"Error learning hotel match: {str(e)}")
        return False


@login_required
def get_attachment_content(request, attachment_id):
    """
    API endpoint to get formatted content of an attachment
    """
    attachment = get_object_or_404(EmailAttachment, id=attachment_id)
    
    # Check if file exists
    if not attachment.file or not os.path.exists(attachment.file.path):
        return JsonResponse({'error': 'File not found'}, status=404)
    
    # Use decoded filename and file_extension property for better file type detection
    html_content = ""
    error = None
    
    try:
        # Handle different file types using the attachment's properties
        if attachment.is_word:
            # Process Word document
            doc = Document(attachment.file.path)
            paragraphs = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(f"<p>{para.text}</p>")
            
            # Process tables (limited to first 100 rows)
            tables_html = []
            for table in doc.tables:
                rows_html = []
                row_count = 0
                
                for row in table.rows:
                    if row_count >= 100:  # Limit to 100 rows
                        rows_html.append('<tr><td colspan="100%">... (additional rows omitted) ...</td></tr>')
                        break
                    
                    cells_html = []
                    for cell in row.cells:
                        cells_html.append(f"<td>{cell.text}</td>")
                    
                    rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
                    row_count += 1
                
                tables_html.append(f'<table class="table table-bordered table-sm my-3">{"".join(rows_html)}</table>')
            
            html_content = ''.join(paragraphs) + ''.join(tables_html)
            
        elif attachment.is_excel:
            # Process Excel spreadsheet
            df = pd.read_excel(attachment.file.path, sheet_name=None)  # Read all sheets
            
            sheets_html = []
            for sheet_name, sheet_data in df.items():
                # Limit to 100 rows per sheet
                if len(sheet_data) > 100:
                    sheet_data = sheet_data.head(100)
                    has_more = True
                else:
                    has_more = False
                
                # Convert to HTML
                table_html = sheet_data.to_html(
                    classes='table table-bordered table-sm table-striped',
                    index=False,
                    na_rep='-'
                )
                
                # Add sheet name and more rows indicator
                sheet_html = f'<div class="sheet-container mb-4">'
                sheet_html += f'<h5 class="sheet-name bg-light p-2 border rounded">{sheet_name}</h5>'
                sheet_html += table_html
                if has_more:
                    sheet_html += '<div class="text-center text-muted my-2">... (showing first 100 rows) ...</div>'
                sheet_html += '</div>'
                
                sheets_html.append(sheet_html)
            
            html_content = ''.join(sheets_html)
        
        elif attachment.file_extension == '.csv':
            # Process CSV file
            df = pd.read_csv(attachment.file.path)
            
            # Limit to 100 rows
            if len(df) > 100:
                df = df.head(100)
                has_more = True
            else:
                has_more = False
            
            # Convert to HTML
            html_content = df.to_html(
                classes='table table-bordered table-sm table-striped',
                index=False,
                na_rep='-'
            )
            
            if has_more:
                html_content += '<div class="text-center text-muted my-2">... (showing first 100 rows) ...</div>'
        
        else:
            error = f"Unsupported file format for HTML preview: {attachment.file_extension}"
            
    except Exception as e:
        error = f"Error processing file: {str(e)}"
        import traceback
        print(traceback.format_exc())
    
    if error:
        return JsonResponse({'error': error}, status=400)
    
    return JsonResponse({'html_content': html_content})


@login_required
def analyze_email_attachments(request, email_id):
    """
    View for analyzing email attachments and extracting stop sale/open sale information
    """
    email = get_object_or_404(Email, id=email_id)
    
    if not email.has_attachments:
        messages.error(request, "This email has no attachments to analyze")
        return redirect('emails:email_detail', email_id=email.id)
    
    # Log more details about the attachments for debugging
    attachments = email.attachments.all()
    for attachment in attachments:
        logger.info(f"Attachment ID {attachment.id}: Raw filename: {attachment.filename}, " 
                   f"Decoded filename: {attachment.decoded_filename}, "
                   f"File extension: {attachment.file_extension}, "
                   f"Content type: {attachment.content_type}, "
                   f"Is PDF: {attachment.is_pdf}")
        
        # Additional debugging for attachment issues
        if attachment.file_extension.endswith('?='):
            logger.warning(f"Problematic file extension detected for Attachment ID {attachment.id}: {attachment.file_extension}")
            # Try to repair the file extension for this analysis session
            real_extension = os.path.splitext(attachment.decoded_filename.lower().replace('?=', ''))[1]
            logger.info(f"Repaired extension: {real_extension} for Attachment ID {attachment.id}")
    
    success = process_email_attachments(email, request.user)
    
    if success:
        messages.success(request, "Attachments analyzed successfully")
    else:
        messages.error(request, "Failed to analyze attachments")
    
    return redirect('emails:email_detail', email_id=email.id)


@login_required
def add_manual_rule(request, email_id=None):
    """
    Manuel kural ekleme sayfası
    Belirli bir email ile ilişkilendirilebilir (email_id verilirse) veya bağımsız kural olarak eklenebilir
    """
    # Eğer email_id verilmişse, ilgili e-postayı al
    email = None
    if email_id:
        email = get_object_or_404(Email, id=email_id)
    
    # Formun gönderilmesi durumunda
    if request.method == 'POST':
        hotel_id = request.POST.get('hotel_id')
        room_types = request.POST.getlist('room_types')
        market_ids = request.POST.getlist('market_ids')
        contracts = request.POST.getlist('contracts')
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        sale_type = request.POST.get('sale_type')
        
        try:
            # Veri doğrulama
            if not hotel_id or not market_ids or not start_date or not end_date or not sale_type:
                messages.error(request, "Tüm zorunlu alanları doldurun")
                return redirect('emails:add_manual_rule', email_id=email_id)
            
            # Otel ve pazarları al
            hotel = Hotel.objects.get(id=hotel_id)
            selected_markets = Market.objects.filter(id__in=market_ids)
            
            # Tarihleri parse et
            start_date_obj = datetime.strptime(start_date, '%Y-%m-%d').date()
            end_date_obj = datetime.strptime(end_date, '%Y-%m-%d').date()
            
            # Otel adını ayarla
            hotel_name = hotel.juniper_hotel_name
            
            # Oda tipini ayarla
            is_all_rooms = 'all' in room_types
            
            if is_all_rooms:
                room_type = "All Room"
                selected_rooms = []
            else:
                selected_rooms = Room.objects.filter(hotel=hotel, id__in=room_types)
                if not selected_rooms.exists():
                    messages.error(request, "En az bir oda tipi seçin veya 'Tüm Odalar' seçeneğini işaretleyin")
                    return redirect('emails:add_manual_rule', email_id=email_id)
                room_type = ", ".join(selected_rooms.values_list('juniper_room_type', flat=True))
            
            # Kontrat bilgisini ayarla
            selected_contracts = ", ".join(contracts) if contracts else None
            
            # EmailRow oluştur
            with transaction.atomic():
                # Eğer email_id yoksa ve bağımsız kural ekleniyorsa, yeni bir Email oluştur
                if not email:
                    email = Email.objects.create(
                        subject="Manuel Eklenen Kural",
                        sender=f"{request.user.get_full_name()} <{request.user.email}>",
                        recipient="Manuel Ekleme",
                        received_date=timezone.now(),
                        message_id=f"manual-rule-{timezone.now().timestamp()}",
                        body_text="Bu kural manuel olarak eklenmiştir.",
                        status="pending"
                    )
                
                # EmailRow oluştur
                row = EmailRow.objects.create(
                    email=email,
                    hotel_name=hotel_name,
                    room_type=room_type,
                    start_date=start_date_obj,
                    end_date=end_date_obj,
                    sale_type=sale_type,
                    status='pending',
                    juniper_hotel=hotel,
                    selected_contracts=selected_contracts,
                    manually_edited=True,
                    processed_by=request.user,
                    processed_at=timezone.now()
                )
                
                # Pazarları ayarla
                row.markets.set(selected_markets)
                
                # Oda tiplerini ayarla (tüm odalar seçilmemişse)
                if not is_all_rooms and selected_rooms:
                    row.juniper_rooms.set(selected_rooms)
                
                # Log kaydı oluştur
                UserLog.objects.create(
                    user=request.user,
                    action_type='add_manual_rule',
                    email=email,
                    email_row=row,
                    ip_address=request.META.get('REMOTE_ADDR'),
                    details=f"Manuel kural eklendi: {hotel_name} - {room_type} - {start_date} ile {end_date} arası {sale_type}"
                )
                
                # Öğrenme sistemine kaydet
                learn_hotel_matching(row, request.user)
                
                # Pazarları öğren
                if original_market_name and selected_markets.exists():
                    row.original_market_name = original_market_name
                    row.save()
                    learn_market_matching(row, request.user)
                
                # Kontratları öğren
                if selected_contracts:
                    learn_contract_matching(row, request.user)
                
                messages.success(request, "Kural başarıyla eklendi")
                
                # E-posta sayfasına yönlendir
                return redirect('emails:email_detail', email_id=email.id)
                
        except Hotel.DoesNotExist:
            messages.error(request, "Seçilen otel bulunamadı")
        except Market.DoesNotExist:
            messages.error(request, "Seçilen pazar bulunamadı")
        except ValueError as e:
            messages.error(request, f"Geçersiz veri formatı: {str(e)}")
        except Exception as e:
            logger.error(f"Manuel kural eklemede hata: {str(e)}", exc_info=True)
            messages.error(request, f"Bir hata oluştu: {str(e)}")
    
    # Form sayfasını göster
    context = {
        'email': email,
        'hotels': Hotel.objects.all().order_by('juniper_hotel_name'),
        'markets': Market.objects.all().order_by('name'),
    }
    
    return render(request, 'emails/add_manual_rule.html', context)


def learn_room_type_group_matching(email_row, user=None):
    """
    E-postada belirtilen oda tipi ile kullanıcının manuel olarak eşleştirdiği Juniper oda tipleri arasında
    otomatik eşleştirme için oda tipi grubu oluşturur veya günceller.
    
    Eğer e-postadaki oda tipi adı (örn. "Ana Bina Kara Manzaralı") kullanıcı tarafından belirli Juniper oda 
    tipleriyle eşleştirilirse, bu fonksiyon:
    1. Bu oda tipi adıyla otel için bir RoomTypeGroup oluşturur/günceller 
    2. Kullanıcının seçtiği odaları bu gruba varyant olarak ekler
    
    Args:
        email_row (EmailRow): Onaylanan e-posta satırı
        user (User, optional): İşlemi yapan kullanıcı
    
    Returns:
        tuple: (bool, str) - (Başarılı oldu mu, Bilgi mesajı)
    """
    from hotels.models import RoomTypeGroup, RoomTypeVariant, RoomTypeGroupLearning
    import logging
    logger = logging.getLogger(__name__)
    
    # Gerekli verileri kontrol et
    if not email_row.juniper_hotel or not email_row.room_type:
        return False, "Otel veya oda tipi bilgisi eksik"
    
    # Tüm odalar durumunu kontrol et
    is_all_room_type = email_row.room_type.strip().upper() in ['ALL ROOM', 'ALL ROOMS', 'ALL ROOM TYPES', 'TÜM ODALAR']
    if is_all_room_type:
        return False, "Tüm odalar seçildiğinde grup oluşturulmaz"
    
    # Juniper odaları kontrol et
    if not email_row.juniper_rooms.exists():
        return False, "Eşleştirilmiş Juniper odası bulunamadı"
    
    try:
        # E-postadaki oda tipini temizle
        email_room_type = email_row.room_type.strip()
        clean_room_type = email_room_type.upper()
        
        # Otel ve temizlenmiş e-posta oda tipi için mevcut grup kontrolü
        room_group = RoomTypeGroup.objects.filter(
            hotel=email_row.juniper_hotel,
            name__iexact=clean_room_type
        ).first()
        
        # Eğer grup yoksa oluştur
        if not room_group:
            room_group = RoomTypeGroup.objects.create(
                hotel=email_row.juniper_hotel,
                name=clean_room_type
            )
            logger.info(f"Yeni oda tipi grubu oluşturuldu: {clean_room_type} - {email_row.juniper_hotel.juniper_hotel_name}")
        
        # Kullanıcının seçtiği odaları al
        selected_juniper_rooms = list(email_row.juniper_rooms.all())
        
        # Her bir oda için grup varyantı oluştur
        added_variants = 0
        for room in selected_juniper_rooms:
            # Juniper oda adı
            room_name = room.juniper_room_type.strip()
            
            # Bu oda adı zaten varyant olarak eklenmiş mi kontrol et
            variant_exists = RoomTypeVariant.objects.filter(
                group=room_group,
                variant_room_name__iexact=room_name
            ).exists()
            
            if not variant_exists:
                # Varyant oluştur
                RoomTypeVariant.objects.create(
                    group=room_group,
                    variant_room_name=room_name
                )
                added_variants += 1
        
        # Ayrıca RoomTypeGroupLearning kaydı oluştur/güncelle
        room_learning, created = RoomTypeGroupLearning.objects.get_or_create(
            hotel=email_row.juniper_hotel,
            mail_room_type=email_room_type,
            defaults={
                'group': room_group,
                'confidence': 0.8
            }
        )
        
        # Eğer kaydedilmiş bir öğrenme varsa güven puanını güncelle
        if not created:
            # Grubu güncelle (farklı olabilir)
            room_learning.group = room_group
            
            # Güven puanını artır
            room_learning.confidence = min(0.99, room_learning.confidence + 0.1)
            room_learning.frequency += 1
            room_learning.save()
        
        # Sonucu döndür
        if added_variants > 0:
            return True, f"{added_variants} yeni oda varyantı '{clean_room_type}' grubuna eklendi"
        else:
            return True, f"Oda tipi grubu '{clean_room_type}' için varyantlar zaten güncel"
            
    except Exception as e:
        logger.error(f"Oda tipi grubu öğrenme hatası: {str(e)}", exc_info=True)
        return False, f"Hata: {str(e)}"
